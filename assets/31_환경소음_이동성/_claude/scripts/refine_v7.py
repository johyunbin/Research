# -*- coding: utf-8 -*-
# v8 정제7(정합성 감사 반영): 계절 과대주장 통일(FDR 후 봄·가을) + Fig9 캡션·§4.3 단일센서 서사 교정.
import os, re
from datetime import datetime, timezone, timedelta
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = r"T:\00_BACKUP\01_한양대학교\00_연구실\00_프로젝트\★_논문화 프로젝트\31_환경소음_이동성"
SRC = os.path.join(ROOT, "01_논문작업", "Manuscript_v8_20260623_102856.docx")
ts = datetime.now(timezone(timedelta(hours=9))).strftime("%Y%m%d_%H%M%S")
OUT = os.path.join(ROOT, "01_논문작업", f"Manuscript_v8_{ts}.docx")

def setfont(run, size=12, sub=False):
    run.font.size = Pt(size); run.font.name = 'Times New Roman'
    if sub: run.font.subscript = True
    rpr = run._element.get_or_add_rPr(); rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'): rf.set(qn(a), 'Times New Roman')
    rf.set(qn('w:eastAsia'), '바탕')

SUB = [('Leq,24h', ('L', 'eq,24h')), ('LAeq', ('L', 'Aeq')), ('Lday', ('L', 'day')), ('Lnight', ('L', 'night')), ('Leq', ('L', 'eq'))]
SRE = re.compile('(' + '|'.join(re.escape(k) for k, _ in SUB) + ')'); SMAP = dict(SUB)
def replace_para(p, text, size=12):
    for r in list(p.runs): r._element.getparent().remove(r._element)
    for part in SRE.split(text):
        if not part: continue
        if part in SMAP:
            b, sb = SMAP[part]; setfont(p.add_run(b), size); setfont(p.add_run(sb), size, sub=True)
        else:
            setfont(p.add_run(part), size)

doc = Document(SRC)
def find(m):
    h = [p for p in doc.paragraphs if p.text.strip().startswith(m)]
    assert len(h) == 1, f"마커 {len(h)}건: {m[:30]}"
    return h[0]

REPL = [
# 초록: "consistent across seasons" → "positive across seasons"(방향만 단정, 유의성 과대주장 제거)
("How strongly the urban acoustic",
 "How strongly the urban acoustic environment responds to human mobility is a first-order question for "
 "mobility, land-use and noise policy, yet it remains difficult to answer because mobility rarely varies "
 "exogenously and is entangled with weather, land use and time of day. We exploit Korea's graded COVID-19 "
 "social-distancing system, which repeatedly tightened and relaxed activity over 2020-2023, as a natural "
 "experiment, paired with a city-scale low-cost IoT noise network, to estimate a measured-mobility dose-response "
 "for urban noise. We assembled 1,248,794 sensor-days from 1,123 Smart Seoul Data-of-Things (S-DoT) sensors "
 "across 421 administrative neighbourhoods and matched each sensor-day to neighbourhood daytime de-facto "
 "population. Because low-cost sensors carry calibration offsets and multi-year drift, we identify the effect "
 "only from within-sensor variation and from cross-neighbourhood variation within each date (two-way fixed "
 "effects). Daytime noise rises and falls with daytime mobility (beta = +0.65 dB per log-unit; a 30% mobility "
 "reduction corresponds to about 0.23 dB), robust to placebo permutation (p = 0.003) and positive across "
 "seasons but specific to daytime: a nighttime association vanishes under two-way fixed effects. No robust "
 "gradient survives in the long-run spatial cross-section, and a calibrated environmental-noise network shows "
 "that S-DoT under-reads standard LAeq by about 12 dB and that its multi-year decline is a sensor-drift "
 "artifact. The response is statistically robust but modest: mobility-demand management alone cannot deliver "
 "large noise reductions and must be paired with source- and propagation-stage measures. It also corrects "
 "binary-lockdown over-estimation and offers a transferable, drift-aware design guide for IoT-based noise "
 "monitoring in smart cities."),
# §3.2 계절: 명목/FDR 구분 명시
("이 효과의 기능적 위치를 분해하면",
 "이 효과의 기능적 위치를 분해하면(Fig. 3, Table 6, RQ2) 효과는 daytime에 특정된다. 결과변수별로 주간만 유의하고 야간·주야 gap은 비유의했다. 토지이용별로는 상업(+0.34)·혼합(+0.86)·주거(+0.85) 모두 양(+)이었으나 표본 분할로 검정력이 떨어져 개별 유의성은 약했고, 상호작용 검정에서도 토지이용 차이는 유의하지 않았다(효과는 대체로 균일). 평일(+0.49)·주말(+0.64)은 유사했다. 특히 계절별로는 네 계절 모두 같은 양(+) 방향이며(DJF +0.66·MAM +0.69·JJA +0.53·SON +0.76), 겨울·봄·가을이 명목 유의했다. 다만 다중비교(FDR) 보정 후에는 봄(MAM)·가을(SON)이 유의하게 남아(§3.6), 효과가 특정 계절의 아티팩트가 아니라 연중 같은 방향으로 일관됨을 보인다."),
# Fig 9 캡션: single-sensor → extreme neighbourhoods
("Fig. 9. Commercial vs residential",
 "Fig. 9. Commercial vs residential neighbourhoods. (a,c) drift-removed relative noise change; (b,d) mobility vs "
 "noise (Theil-Sen, Spearman vs Pearson). The Pearson correlation is inflated by a few extreme neighbourhoods; "
 "the robust Spearman is approximately zero."),
# §4.1: 네 계절 일관 → 같은 방향(보정 후 봄·가을 유의)
("본 연구는 측정 이동량에 대한 도시소음의 단계적",
 "본 연구는 측정 이동량에 대한 도시소음의 단계적(graded) 용량-반응을 처음으로 정량 추정했다. 핵심 결과는 네 가지로 요약된다. 첫째, 한 동의 주간 이동량과 그 동의 주간 소음 사이에 통계적으로 견고한 정(+)의 용량-반응이 존재하지만 그 크기는 작다(이동량 30% 감소 ≈ 소음 −0.2~0.3 dB). 둘째, 효과는 주간에 한정되며 네 계절 모두 같은 방향으로 나타난다(다중비교 보정 후 봄·가을에서 유의). 야간의 겉보기 효과는 엄밀한 식별에서 시간 교란으로 사라진다. 셋째, 효과는 '같은 날 동 사이의 차이'와 차이의 차이 분석에서만 안정적으로 드러나며, 동별 장기 평균을 비교하는 공간 단면에서는 사라진다. 넷째, 저가 IoT 망의 여러 해에 걸친 절대 수준은 센서 표류와 지역별 회복 차이로 교란되어 시계열·절대 비교를 믿을 수 없고, 오직 센서 내 변화와 동 사이 비교에 기반한 설계만이 유효하다."),
# §4.3: 연중 일관 → 같은 방향, 단일센서 서사 → 극단 동 Pearson
("효과가 나타나는 양상은 그 메커니즘",
 "효과가 나타나는 양상은 그 메커니즘과 부합한다. 효과가 주간에 나타나고 연중 같은 방향을 유지하는 것은, 주간 소음이 상거래·통근·방문 같은 사람의 활동에 직접 연동되는 반면 야간 소음은 고정 설비나 간선교통 같은 배경원이 지배해 이동량 변화에 둔감하기 때문이다. 야간의 겉보기 효과가 엄밀한 식별에서 사라진 사례는, 단순한 시간 추세를 인과로 오인할 위험을 보여 준다. 효과가 토지이용에 비교적 균일했다는 점은, 이동량-소음 반응이 특정 장소 유형에 몰려 있지 않고 도시 전반에 얇게 퍼져 있음을 시사한다. 특히 주목할 것은 동 단위 장기 공간 단면에서 깨끗한 경향이 나타나지 않은 점이다. 이는 효과가 없어서가 아니라, 동별 장기 평균 소음 변화가 국지적 공사·도로 변화·센서별 표류 같은 특이 요인에 가려지기 때문이다. 신호는 같은 날 동 사이의 빠른 변동에 있고, 단순한 공간 지도화로는 드러나지 않는다. 실제로 상업 동에서 잠깐 보였던 양의 상관조차 소수의 극단적인 동에 민감한 Pearson이 만든 착시였고, 극단값에 강한 방법으로 보면 0이었다. 이는 고밀도 IoT 자료의 공간 분석이 얼마나 쉽게 잘못된 결론으로 이어질 수 있는지 보여 주는 사례다."),
]
for m, t in REPL:
    replace_para(find(m), t, 12)
print(f"정합성 수정: {len(REPL)}건")
ab = find("How strongly the urban acoustic").text
print("초록 단어수:", len(ab.split()))
doc.save(OUT)
print("저장:", OUT)

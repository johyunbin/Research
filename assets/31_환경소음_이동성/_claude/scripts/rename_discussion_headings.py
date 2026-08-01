# -*- coding: utf-8 -*-
# Discussion 소제목 개편(사용자 승인 2026-07-27): 템플릿형 → 내용 전달형 + 4.5/4.6 분리.
# EN·KR 정본 동시 적용(소제목은 양쪽 공유 영문 — KR만 남기면 재생성 시 구제목 회귀).
import copy
from docx import Document

RENAME = [
    ("4.1. Summary of findings", "4.1. Principal findings"),
    ("4.2. The magnitude in context", "4.2. Comparison with lockdown-era noise studies"),
    ("4.3. Mechanisms behind the patterns", "4.3. Interpreting the daytime response and the night-time null"),
    ("4.4. Methodological implications", "4.4. Implications for low-cost urban noise sensing"),
    ("4.5. Policy implications and limitations", "4.5. Policy implications"),
]
NEW_HEAD = "4.6. Limitations and future research"

FILES = [
    (r"C:\Users\wh850\Research\assets\31_환경소음_이동성\01_논문작업\Manuscript_EN_20260727_005544.docx",
     "Several limitations qualify these conclusions."),
    (r"C:\Users\wh850\Research\assets\31_환경소음_이동성\01_논문작업\Manuscript_20260726_231342.docx",
     "본 연구에는 몇 가지 한계가"),
]

for path, limit_marker in FILES:
    doc = Document(path)
    n = 0
    head_45 = None
    for p in doc.paragraphs:
        t = p.text.strip()
        for old, new in RENAME:
            if t == old:
                for i, r in enumerate(p.runs):
                    r.text = new if i == 0 else ""
                n += 1
                if new.startswith("4.5."):
                    head_45 = p
    assert n == 5, f"{path}: 제목 치환 {n}건 (기대 5)"
    assert head_45 is not None, "4.5 제목 문단 미확보"

    limits = [p for p in doc.paragraphs if p.text.strip().startswith(limit_marker)]
    assert len(limits) == 1, f"{path}: 한계 문단 {len(limits)}건"
    new_p = copy.deepcopy(head_45._p)
    limits[0]._p.addprevious(new_p)
    # 복제한 제목 문단의 텍스트를 4.6으로
    from docx.text.paragraph import Paragraph
    np = Paragraph(new_p, limits[0]._parent)
    for i, r in enumerate(np.runs):
        r.text = NEW_HEAD if i == 0 else ""
    doc.save(path)
    print(f"OK 5+1 제목 적용: {path}")

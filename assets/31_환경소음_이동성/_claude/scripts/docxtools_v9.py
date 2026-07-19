# -*- coding: utf-8 -*-
# v9 재조립 공용 헬퍼: 문단 탐색/치환(첨자 처리)·블록 삭제/삽입·booktabs 표 빌더·그림 blob 교체.
import os, re, struct
from docx.shared import Pt, Inches, Emu
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.table import Table
from docx.text.paragraph import Paragraph

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
FIGD = os.path.join(ROOT, "_claude", "figures_v2")

# ---- 첨자 렌더링 (기존 refine 계열과 동일 규범: TNR, eastAsia '바탕') ----
SUB = [('Leq,24h', ('L', 'eq,24h')), ('LAeq', ('L', 'Aeq')), ('Lday', ('L', 'day')),
       ('Lnight', ('L', 'night')), ('Leq24', ('L', 'eq,24h')), ('Leq', ('L', 'eq'))]
SRE = re.compile('(' + '|'.join(re.escape(k) for k, _ in SUB) + ')')
SMAP = dict(SUB)

def setfont(run, size=12, sub=False, bold=False, italic=False):
    run.font.size = Pt(size); run.font.name = 'Times New Roman'
    run.bold = bold; run.italic = italic
    if sub: run.font.subscript = True
    rpr = run._element.get_or_add_rPr(); rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'): rf.set(qn(a), 'Times New Roman')
    rf.set(qn('w:eastAsia'), '바탕')

def fill_para(p, text, size=12, bold=False, italic=False):
    """문단을 비우고 첨자 인식 run들로 채움. '\n' = 줄바꿈."""
    for r in list(p.runs): r._element.getparent().remove(r._element)
    lines = text.split("\n")
    for li, line in enumerate(lines):
        if li > 0:
            br = p.add_run(); setfont(br, size); br.add_break()
        for part in SRE.split(line):
            if not part: continue
            if part in SMAP:
                b, sb = SMAP[part]
                setfont(p.add_run(b), size, bold=bold, italic=italic)
                setfont(p.add_run(sb), size, sub=True, bold=bold, italic=italic)
            else:
                setfont(p.add_run(part), size, bold=bold, italic=italic)

def iter_blocks(doc):
    for child in doc.element.body.iterchildren():
        if child.tag == qn('w:p'):
            yield Paragraph(child, doc._body)
        elif child.tag == qn('w:tbl'):
            yield Table(child, doc._body)

def find(doc, prefix, unique=True):
    hits = [p for p in doc.paragraphs if p.text.strip().startswith(prefix)]
    if unique:
        assert len(hits) == 1, f"마커 {len(hits)}건: {prefix[:40]}"
        return hits[0]
    return hits

def delete_el(obj):
    el = obj._element if hasattr(obj, "_element") else obj
    el.getparent().remove(el)

def para_after(doc, anchor_el):
    """anchor 요소 뒤에 새 문단 삽입."""
    p = OxmlElement('w:p')
    anchor_el.addnext(p)
    return Paragraph(p, doc._body)

# ---- booktabs 표 ----
def _border(el_name, sz):
    e = OxmlElement(el_name)
    e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), str(sz))
    e.set(qn('w:space'), '0'); e.set(qn('w:color'), '000000')
    return e

def _none_border(el_name):
    e = OxmlElement(el_name); e.set(qn('w:val'), 'none')
    e.set(qn('w:sz'), '0'); e.set(qn('w:space'), '0')
    return e

def style_booktabs(tbl, header_rows=1):
    """상단·하단 굵은 rule + 헤더 아래 가는 rule, 세로선·내부선 제거."""
    tblPr = tbl._tbl.tblPr
    for old in tblPr.findall(qn('w:tblBorders')): tblPr.remove(old)
    borders = OxmlElement('w:tblBorders')
    borders.append(_border('w:top', 12)); borders.append(_border('w:bottom', 12))
    for name in ('w:left', 'w:right', 'w:insideH', 'w:insideV'):
        borders.append(_none_border(name))
    tblPr.append(borders)
    # 헤더 마지막 행 아래 가는 선
    for c in tbl.rows[header_rows - 1].cells:
        tcPr = c._tc.get_or_add_tcPr()
        for old in tcPr.findall(qn('w:tcBorders')): tcPr.remove(old)
        tcB = OxmlElement('w:tcBorders'); tcB.append(_border('w:bottom', 6))
        tcPr.append(tcB)

def build_table(doc, rows, widths, aligns=None, font=9, header_rows=1, anchor_el=None):
    """rows[0..header_rows-1]=헤더(bold). widths=inch 리스트. aligns='l/c/r' 문자열 리스트.
    anchor_el 뒤로 이동 삽입."""
    ncol = len(rows[0])
    tbl = doc.add_table(rows=len(rows), cols=ncol)
    tbl.style = None
    tbl.autofit = False
    AL_MAP = {"l": AL.LEFT, "c": AL.CENTER, "r": AL.RIGHT}
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.rows[i].cells[j]
            cell.width = Inches(widths[j])
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = AL_MAP.get((aligns[j] if aligns else "l"), AL.LEFT)
            fill_para(p, str(val), size=font, bold=(i < header_rows))
    style_booktabs(tbl, header_rows=header_rows)
    if anchor_el is not None:
        el = tbl._element; el.getparent().remove(el)
        anchor_el.addnext(el)
    return tbl

def caption_para(doc, anchor_el, text, size=12):
    """anchor 뒤에 캡션 문단(TNR, 좌정렬) 삽입 후 반환."""
    p = para_after(doc, anchor_el)
    fill_para(p, text, size=size)
    return p

# ---- 그림 blob 교체 ----
def png_size(blob):
    assert blob[:8] == b"\x89PNG\r\n\x1a\n"
    return struct.unpack(">II", blob[16:24])

def replace_image(doc, shape, png_path, width_in=6.5):
    blob = open(png_path, "rb").read()
    wpx, hpx = png_size(blob)
    rId = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
    doc.part.related_parts[rId]._blob = blob
    shape.width = Inches(width_in)
    shape.height = Emu(int(round(Inches(width_in) * hpx / wpx)))
    return wpx, hpx

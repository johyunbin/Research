# -*- coding: utf-8 -*-
# v8 정제8: §1.5 기여 문장 자연화 + RQ 라벨 복원 + §3.5에 '검교정망 상승=용량반응과 같은 방향(corroboration)·다년추세 무관(insulation)' 보강.
import os, re
from datetime import datetime, timezone, timedelta
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = r"T:\00_BACKUP\01_한양대학교\00_연구실\00_프로젝트\★_논문화 프로젝트\31_환경소음_이동성"
SRC = os.path.join(ROOT, "01_논문작업", "Manuscript_v8_20260623_104531.docx")
ts = datetime.now(timezone(timedelta(hours=9))).strftime("%Y%m%d_%H%M%S")
OUT = os.path.join(ROOT, "01_논문작업", f"Manuscript_v8_{ts}.docx")

def setfont(run, size=12, sub=False):
    run.font.size = Pt(size); run.font.name = 'Times New Roman'
    if sub: run.font.subscript = True
    rpr = run._element.get_or_add_rPr(); rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'): rf.set(qn(a), 'Times New Roman')
    rf.set(qn('w:eastAsia'), '바탕')
SUB = [('Leq,24h', ('L', 'eq,24h')), ('LAeq', ('L', 'Aeq')), ('Lday', ('L', 'day')), ('Lnight', ('L', 'night')), ('Leq', ('L', 'eq'))]
SRE = re.compile('(' + '|'.join(re.escape(k) for k, _ in SUB) + ')'); SMAP = dict(SUB)
def replace_para(p, text, size=12):
    for r in list(p.runs): r._element.getparent().remove(r._element)
    for part in SRE.split(text):
        if not part: continue
        if part in SMAP:
            b, sb = SMAP[part]; setfont(p.add_run(b), size); setfont(p.add_run(sb), size, sub=True)
        else:
            setfont(p.add_run(part), size)

doc = Document(SRC)
def find(m):
    h = [p for p in doc.paragraphs if p.text.strip().startswith(m)]
    assert len(h) == 1, f"마커 {len(h)}건: {m[:30]}"
    return h[0]

REPL = [
# RQ 라벨 복원
("측정 이동량 감소는 도시소음을 얼마나 낮추는가",
 "RQ1. 측정 이동량 감소는 도시소음을 얼마나 낮추는가(용량-반응 기울기)?"),
("그 효과는 토지이용·주야·요일·계절에",
 "RQ2. 그 효과는 토지이용·주야·요일·계절에 따라 어떻게 다른가?"),
("거리두기 전환에서 소음은 어떻게 반등",
 "RQ3. 거리두기 전환에서 소음은 어떻게 반등하며 어떤 공간패턴을 갖는가?"),
# 기여 문장 자연화(병렬·동사종결)
("기여는 네 가지다",
 "본 연구의 기여는 네 가지다. 첫째, 측정된 이동량에 대한 도시소음의 단계적(graded) 용량-반응을 처음으로 정량 추정한다. 둘째, 약 1,100개 지점의 고밀도 IoT 센서망과 행정동 생활인구를 결합해 도시 미시공간 해상도에서 그 효과를 식별한다. 셋째, 서구 도시에 편중되어 온 기존 문헌을 서울이라는 고밀도 동아시아 메가시티로 확장한다. 넷째, 저가 IoT 소음망의 검교정·드리프트 한계를 정량적으로 드러내고, 절대 수준 대신 센서 내 상대변화와 '같은 날 동 사이 비교'에 기반해야 한다는 식별 설계 원칙을 제시한다."),
# §3.5: corroboration + insulation 보강
("이 하락이 실제 소음 변화가 아니라",
 "이 하락이 실제 소음 변화가 아니라 센서 드리프트임을, 검교정된 공식 환경소음 측정망(국가소음정보시스템)과 대조해 확인했다. 먼저 절대레벨이 크게 어긋난다. 검교정망 인근의 S-DoT는 표준 LAeq보다 주간 평균 11.7 dB, 도로변에서는 약 16 dB 낮게 읽혀(Fig. 10b), S-DoT 절대값을 표준 소음도로 해석할 수 없음을 정량적으로 보여 준다. 반면 검교정망 자체의 다년 추세는 안정적이거나 오히려 상승한다(Fig. 10a). 일별 자동측정망(도로 9점)은 2020→2023에 +0.04 dB로 사실상 일정했고, 분기 수동측정망은 주거·일반지역 91점이 +1.93 dB, 도로 60점이 +1.91 dB 상승했다(별도의 도로교통 4점 상시 LAeq도 같은 방향이다). 절대레벨의 치우침은 추세 차분에서 상쇄되므로, 검교정망이 안정·상승하는 동안 S-DoT만 2.3 dB 하락했다는 사실은 그 하락이 실제 환경 변화가 아니라 센서의 하향 드리프트임을 뜻한다. 특히 S-DoT 하락이 집중된 비(非)도로 정온지역에서 검교정 주거망이 오히려 상승했다는 점은 그 하락이 실재가 아님을 직접 보여 준다. 한편 이 검교정망의 상승은 우리 용량-반응과 같은 방향임에 주목할 만하다. 검교정 소음은 이동량이 가장 크게 줄었던 2020년에 가장 낮았다가 활동이 회복되며 2023년까지 올라갔는데, 이는 '이동량이 늘면 소음도 는다'는 우리 추정과 부합한다. 즉 다년 절대 수준에서 신뢰할 수 있는 검교정망은 오히려 우리 결과의 방향을 뒷받침하며, 반대로 움직인 것은 드리프트에 오염된 S-DoT의 다년 추세뿐이다. 무엇보다 우리의 용량-반응은 다년 비교가 아니라 '같은 날 동 사이의 차이'로 식별되며(날짜 고정효과가 다년·계절 변동을 모두 흡수한다), S-DoT의 다년 절대 추세를 전혀 사용하지 않는다. 이것이 우리가 절대·시계열 비교 대신 센서 내 상대변화와 '같은 날 동 사이의 차이'(M2)에 의존하는 이유다."),
]
for m, t in REPL:
    replace_para(find(m), t, 12)
print(f"수정: {len(REPL)}건")
doc.save(OUT)
print("저장:", OUT)

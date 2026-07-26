# -*- coding: utf-8 -*-
# v8 정제5(major revision): 정합성 수정 + §3.6 강건성 신설(Fig11·Table7) + 표준섹션 + 결론확장 + SCS 리프레이밍 + 소제목 볼드.
import os, re, struct
from datetime import datetime, timezone, timedelta
from docx import Document
from docx.shared import Pt, Emu, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = r"T:\00_BACKUP\01_한양대학교\00_연구실\00_프로젝트\★_논문화 프로젝트\31_환경소음_이동성"
FIGD = os.path.join(ROOT, "_claude", "figures")
SRC = os.path.join(ROOT, "01_논문작업", "Manuscript_v8_20260622_012515.docx")
ts = datetime.now(timezone(timedelta(hours=9))).strftime("%Y%m%d_%H%M%S")
OUT = os.path.join(ROOT, "01_논문작업", f"Manuscript_v8_{ts}.docx")

def setfont(run, size=12, bold=False, italic=False, sub=False):
    run.font.size = Pt(size); run.bold = bold; run.italic = italic; run.font.name = 'Times New Roman'
    if sub: run.font.subscript = True
    rpr = run._element.get_or_add_rPr(); rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'): rf.set(qn(a), 'Times New Roman')
    rf.set(qn('w:eastAsia'), '바탕')

SUB = [('Leq,24h', ('L', 'eq,24h')), ('LAeq', ('L', 'Aeq')), ('Lday', ('L', 'day')),
       ('Lnight', ('L', 'night')), ('Leq', ('L', 'eq'))]
SRE = re.compile('(' + '|'.join(re.escape(k) for k, _ in SUB) + ')'); SMAP = dict(SUB)

def emit(p, text, size=12, bold=False, italic=False):
    for part in SRE.split(text):
        if not part: continue
        if part in SMAP:
            b, sb = SMAP[part]; setfont(p.add_run(b), size, bold, italic); setfont(p.add_run(sb), size, bold, italic, sub=True)
        else:
            setfont(p.add_run(part), size, bold, italic)

def replace_para(p, text, size=12, bold=False):
    for r in list(p.runs): r._element.getparent().remove(r._element)
    emit(p, text, size, bold)

doc = Document(SRC)
paras = doc.paragraphs
def find(marker):
    hits = [p for p in doc.paragraphs if p.text.strip().startswith(marker)]
    assert len(hits) == 1, f"마커 {len(hits)}건: {marker[:34]}"
    return hits[0]

# ===== ① 본문 정합성 수정 (마커 기반 교체) =====
ABS = ("How strongly the urban acoustic environment responds to human mobility is a first-order question for "
       "mobility, land-use and noise policy, yet it remains difficult to answer because mobility rarely varies "
       "exogenously and is entangled with weather, land use and time of day. We exploit Korea's graded COVID-19 "
       "social-distancing system, which repeatedly tightened and relaxed activity over 2020-2023, as a natural "
       "experiment, paired with a city-scale low-cost IoT noise network, to estimate a measured-mobility dose-response "
       "for urban noise. We assembled 1,248,794 sensor-days from 1,123 Smart Seoul Data-of-Things (S-DoT) sensors "
       "across 421 administrative neighbourhoods and matched each sensor-day to neighbourhood daytime de-facto "
       "population. Because low-cost sensors carry calibration offsets and multi-year drift, we identify the effect "
       "only from within-sensor variation and, in the headline specification, from cross-neighbourhood variation "
       "within each date (two-way fixed effects). Daytime noise rises and falls with daytime mobility (beta = +0.65 dB "
       "per log-unit; a 30% mobility reduction corresponds to about 0.23 dB), robust to placebo permutation "
       "(p = 0.003) and consistent across seasons but specific to daytime: a nighttime association "
       "vanishes under two-way fixed effects. No robust gradient survives in the long-run spatial cross-section, and "
       "official roadside LAeq stations confirm the network tracks real trends while exposing that multi-year absolute "
       "levels are confounded. The response is statistically robust but modest: mobility-demand management alone "
       "cannot deliver large noise reductions and must be paired with source- and propagation-stage measures. It "
       "also corrects binary-lockdown over-estimation and offers a transferable, drift-aware design guide for "
       "IoT-based noise monitoring in smart cities.")
replace_para(find("How strongly the urban acoustic"), ABS, 12)
print("초록 단어수:", len(ABS.split()))

# Fig 1 캡션: 1,165(지오코딩) vs 1,123(분석패널) 명시
replace_para(find("Fig. 1. Study area and sensor network"),
 "Fig. 1. Study area and sensor network. (a) 1,165 geolocated S-DoT noise sensors (1,123 enter the analysis panel "
 "after quality control) across 421 administrative neighbourhoods (dong) of Seoul, colored by activity type "
 "(daytime/nighttime living-population ratio: commercial, mixed, residential). (b) Baseline daytime noise level "
 "(post-lift mean Lday, 2022-07 to 2023-12) by neighbourhood.", 12)

# §2.3: 424개 동 중 421개가 분석대상임을 명시
replace_para(find("이동량 노출은 서울 '생활인구"),
 "이동량 노출은 서울 '생활인구(living / de-facto population)'로 측정했다. 이는 거주 인구가 아니라 모바일 통신 신호로 추정한, 특정 시점에 각 행정동에 실제로 체류하는 인구수다(Table 2). 시간별 체류수(stock)를 일 평균(및 주간 06-21 / 야간 22-05 평균)으로 집약해 424개 동 × 1,461일 = 619,464 동-일을 얻었다(결측 0). 이 가운데 S-DoT 센서가 배정된 421개 동이 분석 대상이 된다. dose 구성에서 한 가지 통찰이 중요하다. 도시 전체 생활인구 총량은 거의 보존된다(사람들이 외출을 줄여도 자택 체류로 계수되기 때문). 즉 거리두기의 신호는 총량이 아니라 공간 재분포에 있으며, 상업·업무 동은 주간 체류가 20-35% 비워지고 주거 동은 최대 37% 증가한다. 따라서 우리는 dose를 동의 자기기준 대비 상대변화로 정의한다. 곧 lp = log(동의 일별 생활인구 / 동의 post-lift 정상기 평균)이다. 이렇게 하면 dose는 각 동이 자기 정상에서 얼마나 벗어났는가를 나타내며, 동·센서별 절대수준 차이에 영향받지 않는다.", 12)

# §2.4: 표본 감쇄 워터폴
replace_para(find("센서 위치정보에는 도로명주소만"),
 "센서 위치정보에는 도로명주소만 있고 행정동 이름이 없어, 각 센서를 그 좌표가 어느 행정동 경계 안에 들어가는지로 동에 배정했다(점-다각형 포함 판정). 행정동 경계는 공개 GeoJSON(vuski/admdongkor, 2022-01 버전)을 사용했다. 한편 생활인구 자료와 경계 자료의 행정동 코드 체계가 일부 지역(강북·강동의 행정구역 개편 동)에서 서로 다르게 매겨져 있었는데, 동 이름을 기준으로 짝지어 바로잡았다. 그 결과 1,170개 센서 중 1,165개(99.6%)가 동에 배정되었고, 주소상의 자치구와 배정된 동의 자치구가 일치하는 비율도 99.6%로 매우 높았다. 이후 소음 QC(하루 12시간 이상·20–95 dB)와 일별 이동량 결합을 거친 최종 분석패널은 1,123개 센서·421개 동(1,248,794 sensor-days)이며, 주력 양방향 고정효과 모형은 주간 이동량이 결측이 아닌 1,122개 센서·420개 동(1,247,546 sensor-days)을 사용한다.", 12)

# §2.6 [클러스터 명세를 코드와 일치]
replace_para(find("저가 S-DoT 센서는 정밀 검교정을 거치지 않아"),
 "저가 S-DoT 센서는 정밀 검교정을 거치지 않아, 똑같은 소리를 들려주어도 센서마다 일정량 높거나 낮게 기록하는 고유의 치우침(상수 오프셋)을 갖는다. 이 치우침의 크기는 센서마다 다르고 알 수 없으므로, 서로 다른 두 센서의 절대 데시벨을 맞비교하는 것은 의미가 없다. 그러나 한 센서가 시간에 따라 얼마나 달라졌는지(예컨대 오늘이 그 센서의 평소보다 얼마나 큰지)만 보면, 그 센서에 늘 똑같이 들어 있는 치우침은 빼는 과정에서 저절로 지워진다. 그래서 우리의 모든 분석은 절대 수준이 아니라 각 센서의 자기 대비 변화(센서 내 상대변화)에 기반한다. 통계적으로 이는 고정효과(fixed-effects) 모형으로 구현된다. 각 센서의 전체 평균을 빼 줌으로써 그 센서 고유의 치우침을 제거하는 것이다(125만 건의 자료에 센서별 더미변수를 일일이 넣는 대신, 집단 평균을 차감하는 수학적으로 동등한 방법을 썼다). 또한 같은 군집에서 나온 관측치들은 서로 닮아 있을 수 있으므로 표준오차를 군집에 강건하게(cluster-robust) 보정했다. 센서 고정효과 모형은 센서 기준, 양방향 고정효과 모형은 동 기준으로 군집화했다 [48]. 전체 식별 논리는 Fig. 2에 요약했다.", 12)

# §2.6 보조분석 [M6 강건성 + 소프트웨어 버전]
replace_para(find("이 두 모형 위에 네 가지 보조 분석을 더했다"),
 "이 두 모형 위에 보조 분석을 더했다. (M3) 기능 분할: M2를 결과변수(주간·야간·주야 차이), 토지이용(동의 주간 대비 야간 인구비를 3분위로 나눈 상업·혼합·주거), 평일/주말, 계절별로 따로 추정해 효과가 어디서 나타나는지 본다. (M4) 차이의 차이(difference-in-differences): 이동량이 크게 줄어든 동(주로 상업지구)과 거의 줄지 않은 동(주로 주거지구)을 나눠, 두 그룹의 소음 변화를 매주 비교한다. 두 그룹은 같은 도시·같은 시기를 공유하므로 날씨·계절·센서 표류 같은 공통 요인은 양쪽에 똑같이 작용해 그 차이를 보면 서로 지워지고, 순수하게 이동량 차이에서 비롯된 소음 차이만 남는다. (M5) 검증: 같은 위치 근처의 S-DoT와 공식 도로교통 4점 상시측정소(표준 LAeq)의 연도별 추세를 맞대어, 센서 표류와 2023년 자료구조 변경의 영향을 점검한다. (M6) 강건성·위약 검정: 식별이 기계적 산물이 아님을 확인하기 위해 ① 같은 날짜 안에서 이동량 dose를 동들 사이에 무작위로 재배치한 placebo 순열검정(300회), ② 이동량이 인과적으로 만들 수 없는 결과변수(기온)에 대한 위약 회귀, ③ 이동량 2차항을 넣은 비선형성 점검을 수행했고, 분할추정(M3)의 다중비교는 Benjamini–Hochberg FDR로 보정했다(§3.6). 끝으로 동 단위 공간 분석에서는 소수의 극단적인 동에 결과가 휘둘리지 않도록 일반 회귀·상관 대신 극단값에 강한(robust) 방법(Theil-Sen 회귀와 Spearman 순위상관)을 쓰고, 추정이 불안정한 단일 센서 동을 제외해 센서가 2개 이상인 동만 사용했다. 모든 분석은 Python 3.13(pandas·NumPy·SciPy·statsmodels)으로 수행했으며, 양방향 고정효과는 센서·날짜 평균을 번갈아 차감하는 반복 within-변환으로, 극단값에 강한 추정은 SciPy의 Theil-Sen·Spearman으로 계산했다.", 12)

# §3.3 ΔLAeq -> ΔLday
replace_para(find("이 기능적 차등은 시간축에서도"),
 "이 기능적 차등은 시간축에서도 뚜렷하다(Fig. 6, RQ3). 상업 동의 주간 이동량은 제한기 내내 기준선 아래로, 주거 동은 기준선 위로 벌어졌다가 2022년 해제 후 수렴한다(수도권 4단계기 중앙값 상업 0.98 vs 주거 1.05). 드리프트를 제거한 상대 소음(동 그룹평균 − 도시평균)에서도 상업 동이 제한기에 도시평균보다 상대적으로 조용한 경향이 관찰된다. 거리두기 전환의 동역학은 차이의 차이(difference-in-differences)로 본다(Fig. 7). 고영향 동과 저영향 동의 주별 ΔLday 차이는 제한기에 음(−)으로 기울고 정상기에 0으로 수렴했으며, 주별 (이동량 차이 vs 소음 차이) 상관은 Pearson +0.44 · Spearman +0.46으로 일치했다.", 12)

# §3.4 단일센서 서사 교정 + '전체' 표본 명시
replace_para(find("반면 동을 하나의 점으로 보고"),
 "반면 동을 하나의 점으로 보고 거리두기 기간 전체의 평균을 비교하는 '장기 공간 단면'에서는 이동량과 소음의 뚜렷한 경향이 나타나지 않았다(Fig. 8). 분석 대상(센서 2개 이상) 동 전체에서, 극단값에 강한 순위상관으로 보면 (이동량 감소 vs 도시평균 대비 소음변화)의 관계는 사실상 0이었다(Spearman ρ≈−0.05). 상업 동만 보면 일반 상관계수(Pearson r=+0.38)가 마치 관계가 있는 듯 보였지만, 이 Pearson 값은 소수의 극단적인 동에 민감해 부풀려진 것이다. 같은 동들을 극단값에 강한 방법으로 다시 보면 관계는 0에 가까웠다(Spearman ρ=+0.08, Theil-Sen 기울기≈0; Fig. 9). 정리하면, 이동량-소음 신호는 '같은 날 동 사이의 차이'(M2·M4)에서만 안정적으로 나타나며, 동별 장기 평균을 단순히 지도에 칠하는 방식으로는 잡히지 않는다.", 12)

# §3.5 ΔLAeq -> ΔLday + City Hall 제외 명시
replace_para(find("각 센서의 자기 대비 변화(ΔLAeq"),
 "각 센서의 자기 대비 변화(ΔLday, 정상기 기준)를 그대로 그려 보면 언뜻 모순처럼 보이는 결과가 나온다. 규제가 가장 강했던 시기의 주간 소음이 오히려 정상기보다 높게 나오는 것이다(+0.85 vs +0.34 dB). 진단해 보니 이는 코로나 효과가 아니라 여러 해에 걸친 절대 수준의 변동(표류) 때문이었다(Fig. 10). 4년 내내 가동된 842개 센서의 연평균 Leq,24h는 49.4(2020)→47.6→47.2→47.1 dB(2023)로 해마다 꾸준히 낮아졌고(85% 센서가 하락 추세), 이 하락이 2020-21년을 상대적으로 '시끄럽게' 보이게 만든 것이다. 다만 2023년 자료구조 변경 지점에서는 수준 도약이 없었다(+0.02 dB). 그런데 이 하락은 모든 센서에 똑같이 일어난 단순한 기계 노후가 아니다. 공식 측정소가 있는 도로변 4지점에서 가까운 S-DoT의 2021→2023 추세는 City Hall +1.7·Sinsa +2.5·Sinchon ≈0·Seongsu −0.3 dB로 하락하지 않았고, 공식 LAeq 추세(각 +7.2[공사 추정]·+0.4·−0.4·+1.5)와 같은 방향이었다(공사로 추정되는 City Hall을 제외해도 나머지 세 지점에서 결론은 같다). 즉 도로변 S-DoT는 공식 측정망을 잘 따라가 신뢰할 만하지만, 도시 전체 평균의 하락(2021→2023 −0.5 dB)은 도로에서 벗어난 조용한 지역에 집중된 위치별 현상이어서, 여러 해에 걸친 절대 수준 비교 자체를 믿기 어렵게 만든다. 이것이 우리가 절대·시계열 비교 대신 '같은 날 동 사이의 차이'(M2)에 의존하는 이유다.", 12)

# §4.1 제목 변경
replace_para(find("4.1. Principal findings"), "4.1. Summary of findings", 12, bold=True)

# ===== ② 소제목 볼드 =====
nb = 0
for p in doc.paragraphs:
    if re.match(r'^\d+\.\d+\.\s', p.text.strip()):
        for r in p.runs: r.bold = True
        nb += 1
print("소제목 볼드:", nb)

# ===== ③ §3.6 강건성 신설 (Fig 11 + Table 7) — Discussion 앞에 삽입 =====
disc = find("4. Discussion")
TBLSTYLE = doc.tables[0].style
def para_before(ref, text=None, size=12, bold=False, align=None):
    p = doc.add_paragraph(); ref._p.addprevious(p._p)
    if text is not None: emit(p, text, size, bold)
    if align is not None: p.alignment = align
    return p
para_before(disc, "3.6. Robustness and falsification", 12, bold=True)
para_before(disc,
 "주력 추정치가 기계적 산물이 아님을 네 가지로 확인했다(Fig. 11, Table 7). 첫째, 같은 날짜 안에서 이동량 dose를 동들 사이에 무작위로 재배치한 placebo 순열검정(300회)에서 위약 계수는 0 근처에 좁게 분포했고(평균 −0.00, SD 0.02), 실제 추정치 +0.65는 그 분포를 완전히 벗어났다(양측 순열 p=0.003). 둘째, 이동량이 인과적으로 만들 수 없는 결과변수인 기온에 대한 위약 회귀는 사실상 0이었다(β=+0.00, p=0.75). 즉 이동량 효과가 날씨를 대리한 것이 아니다. 셋째, 이동량 2차항은 유의하지 않아(β²=−0.59, p=0.44) 선형 용량-반응 근사가 타당했다(Fig. 11b). 넷째, 분할추정(Table 6)의 12개 검정을 Benjamini–Hochberg FDR로 보정해도 주간 효과와 봄(MAM)·가을(SON) 계절 효과는 유의하게 남았다(FDR<0.05). 이로써 '같은 날 동 사이의 차이'라는 식별이 신뢰할 만함을 확인한다.", 12)
# Fig 11 이미지
imgp = para_before(disc, align=WD_ALIGN_PARAGRAPH.CENTER)
imgp.add_run().add_picture(os.path.join(FIGD, "fig_robustness.png"), width=Inches(6.0))
para_before(disc,
 "Fig. 11. Robustness checks. (a) Placebo permutation test: distribution of the dose-response coefficient when "
 "mobility is randomly reshuffled across neighbourhoods within each date (300 shuffles); the actual estimate (+0.65) "
 "lies far outside the null (permutation p = 0.003). (b) Binned dose-response after two-way demeaning: decile means "
 "fall close to the linear fit, supporting a linear approximation.", 11)
# Table 7 캡션 + 표
para_before(disc, "Table 7. Robustness and falsification checks for the headline two-way fixed-effects dose-response.", 11, bold=False)
T7 = [("Check", "Result"),
      ("Placebo permutation (300 within-date shuffles)", "Null beta = -0.00 +/- 0.02; actual = +0.65, permutation p = 0.003"),
      ("Placebo outcome (temperature on mobility)", "beta = +0.00 (p = 0.75); approximately 0, as expected"),
      ("Nonlinearity (quadratic mobility term)", "quadratic beta = -0.59 (p = 0.44), not significant; linear adequate"),
      ("Multiple comparison (BH-FDR, 12 segment tests)", "3 of 4 nominal effects survive (daytime, MAM, SON; FDR < 0.05)")]
tbl7 = doc.add_table(rows=len(T7), cols=2); tbl7.style = TBLSTYLE
disc._p.addprevious(tbl7._tbl)
for ri, (a, b) in enumerate(T7):
    for ci, txt in enumerate((a, b)):
        cell = tbl7.rows[ri].cells[ci]
        for r in list(cell.paragraphs[0].runs): r._element.getparent().remove(r._element)
        emit(cell.paragraphs[0], txt, 10, bold=(ri == 0))

# ===== ④ 결론 확장(2문단) =====
conc = find("우리는 한국의 단계적 거리두기를 자연실험으로 삼아")
replace_para(conc,
 "우리는 한국의 단계적 거리두기를 자연실험으로 삼아, 도시 규모의 IoT 센서망에서 측정 이동량에 대한 도시소음의 graded 용량-반응을 처음으로 정량 추정했다. 그 효과는 통계적으로 견고하지만 작다(이동량 30% 감소 ≈ 주간 소음 0.23 dB, 50% 감소 ≈ 0.45 dB). 효과는 주간에 한정되며 네 계절에 걸쳐 같은 방향·크기로 나타나고(다중비교 보정 후 봄·가을에서 유의), 오직 '같은 날 동 사이의 차이'라는 엄밀한 비교에서만 드러난다. 동별 장기 평균을 비교하는 공간 단면에서는 신호가 사라지고, 저가 IoT 망의 여러 해에 걸친 절대 수준은 센서 표류로 교란된다. 정책적으로 이는, 이동수요를 줄이는 도시정책(차 없는 거리·15분 도시)이 그 자체로는 도시소음을 크게 낮추지 못하며, 소음 목표 달성에는 노면 포장·저소음 차량·차폐 같은 음원·전파 단계 개입이 반드시 병행되어야 함을 실측으로 보여 준다. 또한 봉쇄가 도시를 수 dB 조용하게 만들었다는 통념이 소수 핫스팟의 극단치를 일반화한 과대평가임을 바로잡아, 소음정책의 비용-편익 기대치를 보정한다.", 12)
para2 = doc.add_paragraph(); conc._p.addnext(para2._p)
emit(para2,
 "방법론적으로, 본 연구는 검교정되지 않은 고밀도 IoT 소음망을 도시 환경연구에 신뢰성 있게 쓰는 절차(센서 내 상대변화, 양방향 고정효과, 극단값에 강한 통계, 공식망 교차검증)를 제시하고, 무작위화추론 placebo로 그 식별을 검증했다. 절대 수준을 그대로 정책지표로 삼으면 다년 표류에 오도되므로, 스마트시티 소음 모니터링은 within-sensor 변화량 기반 설계를 채택해야 한다. 향후 과제는 세 가지다. 첫째, 정류장별 승하차 같은 더 직접적인 교통 노출을 동 단위로 결합해 효과를 정밀화하고, 둘째, 검교정된 참조망과 결합해 절대 수준을 복원하며, 셋째, 본 설계를 SONYC·WASN·NoiseCapture 등 다른 저가 소음망과 도쿄·홍콩·싱가포르 같은 고밀도 동아시아 도시로 확장해 일반화를 검증하는 것이다.", 12)

# ===== ⑤ 표준 섹션 삽입 (Data Availability 앞) =====
dav = find("Data Availability Statement")
STD = [("CRediT authorship contribution statement",
        "Hyun In Jo: Conceptualization, Methodology, Software, Formal analysis, Data curation, Writing – original draft, Writing – review & editing, Visualization."),
       ("Declaration of competing interest",
        "The author declares no competing financial interests or personal relationships that could have appeared to influence the work reported in this paper."),
       ("Funding",
        "This research received no specific grant from any funding agency in the public, commercial, or not-for-profit sectors. [저자 확인 후 수정]"),
       ("Ethics",
        "This study analysed only publicly available, aggregated and de-identified data and did not involve human participants directly; institutional review board approval was therefore not required.")]
for h, b in STD:
    para_before(dav, h, 12, bold=True)
    para_before(dav, b, 12)

doc.save(OUT)
print("\n저장:", OUT)

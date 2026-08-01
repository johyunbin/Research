# -*- coding: utf-8 -*-
# 투고 직전 Codex 최종 게이트 High 반영 (2026-08-01):
# CL① 수신 블록 = 직함 생략 안전형(이름 3인+저널) ② 드리프트 표현 hedge ③ 재현성 과장 완화 ④ 서명부 귀속 명확화
# MS⑤ §1.3 "pandemic itself has ended"→"pandemic emergency has passed" ⑥ 결론 They 선행사 교정 (EN+KR)
import os, sys
from docx import Document
sys.path.insert(0, os.path.dirname(__file__))
import docxtools_v9 as T

def para_size(p, default=11.0):
    for r in p.runs:
        if r.font.size is not None:
            return r.font.size.pt
    st = p.style
    while st is not None:
        if st.font.size is not None:
            return st.font.size.pt
        st = st.base_style
    return default

CL = r"C:\Users\wh850\Research\assets\31_환경소음_이동성\01_논문작업\Cover_Letter_SCS_20260727_065746.docx"
EN = r"C:\Users\wh850\Research\assets\31_환경소음_이동성\01_논문작업\Manuscript_EN_20260727_005544.docx"
KR = r"C:\Users\wh850\Research\assets\31_환경소음_이동성\01_논문작업\Manuscript_20260726_231342.docx"

# ---------- Cover letter ----------
doc = Document(CL)
def set_text(p, new):
    for i, r in enumerate(p.runs):
        r.text = new if i == 0 else ""
n = 0
drop = []
for p in doc.paragraphs:
    t = p.text.strip()
    if t == "Prof. Enrico Fabrizio, PhD":
        set_text(p, "Prof. Enrico Fabrizio"); n += 1
    elif t == "Prof. Fariborz Haghighat, Ph.D., P.Eng., Fellow ASHRAE, Fellow ISIAQ":
        set_text(p, "Prof. Fariborz Haghighat"); n += 1
    elif t == "Editors-in-Chief":
        drop.append(p); n += 1
    elif t == "On behalf of all co-authors:":
        set_text(p, "Corresponding author, on behalf of all authors"); n += 1
    elif t == "Hyun In Jo":
        drop.append(p); n += 1
    elif "uncovers a 2.3 dB multi-year drift" in p.text:
        new = p.text.replace(
            "The analysis also uncovers a 2.3 dB multi-year drift in the uncalibrated network, "
            "checked against the official calibrated stations, and turns this into",
            "The analysis also finds a 2.3 dB multi-year decline in the uncalibrated network that, "
            "checked against the official calibrated stations, is consistent with sensor drift, "
            "and turns this into")
        assert new != p.text
        T.fill_para(p, new, size=para_size(p, 12)); n += 1
    elif "reproducible from open sources" in p.text:
        new = p.text.replace(
            "All raw data are public, and the full analysis is reproducible from open sources.",
            "All raw data are openly available from public portals.")
        assert new != p.text
        T.fill_para(p, new, size=para_size(p, 12)); n += 1
for p in drop:
    p._element.getparent().remove(p._element)
assert n == 7, f"커버레터 패치 {n}건 (기대 7)"
doc.save(CL)
print("커버레터 7건 반영")

# ---------- Manuscript EN ----------
doc = Document(EN)
n = 0
for p in doc.paragraphs:
    if "The pandemic itself has ended, but the episode retains" in p.text:
        new = p.text.replace("The pandemic itself has ended, but", "The pandemic emergency has passed, but")
        T.fill_para(p, new, size=para_size(p)); n += 1
    elif "They also read presence as activity and average over" in p.text:
        new = p.text.replace(
            "They also read presence as activity and average over unobserved differences in sensor "
            "installation context.",
            "The analysis also treats presence as a proxy for activity and averages over unobserved "
            "differences in sensor installation context.")
        T.fill_para(p, new, size=para_size(p)); n += 1
assert n == 2, f"EN 패치 {n}건 (기대 2)"
doc.save(EN)
print("EN 2건 반영")

# ---------- Manuscript KR ----------
doc = Document(KR)
n = 0
for p in doc.paragraphs:
    if "팬데믹 자체는 끝났지만" in p.text:
        new = p.text.replace("팬데믹 자체는 끝났지만,", "팬데믹 비상 국면은 지나갔지만,")
        T.fill_para(p, new, size=para_size(p)); n += 1
    elif "또한 이 추정치는 체류(presence)를 활동으로 읽으며" in p.text:
        new = p.text.replace(
            "또한 이 추정치는 체류(presence)를 활동으로 읽으며, 관측되지 않는 센서 설치 맥락의 차이를 평균한 값이다.",
            "또한 이 분석은 체류(presence)를 활동의 대리변수로 다루며, 관측되지 않는 센서 설치 맥락의 차이를 평균한다.")
        T.fill_para(p, new, size=para_size(p)); n += 1
assert n == 2, f"KR 패치 {n}건 (기대 2)"
doc.save(KR)
print("KR 2건 반영")

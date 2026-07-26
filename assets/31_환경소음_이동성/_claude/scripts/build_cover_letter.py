# -*- coding: utf-8 -*-
# SCS 커버레터 생성: 01_Cover_Letter.docx(직전 투고 양식)를 복제해 내용만 교체.
# 서식 = 템플릿 실측(Times New Roman·제목 16pt bold 가운데·날짜 우측·본문 12pt 양쪽정렬).
import os, shutil, sys
from datetime import datetime, timezone, timedelta
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
SRC = os.path.join(ROOT, "01_논문작업", "01_Cover_Letter.docx")
ts = datetime.now(timezone(timedelta(hours=9))).strftime("%Y%m%d_%H%M%S")
OUT = os.path.join(ROOT, "01_논문작업", f"Cover_Letter_SCS_{ts}.docx")
shutil.copy2(SRC, OUT)

doc = Document(OUT)

# (text, align, size, bold, space_after_pt)
C, R, J, L = (WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT,
              WD_ALIGN_PARAGRAPH.JUSTIFY, WD_ALIGN_PARAGRAPH.LEFT)
LINES = [
 ("Cover Letter", C, 16, True, 12),
 ("July 27, 2026", R, 12, False, 12),
 ("Prof. Enrico Fabrizio, PhD", L, 12, False, 0),
 ("Prof. Fariborz Haghighat, Ph.D., P.Eng., Fellow ASHRAE, Fellow ISIAQ", L, 12, False, 0),
 ("Prof. Ryozo Ooka", L, 12, False, 0),
 ("Editors-in-Chief", L, 12, False, 0),
 ("Sustainable Cities and Society", L, 12, False, 12),
 ("Dear Editors:", L, 12, False, 12),
 ("We are pleased to submit our manuscript entitled “The dose-response of urban noise to human "
  "mobility measured with a city-scale IoT sensor network in Seoul” for consideration as an "
  "original research article in Sustainable Cities and Society.", J, 12, False, 12),
 ("Environmental noise is regarded as the second-largest environmental burden on urban health, and "
  "much of sustainable urban policy, from travel-demand management to remote work, acts on human "
  "mobility. How much quieter a city becomes when its people move less has, however, rarely been "
  "measured. The pandemic-era noise literature rests mostly on binary lockdown contrasts at a "
  "handful of monitoring points, with mobility assumed rather than observed, so the graded "
  "dose-response that policy evaluation needs has remained out of reach.", J, 12, False, 12),
 ("Our study estimates this slope directly. We combine 1,123 sensors of Seoul's permanently "
  "operated S-DoT IoT network with daily de-facto population for 421 administrative neighbourhoods "
  "over 2020-2023, a period in which Korea's graded social distancing shifted mobility repeatedly "
  "and in steps. A two-way fixed-effects design that compares neighbourhoods within the same day "
  "yields a statistically solid but small association: a 30% fall in neighbourhood mobility "
  "corresponds to about 0.23 dB less daytime noise. To our knowledge this is the first city-scale "
  "estimate of a graded mobility dose-response of urban noise built on measured, "
  "neighbourhood-level activity. The analysis also uncovers a 2.3 dB multi-year drift in the "
  "uncalibrated network, checked against the official calibrated stations, and turns this into "
  "design principles for low-cost urban sensing: rely on within-sensor change and same-day "
  "cross-neighbourhood comparison, never on absolute levels.", J, 12, False, 12),
 ("We believe this work is well suited to Sustainable Cities and Society. It addresses the "
  "environmental performance of a dense megacity with city-scale smart sensing, gives planners a "
  "quantitative basis for judging the noise co-benefits of mobility-oriented policies such as "
  "travel-demand management and car-free streets, and offers a transferable template for the "
  "growing family of low-cost urban monitoring networks. All raw data are public, and the full "
  "analysis is reproducible from open sources.", J, 12, False, 12),
 ("This manuscript is original, has not been published previously, and is not under consideration "
  "elsewhere. All authors have approved the submission and declare no competing interests.", J, 12, False, 12),
 ("We thank you for considering our work and look forward to your response.", J, 12, False, 12),
 ("Sincerely,", L, 12, False, 12),
 ("Kyung-Tae Lee", L, 12, False, 0),
 ("On behalf of all co-authors:", L, 12, False, 0),
 ("Hyun In Jo", L, 12, False, 12),
 ("Department of Architectural Engineering,", L, 12, False, 0),
 ("Kangwon National University, Gangwon-do 24341, Republic of Korea", L, 12, False, 0),
 ("Email: ktlee@kangwon.ac.kr", L, 12, False, 0),
]

# 기존 문단 전부 제거 후 재구성
for p in list(doc.paragraphs):
    p._element.getparent().remove(p._element)

for text, align, size, bold, after in LINES:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.font.bold = bold

doc.save(OUT)
print("저장:", OUT)

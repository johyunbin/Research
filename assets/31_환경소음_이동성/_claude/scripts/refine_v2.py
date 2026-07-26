# -*- coding: utf-8 -*-
# v8 정제: ①초록 의문형 제거 ②본문 대시(—) 정리(아래첨자·인용 보존) ③Table 1·2 산문 재작성 ④Fig4 재임베드.
import os, re, struct
from datetime import datetime, timezone, timedelta
from docx import Document
from docx.shared import Pt, Emu, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = r"T:\00_BACKUP\01_한양대학교\00_연구실\00_프로젝트\★_논문화 프로젝트\31_환경소음_이동성"
FIGD = os.path.join(ROOT, "_claude", "figures")
SRC = os.path.join(ROOT, "01_논문작업", "Manuscript_v8_20260621_235125.docx")
ts = datetime.now(timezone(timedelta(hours=9))).strftime("%Y%m%d_%H%M%S")
OUT = os.path.join(ROOT, "01_논문작업", f"Manuscript_v8_{ts}.docx")

def setfont(run, size=12, bold=False, italic=False, sub=False):
    run.font.size = Pt(size); run.bold = bold; run.italic = italic; run.font.name = 'Times New Roman'
    if sub: run.font.subscript = True
    rpr = run._element.get_or_add_rPr(); rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'): rf.set(qn(a), 'Times New Roman')
    rf.set(qn('w:eastAsia'), '바탕')

# 아래첨자 토큰(긴 것 우선) — p.text로 읽힌 형태 기준
SUB = [('ΔLAeq', ('ΔL', 'Aeq')), ('Leq,24h', ('L', 'eq,24h')), ('LAeq', ('L', 'Aeq')),
       ('Lday', ('L', 'day')), ('Lnight', ('L', 'night')), ('Leq', ('L', 'eq'))]
SRE = re.compile('(' + '|'.join(re.escape(k) for k, _ in SUB) + ')')
SMAP = dict(SUB)

def emit(p, text, size=12, bold=False, italic=False):
    for part in SRE.split(text):
        if not part: continue
        if part in SMAP:
            b, sb = SMAP[part]
            setfont(p.add_run(b), size, bold=bold, italic=italic)
            setfont(p.add_run(sb), size, bold=bold, italic=italic, sub=True)
        else:
            setfont(p.add_run(part), size, bold=bold, italic=italic)

def replace_para(p, text, size=12):
    for r in list(p.runs): r._element.getparent().remove(r._element)
    emit(p, text, size)

def cell_set(cell, text, size=10, bold=False):
    ps = cell.paragraphs
    for extra in ps[1:]: extra._element.getparent().remove(extra._element)
    p = ps[0]
    for r in list(p.runs): r._element.getparent().remove(r._element)
    emit(p, text, size, bold=bold)

doc = Document(SRC)
paras = doc.paragraphs

# ---------- ① 초록 의문형 제거 ----------
ABS = ("How strongly the urban acoustic environment responds to human mobility is a first-order question for "
       "mobility, land-use and noise policy, yet it remains difficult to answer because mobility rarely varies "
       "exogenously and is entangled with weather, land use and time of day. We exploit Korea's graded COVID-19 "
       "social-distancing system, which repeatedly tightened and relaxed activity over 2020-2023, as a natural "
       "experiment, paired with a city-scale low-cost IoT noise network, to estimate a measured-mobility "
       "dose-response for urban noise. We assembled 1,248,794 sensor-days from 1,123 Smart Seoul Data-of-Things "
       "(S-DoT) sensors across 421 administrative neighbourhoods and matched each sensor-day to neighbourhood "
       "daytime de-facto population. Because low-cost sensors carry calibration offsets and multi-year drift, we "
       "identify the effect only from within-sensor variation and, in the headline specification, from "
       "cross-neighbourhood variation within each date (two-way fixed effects). Daytime noise rises and falls with "
       "daytime mobility (beta = +0.65 dB per log-unit; a 30% mobility reduction corresponds to about 0.23 dB). The "
       "effect is robust and consistent across seasons but specific to daytime: a nighttime association vanishes "
       "under two-way fixed effects, and the day-night gap does not respond. No robust gradient survives in the "
       "long-run spatial cross-section. Official roadside LAeq stations confirm the network tracks real trends "
       "while exposing that multi-year absolute levels are confounded. The dose-response is statistically robust "
       "but modest, qualifying the binary-lockdown literature, and the analysis doubles as a design guide for "
       "IoT-based urban environmental monitoring.")

# ---------- ② 본문 대시 정리: (시작 마커, 새 본문) ----------
REPL = [
("환경소음은 대기오염에 이어",
 "환경소음은 대기오염에 이어 도시 환경이 인간 건강에 부과하는 두 번째로 큰 부담으로 평가된다. 세계보건기구(WHO)는 환경소음을 유럽 도시의 주요 질병부담 원인으로 규정하며, 서유럽에서만 매년 100만 healthy life-years 이상이 교통소음으로 손실되는 것으로 추정한다 [1,2]. 소음의 건강영향은 단순한 불쾌를 넘어선다. 야간소음은 수면을 단편화하고 [3], 만성적 노출은 자율신경·내분비 스트레스 경로를 통해 고혈압·허혈성 심질환 등 심혈관계 질환 위험을 높이며 [4,5], 주간에는 소음성가심(annoyance)과 인지수행 저하를 유발한다 [6,7]. 이러한 영향은 막대한 사회·경제적 비용으로 환산된다 [8]."),
("따라서 소음은 도시의 지속가능성",
 "따라서 소음은 도시의 지속가능성과 거주적합성(livability)을 좌우하는 핵심 물리적 스트레서이며, 그 부담이 사회경제적 약자에게 불균등하게 분포한다는 점에서 환경정의(environmental justice)의 문제이기도 하다. 소음을 단순 데시벨이 아니라 맥락 속 인간 경험으로 다루는 사운드스케이프(soundscape) 패러다임 [9,10,11,12,13,14]은 이 인식을 확장해, 도시 공공공간의 음향쾌적성 [15]과 토지이용·도시형태에 따른 음환경의 시공간 변이 [16]를 연구 대상으로 끌어올렸다. 그러나 소음을 '관리'하려면 한 가지 근본 질문에 답해야 한다. 도시소음은 계획가가 바꿀 수 있는 것, 즉 인간의 활동과 이동량에 얼마나 민감하게 반응하는가? 이 반응성(responsiveness)이야말로 모든 이동·토지이용 기반 소음정책의 효과를 가늠하는 출발점이지만, 놀랍게도 정량적으로 규명된 바가 거의 없다."),
("기존의 교통소음 예측 모델",
 "기존의 교통소음 예측 모델, 예컨대 CNOSSOS-EU 같은 표준식이나 토지이용회귀(land-use regression) 모델 [17,18]은 교통량·도로폭·토지이용으로부터 소음을 추정하지만 본질적으로 횡단적·정적(static)이다. 즉 '교통이 많은 곳이 시끄럽다'는 공간 패턴은 잘 재현하나, '이동량이 줄면 소음이 얼마나 주는가'라는 동적 용량-반응(dynamic dose-response)은 답하지 못한다. 이 빈틈을 메우려면 이동량의 외생적·점진적(graded) 변동과, 그에 대응하는 고밀도 소음 관측이 동시에 필요하다. 이는 정상적인 도시에서는 좀처럼 주어지지 않는 조건이다."),
("코로나-19 대유행은 바로 이 조건",
 "코로나-19 대유행은 바로 이 조건을 전례 없이 만들어 낸 자연실험(natural experiment)이었다. 봉쇄·거리두기는 인간 이동량을 외생적으로 급감시켰고, 다수 연구가 봉쇄기 도시소음의 수 dB 감소를 보고했다(마드리드는 주간 약 3 dB 감소 [19], 더블린 [20], 런던 [21], 스톡홀름 [22], 몬트리올 [23], 그라나다 [24], 부에노스아이레스 [25], 리마 [26], 피렌체 [27], 인도 도시들 [28,29]). 같은 시기 대기질에서도 이산화질소·이산화탄소가 급감해 [30,31,32] 봉쇄가 도시 환경 전반에 미친 충격을 입증했고, 봉쇄기 사운드스케이프 지각의 변화도 관찰되었다 [33,34]."),
("그러나 이 빠르게 축적된 문헌",
 "그러나 이 빠르게 축적된 문헌은 세 가지 구조적 한계를 공유하며, 이는 본 연구의 동기를 직접 규정한다. 첫째, 대부분 '봉쇄 vs 비봉쇄'의 이항적(binary) 비교로, 규제 강도가 단계적으로 달라지는 graded 용량-반응 곡선을 추정할 수 없다. 그 결과 소음이 이동량 변화에 비례적으로 반응하는지, 그 기울기가 얼마인지 알 수 없다. 둘째, 거의 모든 연구가 소음 감소를 '교통이 줄어서'라고 서술적으로 귀인할 뿐, 이동량을 실제로 측정해 소음과 정량적으로 결합하지 않는다. 노출(exposure)이 가정될 뿐 측정되지 않는 것이다. 셋째, 측정점이 소수(대개 4~70점)이고 서구 도시에 편중되어, 고밀도·동아시아 도시의 미시적 공간 이질성과 토지이용별 차등을 포착하지 못한다. 요컨대 graded·measured·high-density·동아시아라는 네 조건을 동시에 충족하는 증거가 부재하며, 이것이 본 연구가 필요한 이유다."),
("최근 두 가지 데이터 혁신",
 "최근 두 가지 데이터 혁신이 이 빈틈을 메울 가능성을 열었다. 첫째는 모바일 통신·신호 기반의 de-facto(체류) 인구다. 이는 거주지가 아니라 특정 시점에 사람들이 '실제로 어디에 있는가'를 도시 미세공간·고빈도로 측정한다 [35,36,37,38,39]. 둘째는 저가 IoT 음향 센서망이다. SONYC [40], 저가 모니터링 기기 [41], 무선음향센서망(WASN) [42,43], 참여형 NoiseCapture [44], 스마트시티 센싱 [45,46] 등은 도시 음환경을 수백~수천 점에서 상시 관측할 수 있게 했다. 한국의 단계적 거리두기, 서울 전역의 S-DoT 소음센서망(약 1,100점), 그리고 행정동 단위 생활인구의 결합은 위 네 조건을 동시에 만족시키는 드문 기회를 제공한다."),
("그러나 저가 IoT 소음망에는",
 "그러나 저가 IoT 소음망에는 기존 COVID-소음 문헌이 대체로 간과해 온 본질적 함정이 있다. 바로 센서마다 다른 검교정 오프셋과, 시간이 지나며 측정값이 서서히 변하는 드리프트(표류)다 [47]. 저가 센서는 절대 음압을 정밀 측정하도록 검교정되지 않으며, 센서마다 고유의 상수 오프셋을 갖고, 시간이 지나며 노후·환경요인으로 측정값이 표류한다. 이를 무시하고 절대레벨을 시계열·공간으로 단순 비교하면 거짓 패턴 또는 거짓 결론에 이른다(본 연구의 §3.5가 그 실례를 보인다). 따라서 고밀도 IoT 자료를 신뢰성 있게 쓰려면 이 교란을 식별 설계에서 정면으로 제거해야 한다. 이것이 본 연구의 방법론적 핵심이자, 스마트시티 환경 데이터 활용 전반에 적용되는 교훈이다."),
("본 연구는 한국의 단계적 사회적 거리두기를 graded",
 "본 연구는 한국의 단계적 사회적 거리두기를 graded 자연실험으로, 서울의 S-DoT IoT 소음센서를 결과변수로, 행정동 생활인구를 측정 이동량 dose로 결합해 도시소음의 이동량 용량-반응을 추정한다. 핵심적으로, 우리는 절대레벨이 아닌 센서-내 상대변화와 같은 날짜 안에서 동(洞) 간 변이만으로 효과를 식별하는 양방향 고정효과(two-way fixed effects) 전략 [48,49]과 자연실험 틀 [50]을 채택하고, 저가 IoT의 드리프트·해상도 한계를 공식 측정망으로 교차검증한다. 이로써 서울 사운드스케이프의 시공간 변이 연구 [16]를 정량적·인과적 이동량-소음 추정으로 확장한다. 본 연구는 세 가지 질문에 답한다. RQ1: 측정 이동량 감소는 도시소음을 얼마나 낮추는가(용량-반응 기울기)? RQ2: 그 효과는 토지이용·주야·요일·계절에 따라 어떻게 다른가? RQ3: 거리두기 전환에서 소음은 어떻게 반등하며 어떤 공간패턴을 갖는가? 기여는 네 가지다. (1) 측정 이동량에 대한 graded 용량-반응의 최초 정량 추정, (2) 약 1,100점 고밀도 IoT 센서망의 활용, (3) 서울·동아시아 고밀도 메가시티라는 미출판 지리, (4) 저가 IoT 소음망의 드리프트·해상도 한계를 드러내고 cross-sectional 식별의 필요성을 입증하는 방법론적 교훈."),
("결과변수는 서울시가 도시 전역",
 "결과변수는 서울시가 도시 전역에 상시 운영하는 IoT 도시데이터 센서망(Smart Seoul Data of Things, S-DoT)의 소음 자료다. 약 1,100개 지점에서 시간별 광대역 음압(dB)을 측정하며, 모든 자료는 별도 인증 없이 공개 다운로드된다(Table 1). 우리는 시간별 값을 하루 단위의 에너지 등가소음도로 합쳐, 전일 Leq,24h, 주간 Lday(06-21시), 야간 Lnight(22-05시)을 각각 10·log10(평균(10^(L/10)))으로 계산하고, 하루 최소 12시간 이상 관측되고 값이 물리적으로 타당한 범위(20-95 dB)에 드는 자료만 남겼다. 2023년에는 자료 구조가 바뀌어 소음이 최대·평균·최소 세 값으로 나뉘었는데, 이 가운데 평균값을 사용했다. 한 가지가 결정적이다. S-DoT는 정밀 검교정을 거치지 않은 광대역 데시벨이어서(소음 분야의 표준 측정량인 A-가중 등가소음도 LAeq라는 라벨이 붙어 있지 않다), 서로 다른 센서의 절대값을 맞비교하는 데에는 쓸 수 없다. 따라서 우리는 절대 수준이 아니라 각 센서의 시간에 따른 변화만 분석에 쓰고, 그 신뢰성은 공식 LAeq 측정망과 대조해 검증한다(§3.5)."),
("이동량 노출은 서울 '생활인구",
 "이동량 노출은 서울 '생활인구(living / de-facto population)'로 측정했다. 이는 거주 인구가 아니라 모바일 통신 신호로 추정한, 특정 시점에 각 행정동에 실제로 체류하는 인구수다(Table 2). 시간별 체류수(stock)를 일 평균(및 주간 06-21 / 야간 22-05 평균)으로 집약해 424개 동 × 1,461일 = 619,464 동-일을 얻었다(결측 0). dose 구성에서 한 가지 통찰이 중요하다. 도시 전체 생활인구 총량은 거의 보존된다(사람들이 외출을 줄여도 자택 체류로 계수되기 때문). 즉 거리두기의 신호는 총량이 아니라 공간 재분포에 있으며, 상업·업무 동은 주간 체류가 20-35% 비워지고 주거 동은 최대 37% 증가한다. 따라서 우리는 dose를 동의 자기기준 대비 상대변화로 정의한다. 곧 lp = log(동의 일별 생활인구 / 동의 post-lift 정상기 평균)이다. 이렇게 하면 dose는 각 동이 자기 정상에서 얼마나 벗어났는가를 나타내며, 동·센서별 절대수준 차이에 영향받지 않는다."),
("거리두기 '단계'(1·2·2.5·4단계",
 "거리두기 '단계'(1·2·2.5·4단계 등)를 노출변수로 그대로 쓰는 것은 적절치 않다. 단계 체계가 두 차례 통째로 재정의되었고(2020-11, 2021-07), 단계 사이의 간격이 일정하지 않으며, 같은 단계라도 시기에 따라 실제 이동량이 달랐기 때문이다. 대신 우리는 매일의 규제를 두 가지 구체적 수치, 곧 식당·카페 영업종료시각과 사적모임 허용 인원으로부터 0(규제 없음)에서 7(최강)까지 이어지는 연속적 강도(stringency) 지수로 다시 코딩했다(Table 4). 다만 이 지수는 보조 변수일 뿐이며, 본 연구의 주된 노출변수는 어디까지나 실측 이동량(생활인구)이다. 기상은 Open-Meteo ERA5 일자료(기온·강수·풍속), 달력 변수는 요일·주말·한국 공휴일·계절을 포함한다."),
("저가 S-DoT 센서는 정밀 검교정을 거치지 않아",
 "저가 S-DoT 센서는 정밀 검교정을 거치지 않아, 똑같은 소리를 들려주어도 센서마다 일정량 높거나 낮게 기록하는 고유의 치우침(상수 오프셋)을 갖는다. 이 치우침의 크기는 센서마다 다르고 알 수 없으므로, 서로 다른 두 센서의 절대 데시벨을 맞비교하는 것은 의미가 없다. 그러나 한 센서가 시간에 따라 얼마나 달라졌는지(예컨대 오늘이 그 센서의 평소보다 얼마나 큰지)만 보면, 그 센서에 늘 똑같이 들어 있는 치우침은 빼는 과정에서 저절로 지워진다. 그래서 우리의 모든 분석은 절대 수준이 아니라 각 센서의 자기 대비 변화(센서 내 상대변화)에 기반한다. 통계적으로 이는 고정효과(fixed-effects) 모형으로 구현된다. 각 센서의 전체 평균을 빼 줌으로써 그 센서 고유의 치우침을 제거하는 것이다(125만 건의 자료에 센서별 더미변수를 일일이 넣는 대신, 집단 평균을 차감하는 수학적으로 동등한 방법을 썼다). 또한 같은 센서나 같은 동에서 나온 관측치들은 서로 닮아 있을 수 있으므로, 표준오차를 센서·동 단위로 보정했다 [48]. 전체 식별 논리는 Fig. 2에 요약했다."),
("주된 분석은 두 단계의 고정효과 모형이다",
 "주된 분석은 두 단계의 고정효과 모형이다. (M1) 센서 고정효과 모형은 각 센서의 평균을 제거한 뒤, 그 센서의 소음 변화를 동 이동량과 날씨(기온·기온²·강수·풍속·강수 여부)·요일·주말·공휴일로 설명한다. 이렇게 하면 눈에 보이는 시간적 교란을 통제하면서도 시간에 따른 변이는 남겨 두어 이동량의 효과를 추정할 수 있다. (M2)는 본 연구의 핵심 모형으로, 여기에 날짜 고정효과를 더한다. 날짜 고정효과는 특정 날짜에 도시 전체가 공통으로 겪은 모든 것을 한꺼번에 흡수한다. 곧 그날의 날씨, 요일·공휴일, 해가 갈수록 변해 온 도시 전반의 소음 추세, 센서 노후로 인한 표류, 전국적 거리두기가 모두 여기에 포함된다. 그러면 남는 정보는 오직 '같은 날, 같은 도시 안에서 동(洞)마다 이동량이 얼마나 달랐는가'뿐이다. 따라서 M2는 다음 질문에 답한다. 같은 날, 자기 평소보다 더 비워진 동에 있는 센서가 덜 비워진 동의 센서보다 더 조용해졌는가? 이 설계의 한 가지 귀결로, 도시 전체에 하루 하나의 값으로만 변하는 변수(예: 그날의 지하철 총승객 수나 거리두기 강도)는 날짜 고정효과에 완전히 흡수되어 따로 효과를 추정할 수 없다. 동마다 값이 다른 실측 이동량만이 유효한 노출변수다. 두 종류의 고정효과는 센서 평균과 날짜 평균을 번갈아 빼는 계산을 수렴할 때까지 반복해 처리했다."),
("분석에 사용한 변수들의 기술통계는",
 "분석에 사용한 변수들의 기술통계는 Table 3에 정리했다. 센서 고정효과 모형(M1, Table 5)에서 한 동의 주간 이동량과 그 동의 주간 소음은 같은 방향으로 움직였다(β=+1.130 dB/log-unit, SE 0.209, p<0.001). 통제변수들은 모두 물리적으로 타당했다. 기온은 U자형(춥거나 더울 때 소음이 커짐)이었고, 강수 +0.045 dB/mm, 비 오는 날 +0.10 dB, 주말 −0.83 dB, 공휴일 −1.18 dB였다(모두 p<0.001). 그러나 날짜 고정효과까지 더해 도시 전체가 그날 공통으로 겪은 변화(날씨·추세·표류 등)를 모두 걷어 낸 핵심 모형(M2)에서는 효과가 작아졌다(주간 β=+0.648 dB/log-unit, 95% CI 0.17-1.13, p=0.008; 전일 β=+0.628, p=0.038). 풀어 보면, 한 동의 주간 이동량이 평소보다 30% 줄면 그 동의 주간 소음은 약 0.23 dB, 50% 줄면 약 0.45 dB 낮아진다. 즉 도시 전체에 공통된 시간 변화를 엄밀히 제거하고 '같은 날 동 사이의 차이'만으로 추정한 용량-반응은 통계적으로 견고하지만 그 크기는 작다(RQ1)."),
("주간과 야간을 나눠 보면",
 "주간과 야간을 나눠 보면 효과가 주간 활동에서 비롯됨이 분명해진다(Fig. 4). 야간 소음은 센서 고정효과만 둔 모형에서는 야간 이동량에 강하게 반응하는 것처럼 보였으나(β=+2.28, p<0.001), 날짜 고정효과까지 더하자 그 관계가 사라졌다(β=+0.50, 95% CI −0.73~+1.74, p=0.42). 이는 야간의 겉보기 효과가 실제 인과가 아니라, 시간에 공통으로 작용한 교란(센서 표류·계절·도시 추세)이 만든 착시였음을 보여 준다. 반면 주간 효과는 날짜 고정효과 아래에서도 살아남는다. 한편 주야 차이(Lday − Lnight)는 같은 센서가 같은 날 기록한 두 값의 차이이므로, 그 센서의 고유 치우침과 그날 공통으로 작용한 표류가 모두 상쇄된다. 이 주야 차이에 대한 이동량 효과는 사실상 0이었다(β=+0.02, p=0.92). 즉 이동량은 낮과 밤의 소음 격차를 바꾸는 것이 아니라 주간 소음 수준 자체를 소폭 낮추는 방식으로 작용한다."),
("이동량이라는 노출이 시간과 공간에서",
 "이동량이라는 노출이 시간과 공간에서 어떻게 움직였는지가 본 설계의 핵심이다(Fig. 5). 도시 전체의 생활인구 총량은 거의 변하지 않지만(외출을 줄여도 자택 체류로 계수되므로), 그 인구가 어디에 머무는지는 거리두기 국면마다 크게 재배치된다. 상업·업무 동은 주간 체류가 최대 약 35% 비워지는 반면 주거 동은 오히려 늘어난다. 규제가 가장 강했던 시기(2020-12 5인 이상 모임금지, 2021-07 수도권 4단계)에 도심 상업동이 가장 비워졌고, 위드코로나(2021-11)와 전면해제(2022-04)를 거치며 평소 수준으로 회복했다."),
("이 기능적 차등은 시간축에서도",
 "이 기능적 차등은 시간축에서도 뚜렷하다(Fig. 6, RQ3). 상업 동의 주간 이동량은 제한기 내내 기준선 아래로, 주거 동은 기준선 위로 벌어졌다가 2022년 해제 후 수렴한다(수도권 4단계기 중앙값 상업 0.98 vs 주거 1.05). 드리프트를 제거한 상대 소음(동 그룹평균 − 도시평균)에서도 상업 동이 제한기에 도시평균보다 상대적으로 조용한 경향이 관찰된다. 거리두기 전환의 동역학은 차이의 차이(difference-in-differences)로 본다(Fig. 7). 고영향 동과 저영향 동의 주별 ΔLAeq 차이는 제한기에 음(−)으로 기울고 정상기에 0으로 수렴했으며, 주별 (이동량 차이 vs 소음 차이) 상관은 Pearson +0.44 · Spearman +0.46으로 일치했다."),
("반면 동을 하나의 점으로 보고",
 "반면 동을 하나의 점으로 보고 거리두기 기간 전체의 평균을 비교하는 '장기 공간 단면'에서는 이동량과 소음의 뚜렷한 경향이 나타나지 않았다(Fig. 8). 동별로 (이동량이 얼마나 줄었는가)와 (도시 평균 대비 소음이 얼마나 달라졌는가)의 관계는, 극단값에 강한 순위상관으로 보면 사실상 0이었다(전체 Spearman ρ≈−0.05). 상업 동에서는 일반 상관계수(Pearson r=+0.38)가 마치 관계가 있는 듯 보였지만, 이는 센서가 하나뿐이라 추정이 불안정한 소수의 동이 만든 착시였다. 그 동들을 빼고 극단값에 강한 방법으로 다시 보면 관계는 0에 가까웠다(Spearman ρ=+0.08, Theil-Sen 기울기≈0; Fig. 9). 요컨대 이동량-소음 신호는 '같은 날 동 사이의 차이'(M2·M4)에서만 안정적으로 나타나며, 동별 장기 평균을 단순히 지도에 칠하는 방식으로는 잡히지 않는다."),
("각 센서의 자기 대비 변화(ΔLAeq, 정상기 기준)를 그대로",
 "각 센서의 자기 대비 변화(ΔLAeq, 정상기 기준)를 그대로 그려 보면 역설이 나타난다. 규제가 가장 강했던 시기의 주간 소음이 오히려 정상기보다 높게 나오는 것이다(+0.85 vs +0.34 dB). 진단해 보니 이는 코로나 효과가 아니라 여러 해에 걸친 절대 수준의 변동(표류) 때문이었다(Fig. 10). 4년 내내 가동된 842개 센서의 연평균 Leq,24h는 49.4(2020)→47.6→47.2→47.1 dB(2023)로 해마다 꾸준히 낮아졌고(85% 센서가 하락 추세), 이 하락이 2020-21년을 상대적으로 '시끄럽게' 보이게 만든 것이다. 다만 2023년 자료구조 변경 지점에서는 수준 도약이 없었다(+0.02 dB). 그런데 이 하락은 모든 센서에 똑같이 일어난 단순한 기계 노후가 아니다. 공식 측정소가 있는 도로변 4지점에서 가까운 S-DoT의 2021→2023 추세는 City Hall +1.7·Sinsa +2.5·Sinchon ≈0·Seongsu −0.3 dB로 하락하지 않았고, 공식 LAeq 추세(각 +7.2[공사 추정]·+0.4·−0.4·+1.5)와 같은 방향이었다. 즉 도로변 S-DoT는 공식 측정망을 잘 따라가 신뢰할 만하지만, 도시 전체 평균의 하락(2021→2023 −0.5 dB)은 도로에서 벗어난 조용한 지역에 집중된 위치별 현상이어서, 여러 해에 걸친 절대 수준 비교 자체를 믿을 수 없게 만든다. 이것이 바로 우리가 절대·시계열 비교 대신 '같은 날 동 사이의 차이'(M2)에 의존하는 이유다."),
("본 연구는 측정 이동량에 대한 도시소음의 단계적",
 "본 연구는 측정 이동량에 대한 도시소음의 단계적(graded) 용량-반응을 처음으로 정량 추정했다. 핵심 결과는 네 가지로 요약된다. 첫째, 한 동의 주간 이동량과 그 동의 주간 소음 사이에 통계적으로 견고한 정(+)의 용량-반응이 존재하지만 그 크기는 작다(이동량 30% 감소 ≈ 소음 −0.2~0.3 dB). 둘째, 효과는 주간에 한정되며 네 계절에 걸쳐 일관된다. 야간의 겉보기 효과는 엄밀한 식별에서 시간 교란으로 사라진다. 셋째, 효과는 '같은 날 동 사이의 차이'와 차이의 차이 분석에서만 안정적으로 드러나며, 동별 장기 평균을 비교하는 공간 단면에서는 사라진다. 넷째, 저가 IoT 망의 여러 해에 걸친 절대 수준은 센서 표류와 지역별 회복 차이로 교란되어 시계열·절대 비교를 믿을 수 없고, 오직 센서 내 변화와 동 사이 비교에 기반한 설계만이 유효하다."),
("우리의 추정치는 봉쇄기 도시소음이 수 dB",
 "우리의 추정치는 봉쇄기 도시소음이 수 dB 줄었다고 보고한 기존 문헌 [19,20,21,22,23,24]과 정면으로 대비된다. 그 보고들은 대개 특정 번화가의 절대 소음을 팬데믹 전후로 단순 비교한 것이라, 가장 시끄럽고 가장 크게 비워진 소수 지점의 극단적 변화를 도시 전체로 일반화하기 쉽다. 반면 본 연구는 약 1,100개 지점 전역에서, 센서 내 변화로·이동량의 정도에 따라·전 기간에 걸쳐 한계효과(이동량이 한 단위 변할 때의 평균적 소음 변화)를 추정했기에, 도시 평균의 반응은 그보다 훨씬 작게 나타난다. 더욱이 우리 추정치는 하한(최소값)이다. 핵심 모형은 '같은 날 동 사이의 차이'만으로 효과를 잡아내는데, 모두가 동시에 덜 움직인 '도시 전체 공통의 이동량 감소' 성분은 센서 표류와 뒤섞여 따로 떼어낼 수 없기 때문이다. 효과가 작은 것은, 도시소음이 통과 교통이나 간선도로처럼 잘 줄지 않는 배경 소음원에 의해 '끈적하게(sticky)' 유지된다는 점, 그리고 S-DoT의 1 dB 단위 거친 해상도가 측정 잡음으로서 표준오차를 넓힌다는 점과도 들어맞는다. 요컨대 '봉쇄가 도시를 수 dB 조용하게 만들었다'는 통념은 엄밀한 이동량-소음 관점에서는 과장이며, 진짜 단계적 용량-반응은 견고하되 작다."),
("효과가 나타나는 양상은 그 메커니즘",
 "효과가 나타나는 양상은 그 메커니즘과 잘 들어맞는다. 효과가 주간에·연중 일관되게 나타나는 것은, 주간 소음이 상거래·통근·방문 같은 사람의 활동에 직접 연동되는 반면 야간 소음은 고정 설비나 간선교통 같은 배경원이 지배해 이동량 변화에 둔감하기 때문이다. 야간의 겉보기 효과가 엄밀한 식별에서 사라진 사례는, 단순한 시간 추세를 인과로 오인할 위험을 잘 보여 준다. 효과가 토지이용에 비교적 균일했다는 점은, 이동량-소음 반응이 특정 장소 유형에 몰려 있지 않고 도시 전반에 얇게 퍼져 있음을 시사한다. 특히 주목할 것은 동 단위 장기 공간 단면에서 깨끗한 경향이 나타나지 않은 점이다. 이는 효과가 없어서가 아니라, 동별 장기 평균 소음 변화가 국지적 공사·도로 변화·센서별 표류 같은 특이 요인에 압도되기 때문이다. 신호는 같은 날 동 사이의 빠른 변동에 있고, 단순한 공간 지도화로는 드러나지 않는다. 실제로 상업 동에서 잠깐 보였던 양의 상관조차 센서 하나뿐인 소수 동의 극단값이 만든 착시였고, 극단값에 강한 방법으로 보면 0이었다. 이는 고밀도 IoT 자료의 공간 분석이 얼마나 쉽게 잘못된 결론으로 이어질 수 있는지 보여 주는 경고다."),
("본 연구의 방법론적 기여는 결과 못지않게",
 "본 연구의 방법론적 기여는 결과 못지않게 중요하다. 저가 고밀도 IoT 소음망으로 도시 환경효과를 측정하려는 시도는 세 가지 함정을 정면으로 드러낸다. (i) 센서마다 다른 고유의 치우침은 센서 내 변화만 쓰는 설계로 상쇄해야 하고, (ii) 여러 해에 걸친 표류·센서 설치 시점 차이·지역별 회복은 절대 수준의 시계열·공간 비교를 불가능하게 하므로 '같은 날 동 사이 비교'가 필수이며, (iii) 도시 전체에 하나의 값으로만 변하는 노출변수는 시간 교란에 취약하므로 동 단위로 측정된 노출이 필요하다. 특히 장기 공간 단면이 안정적으로 영(null)이라는 결과는, 검교정 없이 절대 수준을 그대로 지도에 칠하는 접근이 거짓 패턴(또는 거짓 '무패턴')을 낳을 수 있음을 경고한다. 이는 SONYC [40]·무선음향센서망 [42]·NoiseCapture [44]처럼 확산 중인 저가 소음 모니터링 [41], 나아가 스마트시티 센서 데이터 [45,46]의 정책 활용 전반에 적용된다. 본 연구가 보여 준 고정효과·극단값에 강한 통계·공식망 교차검증 [47]의 절차는, 앞으로의 IoT 기반 도시 환경연구가 따를 수 있는 식별 설계의 본보기가 된다."),
("계획·정책 측면에서, 이동량·교통수요 관리",
 "계획·정책 측면에서, 이동량·교통수요 관리가 소음에 주는 직접적 공편익(co-benefit)은 통념보다 작다. 봉쇄급 이동량 급감조차 동 단위에서 1 dB 안팎의 변화를 낳았을 뿐이다. 따라서 '15분 도시'·차 없는 거리 등 이동량 저감 정책은 그 자체로 정당하나, 소음 저감 수단으로는 노면 포장·저소음차량·차폐 등 음원·전파 단계 개입 [17,51]과 병행되어야 한다. 효과가 토지이용에 비교적 균일했다는 점은 상업/주거 구분에 따른 차등 규제보다 광역적 접근의 여지를 시사한다. 서울은 세계 최고 밀도의 메가시티 중 하나로, 본 결과는 도쿄·홍콩·싱가포르 등 고밀도 동아시아 도시에 일반화 가능성이 높다."),
("본 연구에는 몇 가지 한계가 있다",
 "본 연구에는 몇 가지 한계가 있다. (1) S-DoT는 검교정되지 않은 데다 1 dB 단위의 거친 해상도와 노후에 따른 표류를 가져 절대 수준과 여러 해에 걸친 시계열을 그대로 믿기 어렵다. 그래서 우리는 센서 내 상대변화와 '같은 날 동 사이 비교'로 우회했다. (2) 생활인구는 사람이 그곳에 '있다'는 정보일 뿐, 소음을 실제로 만드는 활동(특히 교통)을 정확히 대변하지는 못한다. 앞으로 버스·지하철 정류장별 승하차를 동 단위로 연결한 더 직접적인 교통 노출 자료 [35,37]를 쓰면 효과가 더 또렷해질 수 있다. (3) 거리두기 강도 코딩의 정확한 시행 일자는 질병관리청 원자료와의 추가 대조가 필요하다. (4) 기상은 재분석 자료(ERA5)로, 기상청 공식 관측으로 대체할 수 있다. (5) 모두가 동시에 덜 움직인 '도시 전체 공통' 성분은 센서 표류와 뒤섞여 따로 떼어낼 수 없으므로 우리 효과는 하한이며, 검교정된 참조망과 결합해 절대 수준을 복원하는 후속연구가 필요하다. (6) 토지이용은 주간 대비 야간 인구비로 근사했으며, 실제 토지이용 분류로 더 정밀화할 수 있다. (7) 분석은 서울 한 도시에 한정되며, 소음의 주파수·시간 미세구조와 사람의 사운드스케이프 지각 [11,16,34]으로의 확장이 남은 과제다."),
]

# --- 초록 교체 ---
abs_done = False
for p in paras:
    if p.text.strip().startswith("How strongly"):
        replace_para(p, ABS, 12); abs_done = True; break
print("초록 교체:", abs_done, "| 단어수:", len(ABS.split()))

# --- 본문 대시 문단 교체 ---
rs = next(i for i, p in enumerate(paras) if p.text.strip() == "References")
body = paras[:rs]
done = 0
for marker, newtext in REPL:
    hits = [p for p in body if p.text.strip().startswith(marker)]
    assert len(hits) == 1, f"마커 매칭 {len(hits)}건: {marker[:30]}"
    replace_para(hits[0], newtext, 12); done += 1
print(f"본문 문단 교체: {done}/{len(REPL)}")

# ---------- ③ Table 1·2 산문 재작성 + Table 4 헤더 명료화 ----------
T1 = [("Item", "Description"),
      ("Dataset", "Smart Seoul Data of Things (S-DoT), the city's permanent urban-sensing network; Seoul Open Data Plaza dataset OA-15969, downloadable without authentication."),
      ("Coverage", "About 1,100 fixed sensors across all 25 districts of Seoul, recorded hourly from April 2020 to December 2023."),
      ("Measurement", "Broadband sound pressure level (dB) at each sensor, combined into energy-equivalent daily levels: whole-day Leq,24h, daytime Lday (06:00–21:00) and night-time Lnight (22:00–05:00)."),
      ("Quality control", "Sensor-days with fewer than 12 valid hours, or with values outside 20–95 dB, were discarded; a 2023 change in the published field (to a single 'mean noise' value) was reconciled to one common definition."),
      ("Use", "The sensors are not calibrated to a shared absolute reference, so they are used only for change over time within each sensor, and are cross-validated against the official four-point roadside LAeq network (§3.5).")]
T2 = [("Item", "Description"),
      ("Mobility exposure", "Daytime de-facto ('living') population for each administrative dong, i.e. the number of people actually present in a neighbourhood at a given time, estimated from mobile-network signalling rather than from registered residence. Seoul Open Data Plaza, dataset OA-14991."),
      ("Coverage", "Hourly counts for every Seoul dong, averaged to daily daytime (06:00–21:00) and night-time (22:00–05:00) values; 424 dong × 1,461 days (619,464 dong-days). The roughly 2.8 GB of raw hourly files were processed and then deleted to save storage."),
      ("Dose variable", "Within-dong log relative change: the natural logarithm of a day's population divided by that dong's own baseline mean for the post-lifting normal period, so each dong is compared against its own usual level."),
      ("Weather", "Daily mean, maximum and minimum temperature, precipitation and maximum wind speed, from the Open-Meteo ERA5 reanalysis at the Seoul city-centre grid point (37.57°N, 126.97°E)."),
      ("Calendar", "Day of week, a weekend indicator, Korean public holidays (including substitute holidays) and season.")]

def find_table(sig):
    for tb in doc.tables:
        if tb.rows[0].cells[0].text.strip() == sig[0] and sig[1] in tb.rows[1].cells[1].text:
            return tb
    return None

t1 = find_table(("Item", "Smart Seoul"))
t2 = find_table(("Item", "Administrative-dong"))
assert t1 is not None and t2 is not None, "Table 1/2 식별 실패"
for tb, data in ((t1, T1), (t2, T2)):
    assert len(data) == len(tb.rows), f"행수 불일치: data {len(data)} vs table {len(tb.rows)}"
    for ri, (item, desc) in enumerate(data):
        cell_set(tb.rows[ri].cells[0], item, size=10, bold=(ri == 0))
        cell_set(tb.rows[ri].cells[1], desc, size=10, bold=(ri == 0))
print("Table 1·2 재작성 완료")

# Table 4(거리두기) 헤더 명료화: Close (h)->Business curfew (h), Gathering->Gathering cap
for tb in doc.tables:
    hdr = [c.text.strip() for c in tb.rows[0].cells]
    if hdr[:2] == ["From", "Regime"]:
        cell_set(tb.rows[0].cells[2], "Business curfew (h)", size=10, bold=True)
        cell_set(tb.rows[0].cells[3], "Gathering cap", size=10, bold=True)
        cell_set(tb.rows[0].cells[4], "Stringency (0–7)", size=10, bold=True)
        print("Table 4 헤더 명료화 완료")
        break

# ---------- ④ Fig 4(fig_daynight.png, index 3) 재임베드 ----------
blob = open(os.path.join(FIGD, "fig_daynight.png"), "rb").read()
wpx, hpx = struct.unpack(">II", blob[16:24])
sh = doc.inline_shapes[3]
rId = sh._inline.graphic.graphicData.pic.blipFill.blip.embed
doc.part.related_parts[rId]._blob = blob
sh.width = Inches(6.3)
sh.height = Emu(int(round(int(sh.width) * hpx / wpx)))
print(f"Fig 4 재임베드: {sh.width.inches:.2f} x {sh.height.inches:.2f}")

doc.save(OUT)
print("\n저장:", OUT)

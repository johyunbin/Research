# -*- coding: utf-8 -*-
# v9 3단계(조립): 그림 11→9 교체(구 Fig4·Fig9 삭제)·캡션 전면 재작성·Highlights 삽입 →
# 최종 Manuscript_타임코드.docx + Supplementary_타임코드.docx 생성.
import os, re, sys
from datetime import datetime, timezone, timedelta
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_BREAK
sys.path.insert(0, os.path.dirname(__file__))
import docxtools_v9 as T

SRC = sys.argv[1] if len(sys.argv) > 1 else os.path.join(T.ROOT, "_claude", "step2_tables.docx")
ts = datetime.now(timezone(timedelta(hours=9))).strftime("%Y%m%d_%H%M%S")
OUT = os.path.join(T.ROOT, "01_논문작업", f"Manuscript_{ts}.docx")
OUT_SUPP = os.path.join(T.ROOT, "01_논문작업", f"Supplementary_{ts}.docx")

doc = Document(SRC)

# ================= 1) 그림 교체/삭제 =================
# 구 이미지 순서(1-11) → (새 PNG, 폭 in) 또는 None=삭제
PLAN = [
    ("fig1_study_area.png", 6.3), ("fig2_identification.png", 6.5), ("fig3_doseresponse.png", 6.0),
    None,                                            # 구 Fig 4 (daynight) — Fig 3b로 흡수
    ("fig4_phase_maps.png", 5.6), ("fig5_trajectory.png", 6.5), ("fig6_did_eventstudy.png", 6.5),
    ("fig7_spatial_null.png", 6.5),
    None,                                            # 구 Fig 9 (상세) — Supplementary S1
    ("fig8_drift_validation.png", 6.2), ("fig9_robustness.png", 6.0),
]
shapes = list(doc.inline_shapes)
assert len(shapes) == len(PLAN), f"이미지 {len(shapes)}개(11 기대)"
for shape, plan in zip(shapes, PLAN):
    p_el = shape._inline.getparent().getparent().getparent()   # wp:inline→w:drawing→w:r→w:p
    assert p_el.tag.endswith('}p')
    if plan is None:
        p_el.getparent().remove(p_el)
    else:
        fn, w = plan
        T.replace_image(doc, shape, os.path.join(T.FIGD, fn), width_in=w)

# ================= 2) 캡션 재작성 (구번호 → 신번호·신규 문안) =================
# 먼저 삭제된 그림의 구 캡션 제거 (신규 "Fig. 4." 캡션 작성 전에 수행해야 충돌 없음)
for gone in ("Fig. 4.", "Fig. 9."):
    hits = [p for p in doc.paragraphs if p.text.strip().startswith(gone)]
    assert len(hits) == 1, f"구 캡션 {len(hits)}건: {gone}"
    T.delete_el(hits[0])

CAPS = {
 "Fig. 1.": "Fig. 1. Study area and the S-DoT sensing network. (a) 1,165 geolocated S-DoT noise sensors "
   "(1,123 enter the analysis panel after quality control) across the 421 administrative neighbourhoods "
   "(dong) of Seoul, coloured by neighbourhood activity type (terciles of the daytime/night-time de-facto "
   "population ratio: commercial, mixed, residential; Supplementary Fig. S3). (b) Baseline daytime noise "
   "level (post-lifting mean Lday, July 2022 to December 2023) by neighbourhood.",
 "Fig. 2.": "Fig. 2. Identification strategy. (a) Conceptual structure: graded social distancing shifts "
   "neighbourhood mobility, whose effect on urban noise (solid arrow) must be separated from weather and "
   "calendar confounding and from sensor-side artefacts (calibration offset and multi-year drift; dashed "
   "arrows). (b) Fixed-effects ladder: sensor fixed effects remove each sensor's calibration offset; date "
   "fixed effects remove all date-common factors (weather, calendar, city-wide trend, network-wide drift); "
   "the remaining identifying variation is the same-day difference in mobility across neighbourhoods.",
 "Fig. 3.": "Fig. 3. Mobility dose-response of urban noise (two-way sensor + date fixed effects). "
   "(a) Segment-specific estimates (markers = β, bars = 95% CI, dong-clustered SEs) by outcome, land use, "
   "day type and season; filled markers survive Benjamini-Hochberg FDR correction (FDR < 0.05), open blue "
   "markers are nominally significant (p < 0.05), grey markers are not significant; the right-hand column "
   "lists β [95% CI]. (b) Sensor-FE versus two-way FE estimates: the night-time association under sensor "
   "FE alone (+2.28) loses significance once date fixed effects absorb common time-varying confounds "
   "(+0.50, ns), whereas the daytime estimate survives (+0.65); note that the wide night-time CI overlaps "
   "the daytime estimate, so a day-night difference in the coefficients is not itself established.",
 "Fig. 5.": "Fig. 4. Spatiotemporal evolution of the mobility dose across four graded distancing phases "
   "(phase-window averages of neighbourhood daytime de-facto population relative to the post-lifting "
   "baseline): (a) December 2020, Level 2.5 with the 5-person gathering ban; (b) July 2021, capital-area "
   "Level 4 (strongest); (c) November 2021, with-COVID relaxation; (d) March 2022, the weeks before full "
   "lifting. The city-wide median stays near 1.0 in every phase, but central commercial/business "
   "neighbourhoods empty by up to about 35% under the strongest restrictions (blue) while residential "
   "neighbourhoods fill (red), and the contrast fades through (c) and (d).",
 "Fig. 6.": "Fig. 5. Functional differentiation of mobility and noise over time. (a) Weekly daytime "
   "mobility by neighbourhood activity type: commercial neighbourhoods fall below and residential "
   "neighbourhoods rise above the post-lifting baseline during restrictions (shaded bands, stringency >= 4), "
   "converging after the April 2022 lifting. (b) Daytime noise relative to the city-wide mean (drift-robust, "
   "4-week moving average) by activity type.",
 "Fig. 7.": "Fig. 6. High- versus low-mobility-loss neighbourhood trajectories (descriptive). Weekly "
   "difference between the two groups in (a) within-sensor ΔLday and (b) daytime mobility; shading shows "
   "pointwise 95% confidence bands from dong-clustered standard errors (groups treated as independent), and "
   "grey bands mark strong-restriction periods. Date-common drift, season and city-wide trends cancel in "
   "the difference. Because groups are defined from realised mobility during the Level-4 period and both "
   "series are normalised to the post-lifting window, near-zero differences after lifting are partly built "
   "in; the panel is read as a descriptive trajectory rather than a formal event study. The noise gap turns "
   "negative during strong restrictions, tracking the mobility gap (weekly correlation r = +0.44, ρ = +0.46).",
 "Fig. 8.": "Fig. 7. No robust long-run spatial gradient. (a) Neighbourhood daytime mobility reduction "
   "(distancing-era average; peak-phase reductions are larger, cf. Fig. 4). (b) Drift-removed relative noise "
   "change over the same period. (c) Neighbourhood-level association (dongs with >= 2 sensors; point size "
   "proportional to sensor count; Theil-Sen fit): the robust correlation is near zero overall, and the "
   "apparently positive Pearson correlation among commercial dongs is driven by a few extreme "
   "neighbourhoods (Supplementary Fig. S1).",
 "Fig. 10.": "Fig. 8. Sensor drift and calibrated-network comparison. (a) Annual network-mean level change "
   "from 2020: the S-DoT network declines by 2.3 dB over 2020-2023 while calibrated environmental-noise "
   "stations (automatic daily road stations; manual quarterly general and road stations) stay flat or rise "
   "– a pattern consistent with a sensor-drift artefact rather than a real citywide quieting. (b) Nearby "
   "S-DoT sensors read on average 11.7 dB below calibrated daytime LAeq (n = 60 station pairs within 500 m; "
   "non-simultaneous years, 2022 S-DoT vs 2024 survey; larger at roadside), so absolute S-DoT levels cannot "
   "be interpreted as standard noise indicators (Supplementary Fig. S2).",
 "Fig. 11.": "Fig. 9. Robustness and sensitivity. (a) Permutation sensitivity check: distribution of the "
   "dose-response coefficient when the dong-level mobility dose is reshuffled across neighbourhoods within "
   "each date and broadcast to all sensors in the dong (300 shuffles); the actual estimate (+0.65, vertical "
   "line on the broken axis) lies far outside the null (two-sided p = (B+1)/(N+1) = 0.003). (b) Binned "
   "dose-response after two-way demeaning: decile means (error bars = 95% CI, dong-clustered) lie close to "
   "the linear fit, supporting the linear approximation.",
}
for old_prefix, newtext in CAPS.items():
    hits = [p for p in doc.paragraphs if p.text.strip().startswith(old_prefix)]
    assert len(hits) == 1, f"캡션 {len(hits)}건: {old_prefix}"
    T.fill_para(hits[0], newtext)

# ================= 3) Highlights 페이지 삽입 (ABSTRACT 앞) =================
abs_p = T.find(doc, "ABSTRACT")
HIGH = [
    "Graded distancing, measured mobility and 1,123 IoT sensors give a dose–response",
    "A 30% fall in neighbourhood mobility is tied to only ~0.23 dB less daytime noise",
    "Only daytime noise shows a detectable association, consistent in sign across seasons",
    "The 2.3 dB multi-year S-DoT decline is consistent with drift, not quieter streets",
    "Within-sensor, same-day designs are essential for low-cost noise networks",
]
new_els = []
ph = doc.add_paragraph(); T.fill_para(ph, "Highlights", bold=True); new_els.append(ph._element)
for h in HIGH:
    p = doc.add_paragraph(); T.fill_para(p, "• " + h)
    p.paragraph_format.left_indent = Inches(0.25); p.paragraph_format.first_line_indent = Inches(-0.25)
    new_els.append(p._element)
pb2 = doc.add_paragraph(); r = pb2.add_run(); T.setfont(r, 12); r.add_break(WD_BREAK.PAGE)
new_els.append(pb2._element)
for el in new_els:
    abs_p._element.addprevious(el)

doc.save(OUT)

# ================= 4) 검증 스캔 =================
doc2 = Document(OUT)
n_img = len(doc2.inline_shapes)
tbls = [b for b in T.iter_blocks(doc2) if hasattr(b, "rows") and not hasattr(b, "runs")]
caps = [p.text[:8] for p in doc2.paragraphs if re.match(r"^Fig\. \d+\.", p.text.strip())]
body_refs = sorted(set(int(m) for p in doc2.paragraphs
                       for m in re.findall(r"Fig\. (\d+)", p.text)))
sec_marks = sum(p.text.count("§") for p in doc2.paragraphs)
tcapn = [p.text[:9] for p in doc2.paragraphs if re.match(r"^Table \d+\.", p.text.strip())]
print(f"이미지 {n_img} (9 기대) · 표 {len(tbls)} (7 기대)")
print("그림 캡션:", caps)
print("표 캡션:", tcapn)
print("본문 Fig 참조 번호:", body_refs, "(1..9만이어야)")
print("잔존 §:", sec_marks)
assert n_img == 9 and len(tbls) == 7 and max(body_refs) <= 9
print("저장:", OUT)

# ================= 5) Supplementary docx =================
sup = Document()
st = sup.styles["Normal"]; st.font.name = "Times New Roman"; st.font.size = Pt(11)
for s in sup.sections:
    s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Inches(1.0)
p = sup.add_paragraph(); T.fill_para(p, "Supplementary material", bold=True, size=14)
p = sup.add_paragraph(); T.fill_para(
    p, "The dose-response of urban noise to human mobility measured with a city-scale IoT sensor network "
       "in Seoul", size=12, bold=True)
p = sup.add_paragraph(); T.fill_para(p, "Hyun In Jo", size=11)
sup.add_paragraph()
SUPP = [
 ("figS1_landuse_detail.png", 6.3,
  "Supplementary Fig. S1. Commercial versus residential detail of the long-run spatial cross-section. "
  "(a, c) Drift-removed relative noise change for commercial and residential dongs (all classified dongs "
  "mapped). (b, d) Neighbourhood-level association between distancing-era mobility reduction and relative "
  "noise change (dongs with >= 2 sensors): the commercial Pearson correlation (+0.38) collapses to "
  "rho = +0.08 under rank-based statistics, and residential dongs show no association either."),
 ("figS2_calibration_detail.png", 6.3,
  "Supplementary Fig. S2. Absolute-level validation detail. (a) Calibrated LAeq (2024 survey) versus nearby "
  "S-DoT Lday (2022): no cross-sectional correlation (r = -0.04), reflecting sensor-specific calibration "
  "offsets. (b) Distribution of the S-DoT minus calibrated LAeq offset (mean -11.7 dB; n = 60 pairs within "
  "500 m)."),
 ("figS3_landuse_classification.png", 6.3,
  "Supplementary Fig. S3. Neighbourhood activity-type classification. (a) Commercial / mixed / residential "
  "terciles of the post-lifting daytime/night-time de-facto population ratio. (b) Distribution of the ratio "
  "with tercile cuts (0.88 / 1.00)."),
]
for fn, w, capt in SUPP:
    sup.add_picture(os.path.join(T.FIGD, fn), width=Inches(w))
    p = sup.paragraphs[-1]; p.alignment = 1
    cp = sup.add_paragraph(); T.fill_para(cp, capt, size=10)
    sup.add_paragraph()
sup.save(OUT_SUPP)
print("저장:", OUT_SUPP)

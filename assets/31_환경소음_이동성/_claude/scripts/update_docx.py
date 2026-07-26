# -*- coding: utf-8 -*-
# 사용자 수정 docx 제자리 업데이트: ①이미지 10개 교체(새 종횡비) ②초록 ≤250단어 ③참고문헌 인용순 재번호.
import os, re, struct
from datetime import datetime, timezone, timedelta
from docx import Document
from docx.shared import Pt, Emu, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = r"T:\00_BACKUP\01_한양대학교\00_연구실\00_프로젝트\★_논문화 프로젝트\31_환경소음_이동성"
FIGD = os.path.join(ROOT, "_claude", "figures")
SRC = os.path.join(ROOT, "01_논문작업", "Manuscript_v7_20260621_215727.docx")
ts = datetime.now(timezone(timedelta(hours=9))).strftime("%Y%m%d_%H%M%S")
OUT = os.path.join(ROOT, "01_논문작업", f"Manuscript_v7_{ts}.docx")
FIGS = ["fig_study_area.png", "fig_identification_concept.png", "fig_segmented_forest.png",
        "fig_daynight.png", "fig_phase_mobility_maps.png", "fig_landuse_trajectory.png",
        "fig_did_eventstudy.png", "fig_change_map.png", "fig_spatial_bylanduse.png", "fig_drift_validation.png"]

def setfont(run, size=12, bold=False, italic=False, sub=False):
    run.font.size = Pt(size); run.bold = bold; run.italic = italic; run.font.name = 'Times New Roman'
    if sub:
        run.font.subscript = True
    rpr = run._element.get_or_add_rPr(); rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rf.set(qn(a), 'Times New Roman')
    rf.set(qn('w:eastAsia'), '바탕')

SUB = [('LAeq', ('L', 'Aeq')), ('L_day', ('L', 'day')), ('L_night', ('L', 'night'))]
SRE = re.compile('(' + '|'.join(re.escape(k) for k, _ in SUB) + ')'); SMAP = dict(SUB)
def emit(p, text, size, italic=False):
    for part in SRE.split(text):
        if not part:
            continue
        if part in SMAP:
            b, sb = SMAP[part]; setfont(p.add_run(b), size, italic=italic); setfont(p.add_run(sb), size, italic=italic, sub=True)
        else:
            setfont(p.add_run(part), size, italic=italic)

doc = Document(SRC)

# ---------- ① 이미지 교체 ----------
def png_size(b):
    return struct.unpack(">II", b[16:24])
for i, shape in enumerate(doc.inline_shapes):
    blob = open(os.path.join(FIGD, FIGS[i]), "rb").read()
    wpx, hpx = png_size(blob)
    rId = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
    doc.part.related_parts[rId]._blob = blob
    if i in (3, 9):                      # Fig 4·10: 정사각 패널 2개 -> 본문 전체폭
        shape.width = Inches(6.3)
    shape.height = Emu(int(round(int(shape.width) * hpx / wpx)))
print(f"이미지 {len(doc.inline_shapes)}개 교체 (Fig4 w/h={doc.inline_shapes[3].width.inches:.1f}/{doc.inline_shapes[3].height.inches:.1f})")

# ---------- ② 초록 ≤250단어 ----------
ABS = ("How strongly does the urban acoustic environment respond to human mobility? The question matters for mobility, "
       "land-use and noise policy, but it is hard to answer because mobility rarely varies exogenously and is entangled "
       "with weather, land use and time of day. We exploit Korea's graded COVID-19 social-distancing system, which "
       "repeatedly tightened and relaxed activity over 2020-2023, as a natural experiment, paired with a city-scale "
       "low-cost IoT noise network, to estimate a measured-mobility dose-response for urban noise. We assembled 1,248,794 "
       "sensor-days from 1,123 Smart Seoul Data-of-Things (S-DoT) sensors across 421 administrative neighbourhoods and "
       "matched each sensor-day to neighbourhood daytime de-facto population. Because low-cost sensors carry calibration "
       "offsets and multi-year drift, we identify the effect only from within-sensor variation and, in the headline "
       "specification, from cross-neighbourhood variation within each date (two-way fixed effects). Daytime noise rises "
       "and falls with daytime mobility (beta = +0.65 dB per log-unit; a 30% mobility reduction corresponds to about "
       "0.23 dB). The effect is robust and consistent across seasons but specific to daytime: a nighttime association "
       "vanishes under two-way fixed effects, and the day-night gap does not respond. No robust gradient survives in the "
       "long-run spatial cross-section. Official roadside LAeq stations confirm the network tracks real trends while "
       "exposing that multi-year absolute levels are confounded. The dose-response is statistically robust but modest, "
       "qualifying the binary-lockdown literature, and the analysis doubles as a design guide for IoT-based urban "
       "environmental monitoring.")
print(f"초록 단어수: {len(ABS.split())}")
for p in doc.paragraphs:
    if p.text.startswith("How strongly does"):
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        emit(p, ABS, 12)
        break

# ---------- ②.5 미인용 ref 본문 인용 추가 (12 axelsson·15 kang19 soundscape, 32 paez mobility) ----------
ADD = [("[9,10,11,13]", "[9,10,11,12,13,15]"), ("[31,33,34,35]", "[31,32,33,34,35]")]
for p in doc.paragraphs:
    for r in p.runs:
        for old, new in ADD:
            if old in r.text:
                r.text = r.text.replace(old, new)

# ---------- ③ 참고문헌 인용순 재번호 ----------
paras = doc.paragraphs
ref_start = next(i for i, p in enumerate(paras) if p.text.strip() == "References")
body = paras[:ref_start]
refs = [p for p in paras[ref_start + 1:] if re.match(r'^\[\d+\]', p.text)]
CITE = re.compile(r'\[(\d+(?:\s*[,–\-]\s*\d+)*)\]')
order, seen = [], set()
for p in body:
    for m in CITE.finditer(p.text):
        for tok in re.split(r'[,–\-]', m.group(1)):
            if tok.strip().isdigit():
                n = int(tok.strip())
                if n not in seen:
                    seen.add(n); order.append(n)
remap = {old: i + 1 for i, old in enumerate(order)}
allnums = [int(re.match(r'^\[(\d+)\]', p.text).group(1)) for p in refs]
uncited = [n for n in sorted(allnums) if n not in remap]
nxt = len(order)
for n in uncited:
    nxt += 1; remap[n] = nxt
print(f"인용된 ref {len(order)}개, 미인용 {len(uncited)}개: {uncited}")

def remap_tok(m):
    out = ""
    for tok in re.split(r'(\s*[,–\-]\s*)', m.group(1)):
        out += str(remap[int(tok)]) if tok.strip().isdigit() else tok
    return "[" + out + "]"
for p in body:
    for r in p.runs:
        if "[" in r.text and CITE.search(r.text):
            r.text = CITE.sub(remap_tok, r.text)

# 재번호 + 재정렬
new_to_el = {}
for p in refs:
    old = int(re.match(r'^\[(\d+)\]', p.text).group(1)); new = remap[old]
    p.runs[0].text = re.sub(r'^\[\d+\]', f'[{new}]', p.runs[0].text, count=1)
    new_to_el[new] = p._p
parent = refs[0]._p.getparent()
for p in refs:
    parent.remove(p._p)
cur = paras[ref_start]._p
for new in sorted(new_to_el):
    cur.addnext(new_to_el[new]); cur = new_to_el[new]

doc.save(OUT)
print(f"\n저장: {OUT}")

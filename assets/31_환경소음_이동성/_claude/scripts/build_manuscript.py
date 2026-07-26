# -*- coding: utf-8 -*-
# Manuscript .docx 생성 (데이터 출처/provenance 중심). python-docx.
# 출처는 아래 구조화 데이터에서 생성 → 갱신 쉬움. 출력=01_논문작업\Manuscript_v1_20260621.docx
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = r"T:\00_BACKUP\01_한양대학교\00_연구실\00_프로젝트\★_논문화 프로젝트\31_환경소음_이동성\01_논문작업\Manuscript_v1_20260621.docx"

doc = Document()
n = doc.styles['Normal']; n.font.name='Arial'; n.font.size=Pt(10.5)

def H(t, lvl=1):
    return doc.add_heading(t, level=lvl)
def P(t, italic=False, size=10.5):
    p=doc.add_paragraph(); r=p.add_run(t); r.italic=italic; r.font.size=Pt(size); return p
def prov(rows, w0=1.7, w1=4.6):
    tbl=doc.add_table(rows=len(rows), cols=2); tbl.style='Table Grid'
    for i,(k,v) in enumerate(rows):
        cells=tbl.rows[i].cells
        rk=cells[0].paragraphs[0].add_run(k); rk.bold=True; rk.font.size=Pt(9.5)
        rv=cells[1].paragraphs[0].add_run(v); rv.font.size=Pt(9.5)
        cells[0].width=Inches(w0); cells[1].width=Inches(w1)
    doc.add_paragraph()
    return tbl

# ---------- Title ----------
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run("Mobility-driven changes in the urban acoustic environment during graded social distancing in Seoul, 2020-2023")
r.bold=True; r.font.size=Pt(15)
sub=doc.add_paragraph(); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER
rs=sub.add_run("Working manuscript (v1, 2026-06-21) - data-provenance scaffold. Body in Korean, headings/Abstract in English (per author workflow).")
rs.italic=True; rs.font.size=Pt(9)
au=doc.add_paragraph(); au.alignment=WD_ALIGN_PARAGRAPH.CENTER
au.add_run("Hyun In Jo, et al. [공저자/소속 추후 확정]   *Corresponding author").font.size=Pt(9)
doc.add_paragraph()

# ---------- Abstract ----------
H("Abstract", 1)
P("[초고 작성 예정] 코로나-19 시기 서울의 단계적 사회적 거리두기를 자연실험으로 삼아, 고밀도 도시데이터 센서망(S-DoT, 약 1,100점)의 센서-내(within-sensor) 소음 변화와 실측 이동량 사이의 용량-반응(dose-response)을 추정한다. 분석창 2020-04~2023.")
P("Keywords: urban noise; environmental acoustics; COVID-19 mobility; dose-response; IoT sensor network; Seoul; natural experiment", italic=True, size=9.5)

# ---------- 1. Introduction ----------
H("1. Introduction", 1)
P("[서론 작성 예정] 핵심 노벨티 4축: (1) 한국 단계적 거리두기 = graded dose-response(서구 binary lockdown과 차별) (2) 측정 이동량-소음 정량 결합(선행연구 미개척) (3) 고밀도 IoT 센서망 (4) 미출판 지리(서울/동아시아). 서울/한국 COVID-소음 음향 측정 논문은 주요 SCIE 범위에서 미출판(확인).")

# ---------- 2. Data and Materials (PROVENANCE) ----------
H("2. Data and Materials", 1)
P("본 절은 모든 데이터의 출처/획득경로/버전/라이선스를 재현 가능하도록 기록한다. 모든 원시데이터는 공개 공공데이터이며, 다운로드 메커니즘을 명시한다. 검증 시각 2026-06-21 KST.")

H("2.1 Study area and period", 2)
P("서울특별시 25개 자치구 전역. 분석 기간 2020-04-01 ~ 2023-12-31 (S-DoT 소음 데이터 개시일 2020-04-01로 하한 강제; 거리두기 해제 2022-04-18 이후 정상기를 무제한 zero-dose 기준선으로 활용).")

H("2.2 Outcome - Urban noise (S-DoT)", 2)
prov([
 ("정식명","스마트서울 도시데이터 센서(S-DoT) 환경정보"),
 ("영문명","Smart Seoul Data of Things (S-DoT), environmental data"),
 ("제공기관","서울특별시 (서울 열린데이터광장)"),
 ("데이터셋","OA-15969 (공공데이터포털 미러 15061244)"),
 ("페이지 URL","https://data.seoul.go.kr/dataList/OA-15969/S/1/datasetView.do"),
 ("다운로드","POST https://datafile.seoul.go.kr/bigfile/iot/inf/nio_download.do?useCache=false  | body: infId=OA-15969&seq=<YEAR>&infSeq=3 (seq=2020/2021/2022/2023, 무인증)"),
 ("기간(분석)","2020-04-01 ~ 2023-12-31 (연도별 zip)"),
 ("시간 해상도","시간별(hourly)"),
 ("공간 해상도","점(센서) 약 1,100개, 서울 25개 구 전수 (위치파일 location.xlsx, 1,170점 WGS84)"),
 ("핵심 변수","소음(dB). 주의: 2020-2022 단일 컬럼 '소음(dB)'; 2023 스키마 변경(22->58컬럼)으로 '소음 최대/평균/최소(dB)' 분할 -> 본 연구는 '소음 평균(dB)' 사용"),
 ("측정 성격","저가 IoT 센서 광대역 dB; LAeq(A-가중 등가소음도) 명시 없음 -> 절대값 아닌 within-sensor 상대변화로 사용, 공식 LAeq 측정소로 교차검증(2.6)"),
 ("라이선스","누구나 다운로드/사적목적(논문 작성 등) 활용 명시 허용; 간이센서 값으로 서울시 공식입장 아님 고지(명세서)"),
 ("본 연구 처리","cp949 파싱; 시각=등록일자(전송시간은 과학표기 손상); 시리얼->location.xlsx 보정시리얼(col1) 좌표조인(~100%); 시간->일 에너지평균(Leq24/주간06-21/야간22-05); QC >=12h, 20-95dB"),
])

H("2.3 Exposure - Mobility (dose)", 2)
P("주(主) dose = 측정 이동량 (거리두기 '단계'는 영업/모임 규제라 이동량과 부분 연동/내생성 -> 보조 scaffold). 다운로드 메커니즘/seq/라이선스/인코딩 = 실측 검증 완료(2026-06-21).")
prov([
 ("공통 다운로드","POST datafile.seoul.go.kr/bigfile/iot/inf/nio_download.do?useCache=false | infId=<OA>&seq=<N>&infSeq=<페이지값> (무인증) | infSeq: 지하철 1 / 버스 3 / 집계구 1 / 행정동 3"),
 ("지하철","서울교통공사_역별 일별 시간대별 승하차인원 | OA-12921 | 서울교통공사 | 일별x시간대x역(1-8호선) | seq 37/38/39/45(2020-23) | CP949 | 스키마 2종(2020-21 25컬럼 / 2022-23 26컬럼) | 라이선스 공공누리 제3유형(출처표시+변경금지)"),
 ("버스","서울시 버스노선별 정류장별 시간대별 승하차 | OA-12913 | 서울시(교통카드 정산) | 월별x시간대x노선x정류장(일별 없음) | seq 35/50/66(zip)/67-78(2023 월별) | CP949 | 공공누리 제1유형"),
 ("생활인구 행정동","행정동 단위 서울 생활인구(내국인) | OA-14991 | 서울특별시 | 일별x시간대x행정동 | 반기ZIP seq 2219-2224(2020-22) + 월별 2301-2312(2023) | UTF-8-BOM | 공공누리 제1유형"),
 ("생활인구 집계구","집계구 단위 서울 생활인구 | OA-14979 | 서울특별시 | 일별x시간대x집계구 | 월ZIP seq=YYMM(2004~2312) | 대용량(파일당 ~수백MB-1.45GB) | CP949 | 공공누리 제1유형"),
 ("본 연구 dose 선택","주 dose = 생활인구 행정동(동 단위 일별 체류, 패널과 공간정합) + 지하철 승하차(구 단위 교통량). 버스=월별 집계라 보조, 집계구=대용량으로 행정동 우선."),
])

H("2.4 Dose scaffold - Social distancing", 2)
prov([
 ("정식명","사회적 거리두기 시행연혁"),
 ("제공기관/출처","질병관리청; HWPX = 공공데이터포털 15106451 (https://www.data.go.kr/data/15106451/fileData.do)"),
 ("교차확인","korea.kr / mohw.go.kr / 서울시 보도자료(수도권 국면별 일자/영업시간/인원)"),
 ("본 연구 처리","ordinal '단계' 직접 사용 지양(체계 2차 재정의/비등간/내생성) -> 준연속 재코딩(영업 종료시각/허용 모임인원). 수도권/서울 기준. 주의: 최종 코딩은 HWPX 원본 대조 필요"),
])

H("2.5 Covariates - Meteorology", 2)
prov([
 ("정식명","종관기상관측(ASOS) 일자료 / 방재기상관측(AWS)"),
 ("제공기관/출처","기상청 기상자료개방포털 (data.kma.go.kr); 서울 종관지점(108) 등"),
 ("변수","기온/강수/풍속/습도(일/시간) - 소음 교란 통제"),
 ("상태","획득 예정(빌드 Phase1)"),
])

H("2.6 Validation reference data", 2)
prov([
 ("(1) 도로교통소음 4점","서울시 도로교통 소음 측정정보(시청/신사/신촌/성수) | 서울 보건환경연구원 | OA-15473 | 시간별 LEQ(=LAeq, dB(A)), 주간/야간 | 2021-2024 | 시간 within-sensor 검증"),
 ("(2) 환경소음 146점","한국환경공단_환경소음 측정망 정보 | data.go.kr 15065396 | 분기별 LAeq(주/야) | 2019/2021-2024(OpenAPI; 2020 누락) | 공간 검증(좌표는 지오코딩 67점)"),
 ("자동측정망 좌표","한국환경공단_환경소음 측정망 지점정보 | data.go.kr 15093409 | 전국 546점 위경도(서울 9)"),
 ("검증 결과(스파이크)","시간 within-sensor r 0.75-0.92(일변동) / 0.60-0.75(일/주) = 강 / 공간 cross-sensor 절대레벨 r~0 = 약 -> 센서별 오프셋 제각각 -> 상대변화/집계 설계 잠금"),
])

H("2.7 Processing and quality control", 2)
P("S-DoT: 시간->일 에너지평균; 일최소 12시간 / 20-95dB QC; 센서 내 상대변화(자기 baseline 대비); 불량 센서 스크리닝; 2020-04~05 가동초기 결측 -> 약 2020-05 이후 사용. 산출물 sensor-days = 1,253,298.")

# ---------- 3. Methods ----------
H("3. Methods", 1)
P("[작성 예정] (1) 선형혼합효과 dose-response: dLAeq ~ mobility + 기상 + 요일 + 계절 + (1|sensor) + (1|dong) (2) 거리두기 전환 event-study (3) 변화량 공간패턴 지도 (4) 공식 LAeq 4점 교차검증 / 센서 드리프트 감시.")

# ---------- 4-5 placeholders ----------
H("4. Results", 1); P("[분석 후 작성]")
H("5. Discussion", 1); P("[작성 예정] 한계: S-DoT 비검교정 / 절대레벨 공간신뢰 / 센서 드리프트 / 1dB 해상도 / 거리두기 != 이동량 -> 선제 방어.")

# ---------- Data Availability ----------
H("Data Availability Statement", 1)
P("모든 원시데이터는 공개 공공데이터이다. S-DoT 환경정보(서울 OA-15969), 지하철/버스/생활인구(서울 OA-12921/12913/14979/14991), 거리두기 시행연혁(질병청/공공데이터포털 15106451), 기상(기상청 ASOS), 검증기준(서울 OA-15473, 환경공단 15065396/15093409). 상세 URL/다운로드 메커니즘/버전은 2절 참조. 분석 코드/처리 스크립트는 [저장소 TBD].")

# ---------- References ----------
H("References", 1); P("[논문 작성 시 Vancouver 번호식]")

doc.save(OUT)
print("saved:", OUT)
print("paragraphs:", len(doc.paragraphs), "tables:", len(doc.tables))

# -*- coding: utf-8 -*-
# Manuscript v7 — 심화 재구성: 서론 5소제목(심화) · Materials+Methods 통합(2.1-2.6) · Results · Discussion 5소제목(심화).
# Elsevier(SCS/LUP) 양식 · 테이블 6 · 그림 10 · 인용 51 · 아래첨자.
import os, re, json
from datetime import datetime, timezone, timedelta
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ts = datetime.now(timezone(timedelta(hours=9))).strftime("%Y%m%d_%H%M%S")
ROOT = r"T:\00_BACKUP\01_한양대학교\00_연구실\00_프로젝트\★_논문화 프로젝트\31_환경소음_이동성"
FIGD = os.path.join(ROOT, "_claude", "figures"); PD = os.path.join(ROOT, "data", "processed")
OUT = os.path.join(ROOT, "01_논문작업", f"Manuscript_v7_{ts}.docx")
REFS = json.load(open(r"C:\Users\wh850\AppData\Local\Temp\refs_final.json", encoding="utf-8"))

doc = Document()
s = doc.sections[0]; s.page_width, s.page_height = Mm(210), Mm(297)
s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Inches(1.0)
nm = doc.styles['Normal']; nm.font.name = 'Times New Roman'; nm.font.size = Pt(12); nm.paragraph_format.line_spacing = 2.0

def setfont(run, size=12, bold=False, italic=False, sub=False, latin='Times New Roman', ea='바탕'):
    run.font.size = Pt(size); run.bold = bold; run.italic = italic; run.font.name = latin
    if sub:
        run.font.subscript = True
    rpr = run._element.get_or_add_rPr(); rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rf.set(qn(a), latin)
    rf.set(qn('w:eastAsia'), ea)

SUB = [('ΔLAeq', ('ΔL', 'Aeq')), ('ΔL_day', ('ΔL', 'day')), ('ΔL_night', ('ΔL', 'night')),
       ('Leq24', ('L', 'eq,24h')), ('LAeq', ('L', 'Aeq')), ('L_day', ('L', 'day')),
       ('L_night', ('L', 'night')), ('L_dn', ('L', 'dn')), ('Leq', ('L', 'eq'))]
SRE = re.compile('(' + '|'.join(re.escape(k) for k, _ in SUB) + ')'); SMAP = dict(SUB)

def emit(p, text, size, bold=False, italic=False):
    for part in SRE.split(text):
        if not part:
            continue
        if part in SMAP:
            b, sb = SMAP[part]
            setfont(p.add_run(b), size, bold, italic)
            setfont(p.add_run(sb), size, bold, italic, sub=True)
        else:
            setfont(p.add_run(part), size, bold, italic)

def P(t, size=12, bold=False, italic=False, indent=True, ls=2.0, align='J'):
    p = doc.add_paragraph(); emit(p, t, size, bold, italic)
    pf = p.paragraph_format; pf.line_spacing = ls
    if indent:
        pf.first_line_indent = Inches(0.2)
    pf.alignment = {'J': AL.JUSTIFY, 'L': AL.LEFT, 'C': AL.CENTER}[align]
    return p

def H(t, lvl=1):
    p = doc.add_paragraph(); setfont(p.add_run(t), 12.5 if lvl == 1 else 11.5, bold=True)
    pf = p.paragraph_format; pf.line_spacing = 1.5; pf.space_before = Pt(12 if lvl == 1 else 8); pf.space_after = Pt(2)
    return p

def figure(fn, num, cap, width=6.0):
    doc.add_picture(os.path.join(FIGD, fn), width=Inches(width)); doc.paragraphs[-1].alignment = AL.CENTER
    c = doc.add_paragraph(); setfont(c.add_run(f"Fig. {num}. "), 10.5, bold=True); emit(c, cap, 10.5)
    c.paragraph_format.line_spacing = 1.5; c.paragraph_format.space_after = Pt(6)

def mtable(num, title, header, rows, widths):
    c = doc.add_paragraph(); setfont(c.add_run(f"Table {num}. "), 10.5, bold=True); emit(c, title, 10.5)
    c.paragraph_format.line_spacing = 1.5
    tbl = doc.add_table(rows=1 + len(rows), cols=len(header)); tbl.style = 'Table Grid'
    for j, h in enumerate(header):
        cell = tbl.rows[0].cells[j]; emit(cell.paragraphs[0], h, 9, bold=True); cell.width = Inches(widths[j])
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            cell = tbl.rows[i + 1].cells[j]; emit(cell.paragraphs[0], str(v), 9); cell.width = Inches(widths[j])
    doc.add_paragraph()

def star(p):
    return "***" if p < .001 else "**" if p < .01 else "*" if p < .05 else ""

# ============ Title / Abstract ============
t = doc.add_paragraph(); setfont(t.add_run(
    "The dose-response of urban noise to human mobility measured with a city-scale IoT sensor network in Seoul"), 14, bold=True)
t.paragraph_format.line_spacing = 1.5
sub = doc.add_paragraph(); setfont(sub.add_run(
    f"Working manuscript ({ts[:8]}). Body in Korean; title, headings, Abstract and captions in English. "
    "Target: Sustainable Cities and Society / Landscape and Urban Planning. References in numbered (Vancouver) style."), 9, italic=True)
au = doc.add_paragraph(); setfont(au.add_run("Hyun In Jo, et al. [공저자/소속 추후 확정]   *Corresponding author"), 10)
doc.add_paragraph()

H("Abstract", 1)
P("How strongly does the urban acoustic environment respond to human mobility? The question is central to mobility, "
  "land-use and noise policy, yet it is hard to answer because mobility rarely varies exogenously and is entangled with "
  "weather, land use and time of day. We exploit Korea's graded COVID-19 social-distancing system - which repeatedly "
  "tightened and relaxed activity over 2020-2023 - as a natural experiment, and pair it with a city-scale low-cost IoT "
  "noise network to estimate a measured-mobility dose-response for urban noise. We assembled 1,248,794 sensor-days from "
  "1,123 Smart Seoul Data-of-Things (S-DoT) sensors across 421 administrative neighbourhoods (dong) and matched each "
  "sensor-day to neighbourhood-level daytime de-facto population as a continuous exposure. Because low-cost sensors carry "
  "idiosyncratic calibration offsets and multi-year drift, we identify the effect only from within-sensor variation and, "
  "in the headline specification, from cross-neighbourhood variation within each date (two-way sensor-and-date fixed "
  "effects). Daytime noise rises and falls with daytime mobility (beta = +0.65 dB per log-unit, 95% CI 0.17-1.13, "
  "p = 0.008): a 30% mobility reduction corresponds to roughly a 0.23 dB reduction in daytime noise. The effect is robust "
  "and consistent across all four seasons but specific to daytime - a nighttime association apparent under sensor fixed "
  "effects vanishes under two-way fixed effects, and the day-night gap does not respond. No robust mobility-noise gradient "
  "survives in the long-run spatial cross-section once outlier-resistant statistics and a reliability filter are applied. "
  "Official roadside LAeq stations confirm that the network tracks real trends while exposing that multi-year absolute "
  "levels are confounded. The mobility-noise dose-response is statistically robust but modest, qualifying the headline "
  "magnitudes of the binary-lockdown literature, and the analysis doubles as a design guide for IoT-based urban "
  "environmental monitoring.", indent=False)
kp = doc.add_paragraph(); setfont(kp.add_run("Keywords: "), 11, bold=True, italic=True)
setfont(kp.add_run("urban noise; environmental acoustics; human mobility; dose-response; IoT sensor network; "
                   "smart-city monitoring; de-facto population; Seoul; natural experiment"), 11, italic=True)
kp.paragraph_format.line_spacing = 1.5

# ============ 1. Introduction ============
H("1. Introduction", 1)
H("1.1. Environmental noise as a sustainability and public-health challenge", 2)
P("환경소음은 대기오염에 이어 도시 환경이 인간 건강에 부과하는 두 번째로 큰 부담으로 평가된다. 세계보건기구(WHO)는 환경소음을 "
  "유럽 도시의 주요 질병부담 원인으로 규정하며, 서유럽에서만 매년 100만 healthy life-years 이상이 교통소음으로 손실되는 것으로 "
  "추정한다 [1,2]. 소음의 건강영향은 단순한 불쾌를 넘어선다 — 야간소음은 수면을 단편화하고 [8], 만성적 노출은 자율신경·내분비 "
  "스트레스 경로를 통해 고혈압·허혈성 심질환 등 심혈관계 질환 위험을 높이며 [3,4], 주간에는 소음성가심(annoyance)과 인지수행 "
  "저하를 유발한다 [6,7]. 이러한 영향은 막대한 사회·경제적 비용으로 환산된다 [5].")
P("따라서 소음은 도시의 지속가능성과 거주적합성(livability)을 좌우하는 핵심 물리적 스트레서이며, 그 부담이 사회경제적 약자에게 "
  "불균등하게 분포한다는 점에서 환경정의(environmental justice)의 문제이기도 하다. 소음을 단순 데시벨이 아니라 맥락 속 인간 경험으로 "
  "다루는 사운드스케이프(soundscape) 패러다임 [9,10,11,13]은 이 인식을 확장해, 도시 공공공간의 음향쾌적성 [14]과 토지이용·도시형태에 "
  "따른 음환경의 시공간 변이 [16]를 연구 대상으로 끌어올렸다. 그러나 소음을 '관리'하려면 한 가지 근본 질문에 답해야 한다 — 도시소음은 "
  "계획가가 바꿀 수 있는 것, 즉 인간의 활동과 이동량에 얼마나 민감하게 반응하는가? 이 반응성(responsiveness)이야말로 모든 이동·"
  "토지이용 기반 소음정책의 효과를 가늠하는 출발점이지만, 놀랍게도 정량적으로 규명된 바가 거의 없다.")
H("1.2. Mobility as the dominant yet under-quantified driver of urban noise", 2)
P("도시소음의 지배적 원천은 도로교통이며, 소음 수준은 교통량과 함께 증가한다 [43]. 그러나 '교통량'을 넘어선 '도시 전체의 인간 활동·"
  "이동량'과 소음의 관계는 인과적으로 측정된 적이 드물다. 그 이유는 식별(identification)의 어려움에 있다. 정상적인 도시에서 이동량은 "
  "외생적으로 변하지 않으며, 토지이용·시간대·요일·날씨 등 소음과 이동량을 동시에 좌우하는 공통 요인에 묶여 있다. 따라서 관측 자료의 "
  "단순한 이동량-소음 상관은 이 공통 요인들에 의해 교란되어, 이동량이 '원인'으로서 소음을 얼마나 변화시키는지를 말해 주지 못한다.")
P("기존의 교통소음 예측 모델 — CNOSSOS-EU 같은 표준식이나 토지이용회귀(land-use regression) 모델 [43,44] — 은 교통량·도로폭·"
  "토지이용으로부터 소음을 추정하지만 본질적으로 횡단적·정적(static)이다. 즉 '교통이 많은 곳이 시끄럽다'는 공간 패턴은 잘 재현하나, "
  "'이동량이 줄면 소음이 얼마나 주는가'라는 동적 용량-반응(dynamic dose-response)은 답하지 못한다. 이 빈틈을 메우려면 이동량의 "
  "외생적·점진적(graded) 변동과, 그에 대응하는 고밀도 소음 관측이 동시에 필요하다 — 정상적인 도시에서는 좀처럼 주어지지 않는 조건이다.")
H("1.3. COVID-19 as a natural experiment: promise and the limits of current evidence", 2)
P("코로나-19 대유행은 바로 이 조건을 전례 없이 만들어 낸 자연실험(natural experiment)이었다. 봉쇄·거리두기는 인간 이동량을 외생적으로 "
  "급감시켰고, 다수 연구가 봉쇄기 도시소음의 수 dB 감소를 보고했다 — 마드리드는 주간 약 3 dB 감소 [19], 더블린 [18], 런던 [20], "
  "스톡홀름 [21], 몬트리올 [22], 그라나다 [23], 부에노스아이레스 [24], 리마 [25], 피렌체 [49], 인도 도시들 [26,27]. 같은 시기 "
  "대기질에서도 이산화질소·이산화탄소가 급감해 [28,29,30] 봉쇄가 도시 환경 전반에 미친 충격을 입증했고, 봉쇄기 사운드스케이프 지각의 "
  "변화도 관찰되었다 [17,51].")
P("그러나 이 빠르게 축적된 문헌은 세 가지 구조적 한계를 공유하며, 이는 본 연구의 동기를 직접 규정한다. 첫째, 대부분 '봉쇄 vs 비봉쇄'의 "
  "이항적(binary) 비교로, 규제 강도가 단계적으로 달라지는 graded 용량-반응 곡선을 추정할 수 없다 — 우리는 소음이 이동량 변화에 "
  "비례적으로 반응하는지, 그 기울기가 얼마인지 알지 못한다. 둘째, 거의 모든 연구가 소음 감소를 '교통이 줄어서'라고 서술적으로 귀인할 뿐, "
  "이동량을 실제로 측정해 소음과 정량적으로 결합하지 않는다 — 노출(exposure)이 가정될 뿐 측정되지 않는다. 셋째, 측정점이 소수(대개 "
  "4~70점)이고 서구 도시에 편중되어, 고밀도·동아시아 도시의 미시적 공간 이질성과 토지이용별 차등을 포착하지 못한다. 요컨대 graded·"
  "measured·high-density·동아시아라는 네 조건을 동시에 충족하는 증거가 부재하며, 이것이 본 연구가 필요한 이유다.")
H("1.4. Measured mobility, high-density IoT sensing, and an under-addressed pitfall", 2)
P("최근 두 가지 데이터 혁신이 이 빈틈을 메울 가능성을 열었다. 첫째, 모바일 통신·신호 기반의 de-facto(체류) 인구 — 거주지가 아니라 "
  "특정 시점에 사람들이 '실제로 어디에 있는가'를 도시 미세공간·고빈도로 측정한다 [31,33,34,35]. 둘째, 저가 IoT 음향 센서망 — SONYC "
  "[37], 저가 모니터링 기기 [36], 무선음향센서망(WASN) [38,39], 참여형 NoiseCapture [50], 스마트시티 센싱 [40,41] — 은 도시 음환경을 "
  "수백~수천 점에서 상시 관측할 수 있게 했다. 한국의 단계적 거리두기, 서울 전역의 S-DoT 소음센서망(약 1,100점), 그리고 행정동 단위 "
  "생활인구의 결합은 위 네 조건을 동시에 만족시키는 드문 기회를 제공한다.")
P("그러나 저가 IoT 소음망에는 기존 COVID-소음 문헌이 대체로 간과해 온 본질적 함정이 있다 — 센서별 검교정 오프셋과 다년에 걸친 "
  "드리프트다 [42]. 저가 센서는 절대 음압을 정밀 측정하도록 검교정되지 않으며, 센서마다 고유의 상수 오프셋을 갖고, 시간이 지나며 "
  "노후·환경요인으로 측정값이 표류한다. 이를 무시하고 절대레벨을 시계열·공간으로 단순 비교하면 거짓 패턴 또는 거짓 결론에 이른다(본 "
  "연구의 §3.5가 그 실례를 보인다). 따라서 고밀도 IoT 자료를 신뢰성 있게 쓰려면 이 교란을 식별 설계에서 정면으로 제거해야 한다 — 이는 "
  "본 연구의 방법론적 핵심이자, 스마트시티 환경 데이터 활용 전반에 적용되는 교훈이다.")
H("1.5. The present study", 2)
P("본 연구는 한국의 단계적 사회적 거리두기를 graded 자연실험으로, 서울의 S-DoT IoT 소음센서를 결과변수로, 행정동 생활인구를 측정 "
  "이동량 dose로 결합해 도시소음의 이동량 용량-반응을 추정한다. 핵심적으로, 우리는 절대레벨이 아닌 센서-내 상대변화와 같은 날짜 안에서 "
  "동(洞) 간 변이만으로 효과를 식별하는 양방향 고정효과(two-way fixed effects) 전략 [46,48]과 자연실험 틀 [47]을 채택하고, 저가 IoT의 "
  "드리프트·해상도 한계를 공식 측정망으로 교차검증한다. 이로써 서울 사운드스케이프의 시공간 변이 연구 [16]를 정량적·인과적 이동량-소음 "
  "추정으로 확장한다. 세 가지 질문에 답한다 — RQ1: 측정 이동량 감소는 도시소음을 얼마나 낮추는가(용량-반응 기울기)? RQ2: 그 효과는 "
  "토지이용·주야·요일·계절에 따라 어떻게 다른가? RQ3: 거리두기 전환에서 소음은 어떻게 반등하며 어떤 공간패턴을 갖는가? 기여는 네 "
  "가지다 — (1) 측정 이동량에 대한 graded 용량-반응의 최초 정량 추정, (2) 약 1,100점 고밀도 IoT 센서망의 활용, (3) 서울/동아시아 "
  "고밀도 메가시티라는 미출판 지리, (4) 저가 IoT 소음망의 드리프트·해상도 한계를 드러내고 cross-sectional 식별의 필요성을 입증하는 "
  "방법론적 교훈.")

# ============ 2. Materials and methods ============
H("2. Materials and methods", 1)
H("2.1. Study area, period and study design", 2)
P("연구지역은 서울특별시 25개 자치구 전역이다. 서울은 약 960만 인구가 605 km²에 거주하는 세계 최고 밀도의 메가시티 중 하나로, "
  "고밀도 혼합토지이용과 조밀한 행정동(行政洞, administrative neighbourhood) 체계(421개)를 갖추어 미시적 공간 이질성을 포착하기에 "
  "이상적이다(Fig. 1). 분석 기간은 S-DoT 소음 자료의 개시일인 2020-04-01부터 2023-12-31까지로, 거리두기 전면해제(2022-04-18) 이후의 "
  "정상기를 충분히 포함해 회복(rebound)을 끝까지 추적할 수 있다. 연구 설계의 핵심은 세 가지다 — (i) 한국의 단계적 거리두기가 만든 "
  "graded·외생적 이동량 변동을 자연실험으로 활용하고, (ii) 저가 센서의 검교정 오프셋을 차분으로 상쇄하기 위해 결과변수를 센서-내 "
  "상대변화로만 사용하며, (iii) 효과를 같은 날짜 안의 동 간 변이로 식별한다(§2.6).")
figure("fig_study_area.png", 1,
       "Study area and sensor network. (a) 1,165 geolocated S-DoT noise sensors across 421 administrative neighbourhoods "
       "(dong) of Seoul, colored by activity type (daytime/nighttime living-population ratio: commercial, mixed, "
       "residential). (b) Baseline daytime noise level (post-lift mean L_day, 2022-07 to 2023-12) by neighbourhood.", 6.0)
H("2.2. Outcome: urban noise from the S-DoT network", 2)
P("결과변수는 서울시가 도시 전역에 상시 운영하는 IoT 도시데이터 센서망(Smart Seoul Data of Things, S-DoT)의 소음 자료다. 약 1,100개 "
  "지점에서 시간별 광대역 음압(dB)을 측정하며, 모든 자료는 무인증으로 공개 다운로드된다(Table 1). 우리는 시간별 값을 일(日) 에너지 "
  "등가소음도로 집약했다 — 전일 Leq24, 주간 L_day(06-21시), 야간 L_night(22-05시)을 각각 10·log10(평균(10^(L/10)))으로 계산하고, "
  "일 최소 12시간 관측·물리적 범위(20-95 dB)의 품질관리를 적용했다. 2023년 스키마 변경(소음이 최대/평균/최소로 분할)은 평균 컬럼을 "
  "선택해 처리했다. 결정적으로 S-DoT는 검교정되지 않은 광대역 dB(LAeq 라벨 없음)이므로 절대값으로 쓰지 않고 오직 센서-내 상대변화로만 "
  "사용하며, 공식 LAeq 측정망으로 교차검증한다(§3.5).")
mtable(1, "Provenance and processing of the urban-noise outcome (S-DoT).",
       ["Item", "Description"], [
        ["Dataset", "Smart Seoul Data of Things (S-DoT) environmental data; Seoul Open Data OA-15969"],
        ["Access", "POST datafile.seoul.go.kr/.../nio_download.do; infId=OA-15969&seq=<YEAR>&infSeq=3 (no authentication)"],
        ["Resolution", "Hourly broadband dB; ~1,100 point sensors; all 25 districts; 2020-04 to 2023-12"],
        ["Processing", "cp949 decoding; timestamp = 등록일자; serial->WGS84 coordinate join; hourly->daily energy means "
                       "(Leq24 / L_day 06-21 / L_night 22-05); QC >=12 h, 20-95 dB; 2023 schema uses 'mean noise (dB)'"],
        ["Use", "Uncalibrated -> within-sensor relative change only; cross-validated against official 4-point LAeq (§3.5)"],
       ], [1.2, 5.1])
H("2.3. Exposure: human mobility from de-facto population", 2)
P("이동량 노출은 서울 '생활인구(living / de-facto population)'로 측정했다. 이는 거주 인구가 아니라 모바일 통신 신호로 추정한, 특정 "
  "시점에 각 행정동에 실제로 체류하는 인구수다(Table 2). 시간별 체류수(stock)를 일 평균(및 주간 06-21 / 야간 22-05 평균)으로 집약해 "
  "424개 동 × 1,461일 = 619,464 동-일을 얻었다(결측 0). dose 구성에서 한 가지 통찰이 중요하다 — 도시 전체 생활인구 총량은 거의 "
  "보존된다(사람들이 외출을 줄여도 자택 체류로 계수되기 때문). 즉 거리두기의 신호는 총량이 아니라 공간 재분포에 있으며, 상업·업무 동은 "
  "주간 체류가 20-35% 비워지고 주거 동은 최대 37% 증가한다. 따라서 우리는 dose를 동의 자기기준 대비 상대변화로 정의한다 — "
  "lp = log(동의 일별 생활인구 / 동의 post-lift 정상기 평균). 이렇게 하면 dose는 각 동이 자기 정상에서 얼마나 벗어났는가를 나타내며, "
  "동·센서별 절대수준 차이에 영향받지 않는다.")
mtable(2, "Provenance of the mobility exposure (de-facto population) and covariates.",
       ["Item", "Description"], [
        ["Exposure", "Administrative-dong de-facto (living) population, Seoul OA-14991; hourly x dong; from mobile-network "
                     "signalling; ~2.8 GB raw (deleted after processing)"],
        ["Aggregation", "Hourly stock -> daily means (day 06-21 / night 22-05); 424 dong x 1461 d = 619,464 dong-days"],
        ["Dose", "Within-dong log relative change lp = log(daily population / dong post-lift baseline mean)"],
        ["Weather", "Open-Meteo ERA5 daily reanalysis (mean/max/min temperature, precipitation, max wind) at 37.57N 126.97E"],
        ["Calendar", "Day-of-week, weekend, Korean public holidays (incl. substitutes), season"],
       ], [1.2, 5.1])
H("2.4. Linking sensors to neighbourhoods", 2)
P("센서 메타데이터는 도로명주소만 제공하고 행정동명을 담지 않으므로, 각 센서를 좌표 기반 점-다각형 포함(point-in-polygon) 공간조인으로 "
  "행정동에 배정했다. 행정동 경계는 vuski/admdongkor의 GeoJSON(ver20220101)을 사용했고, 순수 Python ray-casting으로 다각형 포함을 "
  "판정했다(MultiPolygon·홀 처리, bbox 프리필터). 생활인구 코드체계와 경계 코드체계가 일부 동(강북·강동 구역개편)에서 어긋나는 "
  "code drift는 이름 기반 크로스워크로 해소했다. 결과적으로 1,170개 센서 중 1,165개(99.6%)가 배정되었고, 주소상 자치구와 배정된 동의 "
  "자치구 일치율도 99.6%였다.")
H("2.5. Covariates: social distancing, weather and calendar", 2)
P("거리두기는 ordinal '단계'를 직접 dose로 쓰지 않았다 — 체계가 두 차례 재정의되었고(2020-11, 2021-07), 단계 간격이 비등간이며, "
  "같은 단계라도 시기별 이동량이 다르고 자발적 감소가 내생적이기 때문이다. 대신 매일의 규제를 영업종료시각과 허용 모임인원으로부터 "
  "준연속 stringency 지수(0-7)로 재코딩했다(Table 4). 이는 보조 scaffold 공변량일 뿐 주 dose는 측정 이동량이다. 기상은 Open-Meteo "
  "ERA5 일자료(기온·강수·풍속), 달력은 요일·주말·한국 공휴일·계절을 포함한다.")
mtable(4, "Seoul capital-area social-distancing timeline and its quasi-continuous re-coding (selected regime changes).",
       ["From", "Regime", "Close (h)", "Gathering", "Stringency"], [
        ["2020-01", "Pre-COVID (normal)", "24", "none", "0"],
        ["2020-03-22", "1st intensive distancing", "21", "<=10", "4"],
        ["2020-05-06", "Daily-life distancing", "24", "none", "0"],
        ["2020-08-30", "Capital Lv.2.5 (first 21h cap)", "21", "<=50", "5"],
        ["2020-12-23", "Lv.2.5 + 5-person ban", "21", "<=4", "6"],
        ["2021-02-15", "Lv.2 + 5-person ban", "22", "<=4", "5"],
        ["2021-07-12", "Capital Level 4", "22", "<=4", "5"],
        ["2021-11-01", "With-COVID recovery", "24", "<=10", "1"],
        ["2021-12-18", "Special measures", "21", "<=4", "6"],
        ["2022-03", "Gradual easing", "23", "<=8", "3"],
        ["2022-04-18", "Full lifting", "24", "none", "0"],
       ], [1.1, 2.4, 0.9, 1.0, 1.0])
H("2.6. Statistical analysis and identification", 2)
P("저가 IoT 센서는 센서별 검교정 오프셋이 제각각이어서 센서 간 절대레벨 비교는 신뢰할 수 없으나, 한 센서 내 시간변화에서는 그 오프셋이 "
  "차분으로 상쇄된다. 따라서 모든 분석은 결과변수의 센서-내 상대변화에 기반한다. 1,248,794 sensor-days에 대해 더미변수 대신 "
  "within-transformation(집단평균 차감)으로 고정효과를 흡수하고, 표준오차는 센서(또는 동)로 군집화했다 [46]. 식별 전략은 Fig. 2에 "
  "요약한다.")
P("주력 분석은 두 단계의 고정효과 모형이다. (M1) 센서 고정효과: Leq ~ mobility + 기상(기온·기온²·강수·풍속·강수일) + 요일·주말·공휴일 "
  "[센서 FE, 센서 군집 SE]. 관측 가능한 시간교란을 통제하되 시간변이를 허용해 dose를 식별한다. (M2, 헤드라인) 센서+날짜 양방향 "
  "고정효과: Leq ~ mobility [센서 FE + 날짜 FE]. 날짜에 공통된 모든 요인 — 기상·요일·공휴일·도시 전체 추세·센서 드리프트·도시 전체 "
  "거리두기 — 이 날짜 FE에 흡수되어, 효과가 '같은 날짜에 자기 기준 대비 더 비워진 동의 센서가 더 조용해졌는가'라는 순수 cross-dong "
  "변이만으로 식별된다(Fig. 2). 따라서 도시 단일 dose(예: 지하철 총승객, 거리두기 강도)는 날짜 FE에 흡수되어 식별 불가하며, 동 단위 "
  "측정 이동량만이 유효한 dose다. 양방향 FE는 센서·날짜 평균을 교대로 차감하는 반복 within-transformation으로 추정했다.")
P("이를 바탕으로 네 가지 보조 분석을 수행했다. (M3) 기능분할: M2를 결과변수(주간/야간/주야 gap), 토지이용(상업/혼합/주거; 동의 주간:"
  "야간 인구비 3분위), 평일/주말, 계절별로 반복한다. (M4) 차이의 차이(difference-in-differences) event study: 이동량 고영향(상업) 동과 "
  "저영향(주거) 동의 ΔL_day 차이를 주별로 추적한다 — 공통 교란이 차분에서 상쇄되어 drift-robust하다. (M5) 검증: 근접 S-DoT와 공식 "
  "도로교통 4점 LAeq 상시측정소의 연추세를 대조해 드리프트·스키마 경계를 점검한다. 공간 단면 회귀는 소수 극단값에 끌리지 않도록 "
  "Theil-Sen 회귀와 Spearman 상관(아웃라이어 내성)을 쓰고, 동별 추정 신뢰도를 위해 센서 2개 이상인 동만 추세 분석에 사용했다. 모든 "
  "분석은 Python(statsmodels)으로 수행했다.")
figure("fig_identification_concept.png", 2,
       "Identification strategy. (a) Causal chain (graded distancing -> measured mobility -> urban noise) and confounders "
       "(weather and calendar; sensor offset and multi-year drift). (b) Fixed-effects ladder: sensor FE removes per-sensor "
       "calibration offset; date FE removes all date-common factors (weather, holidays, city trend, drift); the residual "
       "cross-neighbourhood mobility variation within each date identifies the dose-response.", 5.2)

# ============ 3. Results ============
H("3. Results", 1)
dp = pd.read_csv(os.path.join(PD, "analysis_panel.csv"),
                 usecols=["Leq_day", "Leq_night", "Leq24", "lp_day_logrel", "temp_mean", "precip_mm", "wind_max_ms"])
DVARS = [("L_day (dB)", "Leq_day"), ("L_night (dB)", "Leq_night"), ("Leq24 (dB)", "Leq24"),
         ("Daytime mobility (log rel.)", "lp_day_logrel"), ("Mean temperature (C)", "temp_mean"),
         ("Precipitation (mm)", "precip_mm"), ("Max wind (m/s)", "wind_max_ms")]
drows = [[lab, f"{dp[c].mean():.2f}", f"{dp[c].std():.2f}", f"{dp[c].quantile(.05):.2f}",
          f"{dp[c].quantile(.95):.2f}", f"{dp[c].notna().sum():,}"] for lab, c in DVARS]
mtable(3, "Descriptive statistics of the analysis panel (1,248,794 sensor-days; 1,123 sensors; 421 dong; 2020-2023).",
       ["Variable", "Mean", "SD", "P5", "P95", "N"], drows, [2.2, 0.8, 0.8, 0.8, 0.8, 1.0])

H("3.1. A robust but modest daytime dose-response", 2)
P("센서 고정효과(M1, Table 5)에서 주간 동 이동량과 주간 소음은 정(+)의 용량-반응을 보였다(β=+1.130 dB/log-unit, SE 0.209, p<0.001). "
  "통제변수는 모두 물리적으로 타당했다 — 기온은 U자형(저온·고온에서 소음 증가), 강수 +0.045 dB/mm, 강수일 +0.10 dB, 주말 −0.83 dB, "
  "공휴일 −1.18 dB(모두 p<0.001). 그러나 날짜 공통변이를 제거한 양방향 고정효과(M2, 헤드라인)에서 효과는 작아졌다: 주간 β=+0.648 "
  "dB/log-unit(95% CI 0.17-1.13, p=0.008), 전일 β=+0.628(p=0.038). 해석하면 동 주간 이동량 30% 감소가 주간소음 약 −0.23 dB, 50% "
  "감소가 약 −0.45 dB에 대응한다. 즉 도시 공통 시간변이를 엄밀히 제거한 순수 cross-dong 용량-반응은 통계적으로 견고하나 작다(RQ1).")
m1 = pd.read_csv(os.path.join(PD, "phase2a_doseresponse_results.csv"))
def cell(model, prefix):
    sub = m1[m1.model == model]; row = sub[sub.term.str.startswith(prefix)]
    if len(row) == 0:
        return "-"
    r = row.iloc[0]; return f"{r.coef:+.3f}{star(r.pval)} ({r.se:.3f})"
TERMS = [("Mobility (log living pop.)", "lp_"), ("Mean temperature", "temp_mean"), ("Temperature^2", "temp_sq"),
         ("Precipitation (mm)", "precip_mm"), ("Max wind (m/s)", "wind_max_ms"), ("Rain (0/1)", "rain_flag"),
         ("Weekend (0/1)", "is_weekend"), ("Holiday (0/1)", "holiday")]
m1rows = [[lab, cell("Leq_day~lp_day", pf), cell("Leq_night~lp", pf), cell("Leq24~lp_mean", pf)] for lab, pf in TERMS]
m1rows.append(["Sensors / N", "1,122 / 1.25M", "1,122 / 1.25M", "1,122 / 1.25M"])
mtable(5, "Sensor fixed-effects regression (M1): coefficient (SE) with sensor-clustered standard errors. "
          "*p<0.05, **p<0.01, ***p<0.001.", ["Term", "L_day", "L_night", "Leq24"], m1rows, [2.3, 1.4, 1.4, 1.4])
P("이 효과의 기능적 위치를 분해하면(Fig. 3, Table 6, RQ2) 효과는 daytime에 특정된다. 결과변수별로 주간만 유의하고 야간·주야 gap은 "
  "비유의했다. 토지이용별로는 상업(+0.34)·혼합(+0.86)·주거(+0.85) 모두 양(+)이었으나 표본 분할로 검정력이 떨어져 개별 유의성은 약했고, "
  "상호작용 검정에서도 토지이용 차이는 유의하지 않았다(효과는 대체로 균일). 평일(+0.49)·주말(+0.64)은 유사했다. 특히 계절별로는 네 "
  "계절 모두 양(+)이며 겨울(DJF +0.66)·봄(MAM +0.69)·가을(SON +0.76)에서 유의해, 효과가 특정 계절의 아티팩트가 아니라 연중 일관되게 "
  "존재함을 보였다.")
seg = pd.read_csv(os.path.join(PD, "phase3_segmented_results.csv"))
segrows = [[r.panel.split(") ")[-1] if ")" in r.panel else r.panel, r.group, f"{r.beta:+.3f}{star(r.pval)}",
            f"[{r.beta-1.96*r.se:+.2f}, {r.beta+1.96*r.se:+.2f}]", f"{int(r.n):,}"] for _, r in seg.iterrows()]
mtable(6, "Function-segmented dose-response (M2, two-way fixed effects): beta (dB per log-unit mobility), 95% CI and N. "
          "*p<0.05, **p<0.01.", ["Segment", "Group", "beta", "95% CI", "N"], segrows, [1.6, 1.7, 0.9, 1.5, 1.0])
figure("fig_segmented_forest.png", 3,
       "Function-segmented dose-response (two-way sensor+date fixed effects; markers = beta, bars = 95% CI). (a) by "
       "outcome (daytime, nighttime, day-night gap); (b) by land-use; (c) by day type; (d) by season. Green = p<0.05.", 5.2)

H("3.2. The effect is a daytime phenomenon", 2)
P("주간/야간 비교는 효과가 주간 활동에 기인함을 분명히 한다(Fig. 4). 야간 소음은 센서 고정효과 모형에서 야간 이동량에 강하게 반응하는 "
  "듯 보였으나(β=+2.28, p<0.001), 양방향 고정효과에서 그 연관은 사라졌다(β=+0.50, 95% CI −0.73-+1.74, p=0.42). 이는 야간의 겉보기 "
  "효과가 실제 인과가 아니라 시간 공통 교란(드리프트·계절·도시추세)이었음을 드러낸다. 반면 주간 효과는 양방향 고정효과에서도 "
  "살아남는다. 주야 gap(L_day − L_night; 같은 센서·같은 날 차분이라 오프셋과 공통 드리프트가 상쇄)의 용량-반응은 ≈0(β=+0.02, "
  "p=0.92)으로, 효과가 주야 차등 메커니즘이 아니라 주간 절대수준의 modest 감소임을 시사한다.")
figure("fig_daynight.png", 4,
       "Day vs night. (a) two-way FE dose-response for daytime, nighttime, and the day-night gap; (b) the nighttime "
       "association apparent under sensor FE (+2.28) vanishes under sensor+date FE (+0.50, ns), revealing it as a "
       "time-confound, while the daytime effect survives.", 5.6)
H("3.3. Spatiotemporal mobility and functional differentiation", 2)
P("이동량 dose의 시공간 구조는 본 설계의 핵심이다(Fig. 5). 도시 전체 생활인구 총량은 거의 보존되지만, 거리두기 국면마다 상업·업무 동은 "
  "주간 체류가 최대 약 35% 비워지고 주거 동은 오히려 증가하는 공간 재분포가 나타난다. 가장 강한 제한기(2020-12 2.5단계+5인금지, "
  "2021-07 수도권 4단계)에 도심 상업동이 가장 비워졌고, 위드코로나(2021-11)와 전면해제(2022-04)를 거치며 회복했다.")
figure("fig_phase_mobility_maps.png", 5,
       "Spatiotemporal evolution of daytime mobility (neighbourhood-level living population relative to the post-lift "
       "baseline) across four graded distancing phases. The city-wide median stays near 1.0, but commercial/business "
       "neighbourhoods (teal) empty by up to ~35% under the strongest restrictions.", 5.8)
P("이 기능적 차등은 시간축에서도 뚜렷하다(Fig. 6, RQ3). 상업 동의 주간 이동량은 제한기 내내 기준선 아래로, 주거 동은 기준선 위로 "
  "벌어졌다가 2022년 해제 후 수렴한다(수도권 4단계기 중앙값 상업 0.98 vs 주거 1.05). 드리프트를 제거한 상대 소음(동 그룹평균 − "
  "도시평균)에서도 상업 동이 제한기에 도시평균보다 상대적으로 조용한 경향이 관찰된다. 거리두기 전환의 동역학은 차이의 차이로 본다"
  "(Fig. 7) — 고영향 동과 저영향 동의 주별 ΔLAeq 차이는 제한기에 음(−)으로 기울고 정상기에 0으로 수렴했으며, 주별 (이동량 차이 vs "
  "소음 차이) 상관은 Pearson +0.44 · Spearman +0.46으로 일치했다.")
figure("fig_landuse_trajectory.png", 6,
       "Functional differentiation over time. (a) Weekly daytime mobility by land-use type - commercial neighbourhoods "
       "fall below and residential rise above baseline during restrictions, converging after the lift. (b) Daytime noise "
       "relative to the city mean (drift-robust, 4-week moving average) by land-use type.", 5.6)
figure("fig_did_eventstudy.png", 7,
       "Difference-in-differences event study. Weekly difference between high- and low-mobility-loss neighbourhoods in "
       "(top) within-sensor delta-L_day and (bottom) daytime mobility. Common drift, season and city trend cancel; a "
       "negative difference during restrictions indicates high-impact (commercial) neighbourhoods were relatively quieter.", 6.0)
H("3.4. No robust long-run spatial gradient", 2)
P("반면 동 단위 장기 공간 단면에서는 robust한 이동량-소음 gradient가 나타나지 않았다(Fig. 8-9). 동별 (이동량 감소율 vs 도시평균 대비 "
  "상대 소음변화)의 아웃라이어-내성 상관은 Spearman ρ≈−0.05(전체)로 사실상 0이며, 상업 동에서 Pearson r=+0.38처럼 보이는 연관은 소수 "
  "단일센서 동의 극단값에 끌린 아티팩트로(같은 부분집합의 Spearman ρ=+0.08, Theil-Sen 기울기≈0), 신뢰도 필터(센서≥2)와 robust 통계를 "
  "적용하면 사라진다. 즉 이동량-소음 신호는 고빈도 cross-dong/within-date 변이(M2·DiD)에만 robust하게 존재하며, naive한 동 단위 공간 "
  "cross-section으로는 검출되지 않는다.")
figure("fig_change_map.png", 8,
       "Spatial pattern during the distancing era. (a) land-use type; (b) daytime mobility reduction; (c) drift-removed "
       "relative noise change; (d) mobility vs noise by land-use (Theil-Sen fit, Spearman rho; neighbourhoods with >=2 "
       "sensors, point size proportional to sensor count). The long-run spatial cross-section shows no robust gradient.", 6.0)
figure("fig_spatial_bylanduse.png", 9,
       "Commercial vs residential neighbourhoods. (a,c) drift-removed relative noise change; (b,d) mobility vs noise "
       "(Theil-Sen, Spearman vs Pearson). The Pearson correlation is inflated by a few single-sensor outliers; the robust "
       "Spearman is approximately zero.", 6.0)
H("3.5. Sensor drift and official-station validation", 2)
P("원시 within-sensor ΔLAeq(post-lift 기준)는 역설을 보였다 — 최강제한기 ΔL_day가 정상기보다 높았다(+0.85 vs +0.34 dB). 진단 결과 "
  "이는 코로나 효과가 아니라 다년 절대레벨 교란이었다(Fig. 10). 4년 전수 가동 842센서의 연평균 Leq24는 49.4(2020)-47.6-47.2-47.1 "
  "dB(2023)로 평균·중앙값 모두 단조감소했다(85% 센서가 하락추세). 2022-2023 스키마 경계는 깨끗했다(+0.02 dB). 그러나 이 감소는 균일한 "
  "하드웨어 드리프트가 아니다: 검증된 도로 4점에서 근접 S-DoT의 2021-2023 추세는 City Hall +1.7·Sinsa +2.5·Sinchon ≈0·Seongsu "
  "−0.3 dB로 하락하지 않고 공식 LAeq 추세(각 +7.2[공사 추정]·+0.4·−0.4·+1.5)와 같은 방향이었다. 즉 도로변 S-DoT는 공식 측정망을 잘 "
  "추종해 신뢰할 만하나, 도시 전체 평균의 하락(−0.5 dB/2021-23)은 비도로·정온지역에 집중된 위치별 현상으로 다년 절대비교를 신뢰할 수 "
  "없게 한다. 이는 M2의 cross-dong within-date 식별이 유일하게 유효한 이유를 뒷받침한다.")
figure("fig_drift_validation.png", 10,
       "Drift diagnostic and validation. (a) citywide multi-year noise level (mean and median) of 842 sensors active in "
       "all four years declines monotonically. (b) per-station 2021->2023 trends: nearby S-DoT vs official roadside LAeq; "
       "road-site S-DoT does not fall, so the citywide decline is not uniform hardware drift.", 6.0)

# ============ 4. Discussion ============
H("4. Discussion", 1)
H("4.1. Principal findings", 2)
P("본 연구는 측정 이동량에 대한 도시소음의 graded 용량-반응을 처음으로 정량 추정했다. 핵심 결과는 네 가지로 요약된다. 첫째, 주간 동 "
  "이동량과 주간 소음 사이에 통계적으로 견고한 정(+)의 용량-반응이 존재하나 그 크기는 작다(이동량 30% 감소 ≈ −0.2~0.3 dB). 둘째, "
  "효과는 daytime에 특정되며 네 계절에 걸쳐 일관된다 — 야간의 겉보기 효과는 엄밀 식별에서 시간교란으로 소멸한다. 셋째, 효과는 고빈도 "
  "cross-dong 변이와 차이의 차이 분석에서만 robust하게 드러나며, 동 단위 장기 공간 단면에서는 사라진다. 넷째, 저가 IoT 망의 다년 "
  "절대레벨은 드리프트·위치별 회복으로 교란되어 시계열·절대 비교를 신뢰할 수 없고, 오직 within-sensor·cross-dong 설계만이 유효하다.")
H("4.2. The magnitude in context: a modest, sticky response", 2)
P("우리의 추정치는 봉쇄기 도시소음의 수 dB 감소를 보고한 기존 문헌 [18,19,20,21,22,23]과 정면으로 대비된다. 그 보고들은 대개 특정 "
  "번화가의 절대레벨을 팬데믹 전후로 단순 비교한 것으로, 가장 시끄럽고 가장 많이 비워진 소수 지점에서 측정된 극단적 변화를 도시 전체로 "
  "일반화하기 쉽다. 반면 본 연구는 약 1,100점 전역에서 센서 내·이동량 graded·전기간 회귀로 한계효과(marginal effect)를 추정했기에, "
  "도시 평균적 반응은 그보다 훨씬 작게 나타난다. 더욱이 우리 추정치는 하한(lower bound)이다 — 양방향 고정효과는 cross-dong 성분만 "
  "식별하며, 도시 전체에 공통된 이동량 감소 성분(모두가 덜 움직인 효과)은 센서 드리프트와 교란되어 깨끗이 추정할 수 없기 때문이다. "
  "효과가 작은 것은 도시소음이 통과교통·간선도로 등 비탄력적 배경원에 의해 '끈적(sticky)'하다는 점, 그리고 S-DoT의 1 dB 정수 "
  "해상도가 측정오차로서 표준오차를 넓힌다는 점과도 정합한다. 요컨대 '봉쇄가 도시를 수 dB 조용하게 만들었다'는 통념은 엄밀한 "
  "이동량-소음 관점에서는 과장이며, 진짜 graded 용량-반응은 견고하되 작다.")
H("4.3. Why daytime, why uniform, why no spatial map", 2)
P("효과의 패턴은 메커니즘과 정합한다. 효과가 daytime·연중 일관인 것은, 주간 소음이 상거래·통근·방문 등 사람 활동에 직접 연동되는 반면 "
  "야간 소음은 고정 설비·간선교통 등 배경원이 지배해 이동량 변동에 둔감하기 때문이다. 야간의 겉보기 효과가 엄밀 식별에서 소멸한 사례는, "
  "시간 추세를 인과로 오인하는 naive 모형의 위험을 잘 보여준다. 효과가 토지이용에 비교적 균일했다는 점은, 이동량-소음 반응이 특정 장소 "
  "유형에 국한되지 않고 도시 전반에 얇게 퍼져 있음을 시사한다. 특히 주목할 것은 동 단위 장기 공간 단면에서 깨끗한 gradient가 나타나지 "
  "않은 점이다 — 이는 효과가 없어서가 아니라, 동별 장기평균 소음변화가 국지적 공사·도로 변화·센서별 드리프트 같은 특이요인에 압도되기 "
  "때문이다. 신호는 같은 날짜 안의 고빈도 동 간 변이에 있으며, 단순한 공간 지도화로는 드러나지 않는다. 실제로 상업 동에서 관측된 양의 "
  "Pearson 상관조차 소수 단일센서 동의 극단값이 만든 아티팩트였고, robust 통계에서는 0이었다 — 고밀도 IoT 자료의 공간 분석이 얼마나 "
  "쉽게 오도될 수 있는지를 보여주는 경고다.")
H("4.4. Methodological and monitoring implications", 2)
P("본 연구의 방법론적 기여는 결과 자체만큼 중요하다. 저가 고밀도 IoT 소음망으로 도시 환경효과를 측정하려는 시도는 세 가지 함정을 "
  "정면으로 드러낸다 — (i) 센서별 검교정 오프셋은 within-sensor 설계로 상쇄하고, (ii) 다년 드리프트·배치시점·위치별 회복은 절대/"
  "시계열 비교를 불가능하게 하므로 cross-dong 식별이 필수이며, (iii) 도시 단일 dose의 시간교란은 동 단위 측정 dose를 요구한다. 특히 "
  "장기 공간 단면이 robust하게 null이라는 결과는, 검교정 없이 절대레벨을 그대로 지도화하는 접근이 거짓 패턴 또는 거짓 무패턴을 낳을 수 "
  "있음을 경고한다. 이는 SONYC [37]·WASN [38]·NoiseCapture [50] 등 확산 중인 저가 소음 모니터링 [36], 나아가 스마트시티 센서 데이터 "
  "[40,41]의 정책 활용 전반에 적용된다. 본 연구가 제시한 고정효과·아웃라이어 내성·공식망 교차검증 [42]의 절차는 후속 IoT 기반 도시 "
  "환경연구의 식별 설계 지침이 될 수 있다.")
H("4.5. Planning implications, generalizability, and limitations", 2)
P("계획·정책 측면에서, 이동량·교통수요 관리가 소음에 주는 직접적 공편익(co-benefit)은 통념보다 작다 — 봉쇄급 이동량 급감조차 동 "
  "단위에서 1 dB 안팎의 변화를 낳았다. 따라서 '15분 도시'·차 없는 거리 등 이동량 저감 정책은 그 자체로 정당하나, 소음 저감 수단으로는 "
  "노면 포장·저소음차량·차폐 등 음원·전파 단계 개입 [43,45]과 병행되어야 한다. 효과가 토지이용에 비교적 균일했다는 점은 상업/주거 "
  "구분에 따른 차등 규제보다 광역적 접근의 여지를 시사한다. 서울은 세계 최고 밀도의 메가시티 중 하나로, 본 결과는 도쿄·홍콩·싱가포르 "
  "등 고밀도 동아시아 도시에 일반화 가능성이 높다.")
P("본 연구는 몇 가지 한계를 가진다. (1) S-DoT는 비검교정·1 dB 해상도·노후 드리프트로 절대/시간축 신뢰도가 낮아 상대변화·cross-dong로 "
  "우회했다. (2) 생활인구는 체류(presence)로 소음생성 활동(특히 교통)의 불완전한 proxy이며, 향후 버스·지하철 정류장 단위 승하차를 동 "
  "단위로 매핑한 더 직접적인 교통 dose [31,33]가 효과를 선명히 할 수 있다. (3) 거리두기 준연속 코딩의 정밀 일자경계는 질병청 HWPX "
  "원본 대조가 필요하다. (4) 기상은 ERA5 재분석으로 기상청 공식 ASOS로 대체 가능하다. (5) 도시 전체 공통 이동량 효과는 드리프트 "
  "교란으로 미식별되어 추정치는 하한이며, 검교정 참조망과의 결합으로 절대레벨을 복원하는 후속연구가 필요하다. (6) 토지이용 분류는 "
  "주간/야간 인구비 proxy로 형태학적 분류로 정밀화할 수 있다. (7) 분석은 단일 도시에 한정되며, 소음의 주파수·시간 미세구조와 "
  "사운드스케이프 지각 [11,16,51]으로의 확장이 향후 과제다.")

# ============ 5. Conclusions ============
H("5. Conclusions", 1)
P("우리는 한국의 단계적 거리두기를 graded 자연실험으로 삼아, 도시 규모 IoT 센서망에서 측정 이동량에 대한 도시소음의 용량-반응을 "
  "추정했다. 효과는 통계적으로 견고하나 작고(주간 약 0.2-0.4 dB / 이동량 30% 감소), daytime·연중 일관되며, 엄밀한 cross-dong 식별에서만 "
  "드러난다. 동 단위 장기 공간 단면은 robust하게 null이며, 저가 IoT 망의 다년 절대레벨은 드리프트로 교란된다. 이는 과장된 봉쇄 문헌에 "
  "대한 정직한 반례이자, IoT 기반 도시 환경연구의 식별 설계와 스마트시티 모니터링에 대한 실천적 지침이다.")

H("Data Availability Statement", 1)
P("모든 원시데이터는 공개 공공데이터이다(S-DoT 소음 OA-15969, 생활인구 OA-14991, 지하철 OA-12921, 거리두기 시행연혁 공공데이터포털 "
  "15106451, 기상 Open-Meteo ERA5, 행정동 경계 vuski/admdongkor, 검증 4점 LAeq OA-15473). 지하철 데이터는 공공누리 제3유형"
  "(출처표시+변경금지). 분석 코드는 [저장소 TBD]에 공개 예정이다.", indent=False)

H("References", 1)
for r in sorted(REFS, key=lambda x: x["num"]):
    p = doc.add_paragraph()
    setfont(p.add_run(f"[{r['num']}] "), 10); setfont(p.add_run(r["vancouver"]), 10)
    pf = p.paragraph_format; pf.line_spacing = 1.2; pf.space_after = Pt(2)
    pf.left_indent = Inches(0.3); pf.first_line_indent = Inches(-0.3)

doc.save(OUT)
print("saved:", OUT)
print("paragraphs:", len(doc.paragraphs), "tables:", len(doc.tables), "images:", sum(1 for _ in doc.inline_shapes), "refs:", len(REFS))

# -*- coding: utf-8 -*-
# v8 정제9: 서론·논의 소제목 간결화(AI·장황 표현 제거). 볼드 12pt TNR 유지, 본문/구조 불변.
import os
from datetime import datetime, timezone, timedelta
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = r"T:\00_BACKUP\01_한양대학교\00_연구실\00_프로젝트\★_논문화 프로젝트\31_환경소음_이동성"
SRC = os.path.join(ROOT, "01_논문작업", "Manuscript_v8_20260623_114448.docx")
ts = datetime.now(timezone(timedelta(hours=9))).strftime("%Y%m%d_%H%M%S")
OUT = os.path.join(ROOT, "01_논문작업", f"Manuscript_v8_{ts}.docx")

doc = Document(SRC)
def find(m):
    h = [p for p in doc.paragraphs if p.text.strip().startswith(m)]
    assert len(h) == 1, f"마커 {len(h)}건: {m[:30]}"
    return h[0]

def set_heading(p, text):
    # 기존 헤딩 서식(TNR 12pt bold, eastAsia 없음)에 맞춰 한 run으로 재작성
    for r in list(p.runs): r._element.getparent().remove(r._element)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'; run.font.size = Pt(12); run.bold = True
    rpr = run._element.get_or_add_rPr(); rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'): rf.set(qn(a), 'Times New Roman')
    # eastAsia 미설정(영문 헤딩·기존 None 규범에 맞춤)

H = [
    ("1.1. Environmental noise as a sustainability", "1.1. Environmental noise and urban sustainability"),
    ("1.2. Mobility as the dominant", "1.2. Mobility as a driver of urban noise"),
    ("1.3. COVID-19 as a natural experiment", "1.3. COVID-19 as a natural experiment"),
    ("1.4. Measured mobility", "1.4. Measured mobility and IoT noise sensing"),
    ("4.2. The magnitude in context", "4.2. The magnitude in context"),
    ("4.3. Why daytime", "4.3. Mechanisms behind the patterns"),
    ("4.4. Methodological and monitoring", "4.4. Methodological implications"),
    ("4.5. Planning implications", "4.5. Policy implications and limitations"),
]
for marker, newt in H:
    set_heading(find(marker), newt)
print(f"소제목 간결화: {len(H)}건")
doc.save(OUT)
print("저장:", OUT)

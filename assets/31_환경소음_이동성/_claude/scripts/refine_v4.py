# -*- coding: utf-8 -*-
# v8 정제4: ①Table 1·2 줄단위 정리(사용자 형식 유지·간결 병렬) ②수정 Fig 3·4·8·10 재임베드(사용자 Fig1·2 보존).
import os, re, struct
from datetime import datetime, timezone, timedelta
from docx import Document
from docx.shared import Pt, Emu
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = r"T:\00_BACKUP\01_한양대학교\00_연구실\00_프로젝트\★_논문화 프로젝트\31_환경소음_이동성"
FIGD = os.path.join(ROOT, "_claude", "figures")
SRC = os.path.join(ROOT, "01_논문작업", "Manuscript_v8_20260622_003823.docx")
ts = datetime.now(timezone(timedelta(hours=9))).strftime("%Y%m%d_%H%M%S")
OUT = os.path.join(ROOT, "01_논문작업", f"Manuscript_v8_{ts}.docx")

def setfont(run, size=10, bold=False, sub=False):
    run.font.size = Pt(size); run.bold = bold; run.font.name = 'Times New Roman'
    if sub: run.font.subscript = True
    rpr = run._element.get_or_add_rPr(); rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'): rf.set(qn(a), 'Times New Roman')
    rf.set(qn('w:eastAsia'), '바탕')

SUB = [('Leq,24h', ('L', 'eq,24h')), ('LAeq', ('L', 'Aeq')), ('Lday', ('L', 'day')),
       ('Lnight', ('L', 'night')), ('Leq', ('L', 'eq'))]
SRE = re.compile('(' + '|'.join(re.escape(k) for k, _ in SUB) + ')'); SMAP = dict(SUB)

def emit(p, text, size=10, bold=False):
    for part in SRE.split(text):
        if not part: continue
        if part in SMAP:
            b, sb = SMAP[part]; setfont(p.add_run(b), size, bold); setfont(p.add_run(sb), size, bold, sub=True)
        else:
            setfont(p.add_run(part), size, bold)

def cell_lines(cell, lines, size=10, bold=False):
    for extra in cell.paragraphs[1:]: extra._element.getparent().remove(extra._element)
    p0 = cell.paragraphs[0]
    for r in list(p0.runs): r._element.getparent().remove(r._element)
    emit(p0, lines[0], size, bold)
    p0.paragraph_format.space_before = Pt(0); p0.paragraph_format.space_after = Pt(2)
    for ln in lines[1:]:
        p = cell.add_paragraph(); emit(p, ln, size, bold)
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)

# ---------- Table 1 ----------
T1 = [(["Item"], ["Description"]),
      (["Dataset"], ["Smart Seoul Data of Things (S-DoT)",
                     "City-wide permanent urban-sensing network",
                     "Seoul Open Data Plaza, dataset OA-15969 (no authentication required)"]),
      (["Coverage"], ["~1,100 fixed sensors across all 25 districts of Seoul",
                      "Hourly records, April 2020 – December 2023"]),
      (["Measurement"], ["Broadband sound pressure level (dB) at each sensor",
                         "Aggregated to energy-equivalent daily levels:",
                         "Leq,24h (whole day), Lday (06:00–21:00), Lnight (22:00–05:00)"]),
      (["Quality control"], ["Sensor-days with <12 valid hours or values outside 20–95 dB discarded",
                             "2023 field change (single 'mean noise' value) reconciled to one definition"]),
      (["Use"], ["Not calibrated to a shared absolute reference",
                 "Used only for within-sensor change over time",
                 "Cross-validated against the official four-point roadside LAeq network (§3.5)"])]
# ---------- Table 2 ----------
T2 = [(["Item"], ["Description"]),
      (["Mobility exposure"], ["Daytime de-facto ('living') population per administrative dong",
                               "People actually present in a neighbourhood (mobile-network signalling, not registered residence)",
                               "Seoul Open Data Plaza, dataset OA-14991"]),
      (["Coverage"], ["Hourly counts per dong, averaged to daily daytime (06:00–21:00) and night-time (22:00–05:00) values",
                      "424 dong × 1,461 days = 619,464 dong-days",
                      "~2.8 GB of raw hourly files processed, then deleted to save storage"]),
      (["Dose variable"], ["Within-dong log relative change",
                           "lp = log(daily population / dong's post-lifting baseline mean)",
                           "Each dong compared against its own usual level"]),
      (["Weather"], ["Daily mean / max / min temperature, precipitation, max wind speed",
                     "Open-Meteo ERA5 reanalysis at Seoul city centre (37.57°N, 126.97°E)"]),
      (["Calendar"], ["Day of week, weekend indicator",
                      "Korean public holidays (incl. substitutes), season"])]

doc = Document(SRC)

def find_table(sig0, sig1sub):
    for tb in doc.tables:
        if tb.rows[0].cells[0].text.strip() == "Item" and sig1sub in tb.rows[1].cells[1].text:
            return tb
    return None
t1 = find_table("Item", "Smart Seoul")
t2 = find_table("Item", "Daytime de-facto")
assert t1 and t2, "Table 1/2 식별 실패"
for tb, data in ((t1, T1), (t2, T2)):
    assert len(data) == len(tb.rows), f"행수 {len(data)} vs {len(tb.rows)}"
    for ri, (icell, dcell) in enumerate(data):
        cell_lines(tb.rows[ri].cells[0], icell, 10, bold=True)            # Item 열 = 볼드
        cell_lines(tb.rows[ri].cells[1], dcell, 10, bold=(ri == 0))       # 헤더만 볼드
print("Table 1·2 줄단위 정리 완료")

# ---------- 수정 Fig 3·4·8·10 재임베드 (사용자 Fig1·2=idx0,1 보존) ----------
EMBED = {2: "fig_segmented_forest.png", 3: "fig_daynight.png", 7: "fig_change_map.png", 9: "fig_drift_validation.png"}
for idx, fn in EMBED.items():
    blob = open(os.path.join(FIGD, fn), "rb").read()
    wpx, hpx = struct.unpack(">II", blob[16:24])
    sh = doc.inline_shapes[idx]
    rId = sh._inline.graphic.graphicData.pic.blipFill.blip.embed
    doc.part.related_parts[rId]._blob = blob
    sh.height = Emu(int(round(int(sh.width) * hpx / wpx)))   # 폭 유지(사용자 리사이즈 존중), 높이 재계산
    print(f"  Fig idx{idx} ({fn}): {sh.width.inches:.2f} x {sh.height.inches:.2f}")

doc.save(OUT)
print("\n저장:", OUT)

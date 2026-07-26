# -*- coding: utf-8 -*-
# 영역 적용: 한국어 정본(서식 통일본)의 사본을 만들어 문단 텍스트만 영문 교체.
# 서식 보존: 문단 pPr 유지 + 기존 첫 run의 폰트 크기 상속. 첨자 엔진 재사용.
# 마지막에 전 문단·표 셀의 em/en dash 제거(" – "→", ", 나머지 –—→-).
import os, sys, shutil
from datetime import datetime, timezone, timedelta
from docx import Document
from docx.shared import Pt
sys.path.insert(0, os.path.dirname(__file__))
import docxtools_v9 as T
from translate_en_map import PARA

SRC = os.path.join(T.ROOT, "01_논문작업", "Manuscript_20260726_231342.docx")
ts = datetime.now(timezone(timedelta(hours=9))).strftime("%Y%m%d_%H%M%S")
OUT = os.path.join(T.ROOT, "01_논문작업", f"Manuscript_EN_{ts}.docx")
shutil.copy2(SRC, OUT)

doc = Document(OUT)

def para_size(p, default=11.0):
    for r in p.runs:
        if r.font.size is not None:
            return r.font.size.pt
    st = p.style
    while st is not None:
        if st.font.size is not None:
            return st.font.size.pt
        st = st.base_style
    return default

n = 0
for marker, en in PARA:
    hits = [p for p in doc.paragraphs if p.text.strip().startswith(marker)]
    assert len(hits) == 1, f"마커 {len(hits)}건: {marker}"
    p = hits[0]
    sz = para_size(p)
    T.fill_para(p, en, size=sz)
    n += 1
print(f"영역 적용 {n}문단")

# ---- Table 1 사실 교정 (Bochum [54] = 23개 녹음지점 + 계수지점 1, PMC10048852 실측) ----
TABLE_FIX = [("Bochum (1 long-term site)", "Bochum (23 sound monitors, 1 traffic-count station)")]
nt = 0
for tbl in doc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    for old, new in TABLE_FIX:
                        if old in r.text:
                            r.text = r.text.replace(old, new); nt += 1
print(f"표 셀 교정 {nt}건")

# ---- em/en dash 정리 (전 문단 + 표 셀) ----
def clean_dashes(runs):
    c = 0
    for r in runs:
        t = r.text
        if "–" in t or "—" in t:
            t = t.replace(" – ", ", ").replace(" — ", ", ").replace("–", "-").replace("—", "-")
            r.text = t; c += 1
    return c

nd = sum(clean_dashes(p.runs) for p in doc.paragraphs)
for tbl in doc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            nd += sum(clean_dashes(p.runs) for p in cell.paragraphs)
print(f"dash 정리 run {nd}건")

doc.save(OUT)
print("저장:", OUT)

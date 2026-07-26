# -*- coding: utf-8 -*-
# v8 정제10: 본문 내 섹션 자기참조 (§3.5) 등 제거(AI 표현). run 텍스트만 수술 — 서식·첨자 불변.
import os, re
from datetime import datetime, timezone, timedelta
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = r"T:\00_BACKUP\01_한양대학교\00_연구실\00_프로젝트\★_논문화 프로젝트\31_환경소음_이동성"
SRC = os.path.join(ROOT, "01_논문작업", "Manuscript_v8_20260630_173352.docx")
ts = datetime.now(timezone(timedelta(hours=9))).strftime("%Y%m%d_%H%M%S")
OUT = os.path.join(ROOT, "01_논문작업", f"Manuscript_v8_{ts}.docx")

doc = Document(SRC)
PAT = re.compile(r'\s?\(§[\d.]+\)')   # "(§3.5)" 형태(앞 공백 옵션)
n = 0

def set_body_run(p, text):   # 본문 1 run으로 재구성(TNR 12pt, eastAsia 미설정=None 규범)
    for r in list(p.runs): r._element.getparent().remove(r._element)
    run = p.add_run(text); run.font.name = 'Times New Roman'; run.font.size = Pt(12)
    rpr = run._element.get_or_add_rPr(); rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'): rf.set(qn(a), 'Times New Roman')

# ① [33] 커스텀: run이 잘게 쪼개져 있어 문단을 한 run으로 재구성하며 재서술(첨자 없음)
for p in doc.paragraphs:
    if p.text.strip().startswith("그러나 저가 IoT 소음망에는 기존 COVID-소음"):
        new = p.text.replace("거짓 결론에 이른다(본 연구의 §3.5가 그 실례를 보인다)",
                             "거짓 결론에 이르며, 본 연구가 그 실례를 직접 보인다")
        set_body_run(p, new); n += 1; break

# ② 일반 "(§X)" 제거 — 문단
for p in doc.paragraphs:
    for run in p.runs:
        if "§" in run.text:
            new = PAT.sub('', run.text)
            if new != run.text:
                run.text = new; n += 1

# ③ 표 셀 "(§X)" 제거 (Table 1 Use 등)
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            for p in c.paragraphs:
                for run in p.runs:
                    if "§" in run.text:
                        new = PAT.sub('', run.text)
                        if new != run.text:
                            run.text = new; n += 1

print("§ref 수정:", n)
left = (sum(p.text.count("§") for p in doc.paragraphs)
        + sum(c.text.count("§") for t in doc.tables for r in t.rows for c in r.cells))
print("잔존 § (0이어야):", left)
doc.save(OUT)
print("저장:", OUT)

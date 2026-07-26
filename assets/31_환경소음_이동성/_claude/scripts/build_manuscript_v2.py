# -*- coding: utf-8 -*-
# Manuscript v2 — Phase 1-2 결과·드리프트 진단·provenance 통합. python-docx.
# 정직한 1+3 프레이밍: robust하나 modest한 graded dose-response + 방법론(드리프트·해상도→cross-dong만 유효).
import os
from datetime import datetime, timezone, timedelta
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ts = datetime.now(timezone(timedelta(hours=9))).strftime("%Y%m%d_%H%M%S")
ROOT = r"T:\00_BACKUP\01_한양대학교\00_연구실\00_프로젝트\★_논문화 프로젝트\31_환경소음_이동성"
OUT = os.path.join(ROOT, "01_논문작업", f"Manuscript_v2_{ts}.docx")

doc = Document()
n = doc.styles['Normal']; n.font.name = 'Arial'; n.font.size = Pt(10.5)

def H(t, lvl=1): return doc.add_heading(t, level=lvl)
def P(t, italic=False, size=10.5, bold=False):
    p = doc.add_paragraph(); r = p.add_run(t); r.italic = italic; r.bold = bold; r.font.size = Pt(size); return p
def prov(rows, w0=1.7, w1=4.6):
    tbl = doc.add_table(rows=len(rows), cols=2); tbl.style = 'Table Grid'
    for i, (k, v) in enumerate(rows):
        c = tbl.rows[i].cells
        rk = c[0].paragraphs[0].add_run(k); rk.bold = True; rk.font.size = Pt(9.5)
        rv = c[1].paragraphs[0].add_run(v); rv.font.size = Pt(9.5)
        c[0].width = Inches(w0); c[1].width = Inches(w1)
    doc.add_paragraph(); return tbl

# ---------- Title ----------
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Measured mobility and the urban acoustic environment during graded COVID-19 distancing in Seoul, 2020-2023")
r.bold = True; r.font.size = Pt(15)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run(f"Working manuscript v2 ({ts[:8]}). Body in Korean, headings/Abstract in English (author workflow).")
rs.italic = True; rs.font.size = Pt(9)
au = doc.add_paragraph(); au.alignment = WD_ALIGN_PARAGRAPH.CENTER
au.add_run("Hyun In Jo, et al. [공저자/소속 추후 확정]   *Corresponding author").font.size = Pt(9)
doc.add_paragraph()

# ---------- Abstract ----------
H("Abstract", 1)
P("코로나-19 시기 서울의 단계적 사회적 거리두기를 자연실험으로 삼아, 고밀도 IoT 도시데이터 센서망(S-DoT, 1,123개 가동센서)의 "
  "센서-내(within-sensor) 소음과 행정동 단위 실측 이동량(생활인구) 사이의 용량-반응(dose-response)을 추정했다(2020-04~2023-12, "
  "1,248,794 sensor-days). 센서 고정효과 회귀에서 주간 동(洞) 이동량과 주간 소음은 정(+)의 용량-반응을 보였다(β≈+1.13 dB/log-unit, p<10⁻⁷). "
  "그러나 다년 절대레벨에는 센서 노후 드리프트·배치시점·위치별 교통회복이 교란되어, 센서+날짜 양방향 고정효과로 이를 모두 제거한 "
  "엄밀 식별에서는 효과가 작아진다(주간 β=+0.65 dB/log-unit, p=0.008; 이동량 30% 감소 ≈ −0.23 dB). 효과는 토지이용(상업/주거)에 "
  "무관하게 균일했고(상호작용 비유의), 도시 전체 공통 이동량 성분은 드리프트와 교란되어 식별 불가했다. "
  "결과는 '봉쇄→수 dB 급감'식 기존 보고와 달리, 측정 이동량에 대한 도시소음의 반응이 통계적으로 견고하나 작음(graded·modest)을 보이며, "
  "저가 IoT 소음망으로 코로나 소음효과를 측정할 때의 드리프트·해상도 한계와 cross-sectional 식별의 필요성을 함께 제시한다.")
P("Keywords: urban noise; environmental acoustics; COVID-19 mobility; dose-response; IoT sensor network; sensor drift; Seoul; natural experiment",
  italic=True, size=9.5)

# ---------- 1. Introduction ----------
H("1. Introduction", 1)
P("[서론 작성 예정] 노벨티 4축: (1) 한국 단계적 거리두기 = graded dose-response(서구 binary lockdown과 차별) "
  "(2) 측정 이동량–소음 정량 결합(선행 미개척: 대부분 '교통 줄어서'식 서술적 귀인) (3) 고밀도 IoT 센서망(~1,100점, 선행 대부분 4~70점) "
  "(4) 미출판 지리(서울/동아시아 고밀도 도시). 본 연구의 추가 기여 = 저가 IoT 망의 비검교정·드리프트·해상도 한계를 정면으로 다루어, "
  "절대레벨 시계열이 아닌 cross-dong within-date 식별만이 유효함을 보이고 효과크기를 정직하게 보고함.")

# ---------- 2. Data ----------
H("2. Data and Materials", 1)
P("모든 원시데이터는 공개 공공데이터이며 다운로드 메커니즘을 명시한다. 취득·검증 시각 2026-06-21 KST. "
  "산출 분석패널 = 1,248,794 sensor-days (센서 1,123 · 행정동 421 · 2020-04-01~2023-12-31).")

H("2.1 Outcome — Urban noise (S-DoT)", 2)
prov([
 ("데이터셋", "스마트서울 도시데이터 센서(S-DoT) 환경정보 | 서울 OA-15969 (공공데이터포털 미러 15061244)"),
 ("다운로드", "POST datafile.seoul.go.kr/bigfile/iot/inf/nio_download.do?useCache=false | infId=OA-15969&seq=<YEAR>&infSeq=3 (무인증)"),
 ("해상도", "시간별 dB · 점센서 약 1,100개(가동 1,123) · 서울 25개구 전수 · 2020-04-01~2023-12-31"),
 ("핵심 변수", "소음(dB). 2020-22 단일 '소음(dB)'; 2023 스키마변경(22→58컬럼)으로 '소음 평균(dB)' 사용(최대/최소 회피)"),
 ("측정 성격", "저가 IoT 광대역 dB; LAeq 라벨 없음 → within-sensor 상대변화로만 사용, 공식 LAeq 4점으로 교차검증(§4.2)"),
 ("처리", "cp949; 시각=등록일자; 시리얼→location.xlsx 보정시리얼 좌표조인; 시간→일 에너지평균(Leq24/주간06-21/야간22-05); QC ≥12h·20-95dB"),
 ("산출 일패널", "data/processed/sdot_daily_panel_2020-2023.csv (1,253,298 sensor-days)"),
])

H("2.2 Exposure — Mobility (생활인구, 주 dose)", 2)
prov([
 ("데이터셋", "행정동 단위 서울 생활인구(내국인) | 서울 OA-14991 | 일별×시간대×행정동 32컬럼 UTF-8-BOM"),
 ("다운로드", "POST nio_download.do | infId=OA-14991&seq=<N>&infSeq=3 | 반기ZIP seq 2219-2224(2020-22) + 월별 2301-2312(2023), 총 ~2.8GB(가공 후 삭제)"),
 ("집약", "총생활인구수=시간별 체류(stock) → 일 대표=시간평균(주간06-21/야간22-05/전일) → data/processed/livingpop_dong_daily_2020-2023.csv (424동×1461일=619,464행, 결측0)"),
 ("dose 정의", "within-dong 상대변화: lp_*_logrel = log(동 일별 생활인구 / 동 post-lift(2022-07~2023-12) 기준평균). 도시 총량은 보존되나 동간 재분포가 큼(상업동 주간 −20~30%, 주거동 +최대37%)"),
 ("보조 dose", "지하철 승하차(OA-12921, 시간축·도시) data/processed/subway_daily_seoul_2020-2023.csv; 거리두기 stringency(§2.4)"),
])

H("2.3 Sensor-to-dong spatial join", 2)
prov([
 ("문제", "location.xlsx 주소가 도로명주소(동명 없음) → 좌표 기반 공간조인 필요"),
 ("경계", "vuski/admdongkor 행정동 GeoJSON ver20220101 (adm_cd2 10자리; 조인키 = adm_cd2[:8] = 생활인구 행정동코드)"),
 ("방법", "순수 Python ray-casting point-in-polygon(MultiPolygon·홀·bbox 프리필터). 코드 드리프트(강북 6동·강동 상일분할) 이름기반 검증 크로스워크"),
 ("결과", "1,165/1,170 매칭(99.6%) · 주소구 vs 동의구 일치 99.6% → data/processed/sensor_dong_map.csv"),
])

H("2.4 Covariates", 2)
prov([
 ("거리두기(준연속)", "수도권/서울 국면 → 일별 (영업종료시각·허용모임인원·stringency 0-7) 재코딩. data/processed/distancing_daily_2020-2023.csv. ⚠️정밀 일자경계는 질병청 HWPX(공공데이터포털 15106451) 원본 대조 필요(open item)"),
 ("기상", "Open-Meteo ERA5 재분석 일자료(서울 종관108 좌표 37.57,126.97): 평균/최고/최저기온·강수·최대풍속. 무인증·재현가능. KMA ASOS 공식자료는 API키 보유 시 대체 가능. data/processed/weather_seoul_daily_2020-2023.csv"),
 ("달력", "요일·주말·한국 공휴일(대체/임시 포함)·월·계절"),
])

H("2.5 Validation reference", 2)
prov([
 ("도로교통 4점 LAeq", "서울 OA-15473 (시청·신사·신촌·성수) 시간별 LAeq 2021-2024 → 드리프트·스키마 교차검증(§4.2)"),
 ("환경소음 146점", "한국환경공단 data.go.kr 15065396 (분기 LAeq, 공간검증 — 스파이크 단계)"),
])

# ---------- 3. Methods ----------
H("3. Methods", 1)
P("결과변수는 센서 내 상대변화로만 사용한다(저가 IoT의 센서별 검교정 오프셋을 차분으로 상쇄). 1,248,794행에 대해 더미 대신 "
  "demeaning(within transformation) + 센서/동 clustered SE로 고정효과를 흡수한다.")
P("(M1) 센서 고정효과 회귀: Leq ~ mobility + 기상(기온·기온²·강수·풍속·강수일) + 요일·주말·공휴일 [센서 FE, 센서 clustered SE]. "
  "관측 가능한 시간교란을 통제하되 시간변이를 허용해 dose를 식별.")
P("(M2, 헤드라인) 센서+날짜 양방향 고정효과: Leq ~ mobility [센서 FE + 날짜 FE]. 날짜 공통요인(기상·요일·도시추세·센서드리프트·계절·"
  "도시 전체 거리두기) 전부 흡수 → '같은 날, 자기 baseline 대비 더 비워진 동의 센서가 더 조용해졌나'라는 순수 cross-dong 식별. "
  "도시 단일 dose(지하철·stringency)는 날짜 FE에 흡수되어 식별 불가(시간추세·드리프트와 교란).")
P("(M3) 이질성: dose × 상업도(동 주간/야간 인구비) 상호작용 및 토지이용 그룹별 M2. "
  "(M4) DiD event-study: 이동량 고영향(상업) vs 저영향(주거) 동 ΔLAeq 차이의 시간궤적(공통교란 차분상쇄). "
  "(M5) 검증: 공식 4점 LAeq 연추세 대조(드리프트·스키마 경계).")

# ---------- 4. Results ----------
H("4. Results", 1)

H("4.1 Dose-response: robust but modest", 2)
P("센서 FE(M1)에서 주간 동 이동량(lp_day_logrel)과 주간 소음은 정(+)의 용량-반응을 보였다: β=+1.130 dB/log-unit (SE 0.209, p=6.6×10⁻⁸). "
  "통제변수는 물리적으로 타당했다(기온 U자형, 강수 +0.045 dB/mm, 주말 −0.83 dB, 공휴일 −1.18 dB; 모두 p<0.001).")
P("그러나 센서+날짜 양방향 FE(M2, 헤드라인)에서 효과는 작아진다 — 주간 β=+0.648 dB/log-unit (SE 0.245, 95% CI +0.17~+1.13, p=0.008), "
  "전일 β=+0.628 (p=0.038). 즉 도시 공통 시간변이를 제거하면 순수 cross-dong dose-response는 modest하다: "
  "동 주간 이동량 30% 감소 ≈ 주간소음 −0.23 dB, 50% 감소 ≈ −0.45 dB.")
P("주간/야간 비교는 효과가 daytime 현상임을 보인다(Fig. day-night). 야간은 센서 FE에서 야간 이동량에 강하게 반응하는 듯 보이나(β=+2.28, p<0.001) "
  "양방향 FE에서 소멸한다(β=+0.50, 95% CI −0.73~+1.74, p=0.42) — 야간의 겉보기 연관은 시간교란(드리프트·계절·도시추세)이었다. "
  "주간-야간 gap(L_day−L_night; 같은 센서·같은 날 차분이라 오프셋·공통드리프트 상쇄) dose-response는 ≈0(β=+0.02, p=0.92)으로, "
  "효과가 day-night 차등 메커니즘이 아니라 주간 절대수준의 modest 감소임을 시사한다.")
P("도시 단일 dose는 신뢰 불가였다: 센서 FE만 둔 모형에서 지하철 상대승객은 부호가 반대(β=−2.0, 시간추세·드리프트와 교란), "
  "거리두기 stringency도 미소·역부호. 이는 도시 단일 시계열 dose의 내생성·교란을 보인 것으로, cross-dong 식별의 필요성을 뒷받침한다.")

H("4.2 Sensor drift and validation (핵심 진단)", 2)
P("원시 within-sensor ΔLAeq(post-lift 기준)는 역설을 보였다 — 최강제한기 ΔL_day가 정상기보다 높음(+0.85 vs +0.34 dB). "
  "진단 결과 이는 코로나 효과가 아니라 다년 절대레벨 교란이었다: (a) 4년 전수 가동 842센서의 연평균 Leq24가 "
  "49.4(2020)→47.6→47.2→47.1 dB(2023)로 단조감소(평균·중앙값 모두; 85% 센서 하락추세). "
  "(b) 2022↔2023 스키마 경계는 깨끗함(2022Q4 47.06 vs 2023Q1 47.08, +0.02 dB) — '소음 평균(dB)' 교정 성공. "
  "(c) 그러나 이 감소는 균일한 하드웨어 드리프트가 아니다(Fig. drift): 검증된 도로 4점에서 근접 S-DoT 추세(2021→23)는 "
  "City Hall +1.7·Sinsa +2.5·Sinchon ≈0·Seongsu −0.3 dB로 하락하지 않고 공식 LAeq 추세(각 +7.2[공사 추정 outlier]·+0.4·−0.4·+1.5)와 같은 방향이다 — "
  "즉 도로변 S-DoT는 신뢰할 만하나, 도시 전체 평균의 하락(−0.5 dB/2021→23)은 비도로·정온지역에 집중된 위치별 현상(배치시점·국지변화·선택적 드리프트 혼재)으로 "
  "다년 절대비교를 신뢰할 수 없게 한다. → 절대/시간축 비교는 신뢰 불가, M2의 cross-dong within-date 식별만 유효(모든 날짜공통 교란에 면역).")

H("4.3 Heterogeneity (효과는 균일·modest)", 2)
P("토지이용별로 효과가 다를 것이라는 가설은 기각되었다 — 상업 동(주간/야간 인구비 상위1/3) β=+0.335, 주거 동(하위1/3) β=+0.847, "
  "혼합 동 β=+0.849(개별 비유의), dose×상업도 상호작용 β=−0.056 (p=0.49, 비유의). 동×주 집계에서도 β=0.67로 안정 — "
  "효과크기가 측정오차로 감쇠된 것이 아니라(결과변수 측정오차는 계수를 편의시키지 않음) 진정 modest하며 토지이용에 균일함을 시사한다.")

H("4.4 Spatial pattern and dynamics", 2)
P("동 단위 장기 공간패턴은 robust하게 null이었다. 단일센서 동(그 센서의 드리프트가 곧 동값)을 제외하고(센서 ≥2) "
  "아웃라이어-내성 통계로 보면 (이동량 감소율 vs 도시평균 대비 상대 소음변화) Spearman ρ ≈ −0.05(전체)·+0.08(상업)·−0.10(주거)로 "
  "사실상 0이다. Pearson r은 소수 극단 동(다수가 단일센서·국지 공사 등)에 끌려 +0.1~+0.4로 부풀려지므로 신뢰하지 않는다 — "
  "Theil-Sen 회귀 기울기도 ≈0. 즉 동별 장기평균 소음변화는 국지 요인(공사·도로·센서별 드리프트)에 묻혀 깨끗한 공간 gradient를 내지 않는다(Fig. 변화지도·토지이용별). "
  "반면 DiD event-study(고영향 vs 저영향 동)에서는 주별 (이동량 차이 vs 소음 차이) 상관 r=+0.44로 방향이 일관 — "
  "제한기에 고영향(상업)동이 저영향(주거)동보다 상대적으로 조용했다(차이 −0.07 dB). "
  "핵심 함의: 이동량-소음 신호는 고빈도 cross-dong/within-date 변이(양방향 FE·DiD)에만 robust하게 존재하고, "
  "naive 동 단위 공간 cross-section으로는 검출되지 않는다 — 후속 IoT 소음연구에 대한 방법론적 경고.")

# ---------- 5. Discussion ----------
H("5. Discussion", 1)
P("측정 이동량에 대한 도시소음의 용량-반응은 통계적으로 견고하나 작다(주간 ~0.2-0.3 dB / 이동량 30% 감소). "
  "이는 '봉쇄→평균 5 dB 감소'식 기존 보고와 대비된다 — 그 보고들은 특정 번화가의 절대레벨을 팬데믹 전후로 단순비교한 반면, "
  "본 연구는 센서 내·이동량 graded·전기간 회귀로 한계효과를 엄밀히 추정했다. 더욱이 본 추정치는 하한이다: 양방향 FE는 cross-dong 성분만 "
  "식별하고 도시 전체 공통 이동량 감소 성분은 센서 드리프트와 교란되어 깨끗이 추정할 수 없다. 도시소음이 통과교통 등 "
  "비탄력적 배경원에 의해 '끈적'하다는 점, S-DoT 1 dB 정수해상도가 SE를 넓힌다는 점도 modest 효과와 정합한다.")
P("방법론적 기여: 저가 고밀도 IoT 소음망으로 코로나 소음효과를 측정하려는 시도는 (i) 센서별 검교정 오프셋 → within-sensor, "
  "(ii) 다년 드리프트·배치시점 → 절대/시간축 비교 불가·cross-dong 식별 필수, (iii) 도시 단일 dose의 시간교란 → 동 단위 측정 dose 필요, "
  "를 정면으로 보여준다. 후속 IoT 기반 환경소음 연구의 설계 지침이 된다.")

# ---------- 6. Limitations ----------
H("6. Limitations", 1)
P("(1) S-DoT 비검교정·1 dB 해상도·노후 드리프트 → 절대/시간축 신뢰 낮음(상대변화·cross-dong로 우회). "
  "(2) 생활인구는 체류(presence)로 소음생성 활동(교통)의 불완전 proxy. (3) 거리두기 준연속 코딩의 정밀 일자경계는 HWPX 원본 대조 필요(open). "
  "(4) 기상은 ERA5 재분석(KMA 공식 ASOS 대체 가능). (5) 도시 공통 이동량 효과는 드리프트 교란으로 미식별 → 추정치는 하한. "
  "(6) 4점 검증소가 도로변이라 S-DoT(도로에서 후퇴)와 절대레벨 비교 불가, 추세만 대조.")

# ---------- Data Availability ----------
H("Data Availability Statement", 1)
P("모든 원시데이터는 공개 공공데이터이다. S-DoT 소음(서울 OA-15969), 생활인구 행정동(OA-14991)·지하철(OA-12921), "
  "거리두기 시행연혁(질병청/공공데이터포털 15106451), 기상(Open-Meteo ERA5; 또는 기상청 ASOS), 행정동 경계(vuski/admdongkor, ver20220101), "
  "검증 4점 LAeq(서울 OA-15473)·환경소음(환경공단 15065396). "
  "라이선스: 생활인구·S-DoT 등 서울 열린데이터 = 공공누리 제1유형(출처표시) 다수; "
  "지하철 승하차(OA-12921) = 공공누리 제3유형(출처표시+변경금지)이므로 가공 결과 재배포 시 원자료 미변경 고지 필요. "
  "행정동 경계 vuski/admdongkor = 원저장소 라이선스 준수. 분석 코드/처리 스크립트 = [저장소 TBD].")

H("References", 1); P("[Vancouver 번호식 — 작성 예정]")

doc.save(OUT)
print("saved:", OUT)
print("paragraphs:", len(doc.paragraphs), "tables:", len(doc.tables))

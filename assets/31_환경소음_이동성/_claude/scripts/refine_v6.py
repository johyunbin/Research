# -*- coding: utf-8 -*-
# v8 정제6: 검교정 환경소음망(noiseinfo) 검증 통합 — §3.5 강화 + 새 Fig10 + Methods M5 + Data Availability + 초록.
import os, re, struct
from datetime import datetime, timezone, timedelta
from docx import Document
from docx.shared import Pt, Emu, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = r"T:\00_BACKUP\01_한양대학교\00_연구실\00_프로젝트\★_논문화 프로젝트\31_환경소음_이동성"
FIGD = os.path.join(ROOT, "_claude", "figures")
SRC = os.path.join(ROOT, "01_논문작업", "Manuscript_v8_20260622_015928.docx")
ts = datetime.now(timezone(timedelta(hours=9))).strftime("%Y%m%d_%H%M%S")
OUT = os.path.join(ROOT, "01_논문작업", f"Manuscript_v8_{ts}.docx")

def setfont(run, size=12, bold=False, sub=False):
    run.font.size = Pt(size); run.bold = bold; run.font.name = 'Times New Roman'
    if sub: run.font.subscript = True
    rpr = run._element.get_or_add_rPr(); rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'): rf.set(qn(a), 'Times New Roman')
    rf.set(qn('w:eastAsia'), '바탕')

SUB = [('Leq,24h', ('L', 'eq,24h')), ('LAeq', ('L', 'Aeq')), ('Lday', ('L', 'day')),
       ('Lnight', ('L', 'night')), ('Leq', ('L', 'eq'))]
SRE = re.compile('(' + '|'.join(re.escape(k) for k, _ in SUB) + ')'); SMAP = dict(SUB)

def emit(p, text, size=12, bold=False):
    for part in SRE.split(text):
        if not part: continue
        if part in SMAP:
            b, sb = SMAP[part]; setfont(p.add_run(b), size, bold); setfont(p.add_run(sb), size, bold, sub=True)
        else:
            setfont(p.add_run(part), size, bold)

def replace_para(p, text, size=12, bold=False):
    for r in list(p.runs): r._element.getparent().remove(r._element)
    emit(p, text, size, bold)

doc = Document(SRC)
def find(marker):
    hits = [p for p in doc.paragraphs if p.text.strip().startswith(marker)]
    assert len(hits) == 1, f"마커 {len(hits)}건: {marker[:34]}"
    return hits[0]

# ===== 초록: LAeq-stations 문장을 검교정망 결과로 교체 + 길이 보정 =====
ABS = ("How strongly the urban acoustic environment responds to human mobility is a first-order question for "
       "mobility, land-use and noise policy, yet it remains difficult to answer because mobility rarely varies "
       "exogenously and is entangled with weather, land use and time of day. We exploit Korea's graded COVID-19 "
       "social-distancing system, which repeatedly tightened and relaxed activity over 2020-2023, as a natural "
       "experiment, paired with a city-scale low-cost IoT noise network, to estimate a measured-mobility dose-response "
       "for urban noise. We assembled 1,248,794 sensor-days from 1,123 Smart Seoul Data-of-Things (S-DoT) sensors "
       "across 421 administrative neighbourhoods and matched each sensor-day to neighbourhood daytime de-facto "
       "population. Because low-cost sensors carry calibration offsets and multi-year drift, we identify the effect "
       "only from within-sensor variation and from cross-neighbourhood variation within each date (two-way fixed "
       "effects). Daytime noise rises and falls with daytime mobility (beta = +0.65 dB per log-unit; a 30% mobility "
       "reduction corresponds to about 0.23 dB), robust to placebo permutation (p = 0.003) and consistent across "
       "seasons but specific to daytime: a nighttime association vanishes under two-way fixed effects. No robust "
       "gradient survives in the long-run spatial cross-section, and a calibrated environmental-noise network shows "
       "that S-DoT under-reads standard LAeq by about 12 dB and that its multi-year decline is a sensor-drift "
       "artifact. The response is statistically robust but modest: mobility-demand management alone cannot deliver "
       "large noise reductions and must be paired with source- and propagation-stage measures. It also corrects "
       "binary-lockdown over-estimation and offers a transferable, drift-aware design guide for IoT-based noise "
       "monitoring in smart cities.")
replace_para(find("How strongly the urban acoustic"), ABS, 12)
print("초록 단어수:", len(ABS.split()))

# ===== §2.6 M5 갱신(전체 보조분석 문단 재작성) =====
replace_para(find("이 두 모형 위에 보조 분석을 더했다"),
 "이 두 모형 위에 보조 분석을 더했다. (M3) 기능 분할: M2를 결과변수(주간·야간·주야 차이), 토지이용(동의 주간 대비 야간 인구비를 3분위로 나눈 상업·혼합·주거), 평일/주말, 계절별로 따로 추정해 효과가 어디서 나타나는지 본다. (M4) 차이의 차이(difference-in-differences): 이동량이 크게 줄어든 동(주로 상업지구)과 거의 줄지 않은 동(주로 주거지구)을 나눠, 두 그룹의 소음 변화를 매주 비교한다. 두 그룹은 같은 도시·같은 시기를 공유하므로 날씨·계절·센서 표류 같은 공통 요인은 양쪽에 똑같이 작용해 그 차이를 보면 서로 지워지고, 순수하게 이동량 차이에서 비롯된 소음 차이만 남는다. (M5) 검증: S-DoT의 절대레벨과 다년 시계열의 신뢰성을 검교정된 공식 환경소음 측정망(국가소음정보시스템의 자동·수동 측정망)과 도로교통 4점 상시측정소(표준 LAeq)에 대조해 점검한다. 측정점 인근 S-DoT와의 레벨 차(offset) 및 2020–2023 연추세를 비교해 센서 표류와 2023년 자료구조 변경의 영향을 진단한다. (M6) 강건성·위약 검정: 식별이 기계적 산물이 아님을 확인하기 위해 ① 같은 날짜 안에서 이동량 dose를 동들 사이에 무작위로 재배치한 placebo 순열검정(300회), ② 이동량이 인과적으로 만들 수 없는 결과변수(기온)에 대한 위약 회귀, ③ 이동량 2차항을 넣은 비선형성 점검을 수행했고, 분할추정(M3)의 다중비교는 Benjamini–Hochberg FDR로 보정했다(§3.6). 끝으로 동 단위 공간 분석에서는 소수의 극단적인 동에 결과가 휘둘리지 않도록 일반 회귀·상관 대신 극단값에 강한(robust) 방법(Theil-Sen 회귀와 Spearman 순위상관)을 쓰고, 추정이 불안정한 단일 센서 동을 제외해 센서가 2개 이상인 동만 사용했다. 모든 분석은 Python 3.13(pandas·NumPy·SciPy·statsmodels)으로 수행했으며, 양방향 고정효과는 센서·날짜 평균을 번갈아 차감하는 반복 within-변환으로, 극단값에 강한 추정은 SciPy의 Theil-Sen·Spearman으로 계산했다.", 12)

# ===== §3.5: para A(진단) 재작성 + para B(검교정망 검증) 삽입 =====
pA = find("각 센서의 자기 대비 변화(ΔLday")
replace_para(pA,
 "각 센서의 자기 대비 변화(ΔLday, 정상기 기준)를 그대로 그려 보면 언뜻 모순처럼 보이는 결과가 나온다. 규제가 가장 강했던 시기의 주간 소음이 오히려 정상기보다 높게 나오는 것이다(+0.85 vs +0.34 dB). 진단해 보니 이는 코로나 효과가 아니라 여러 해에 걸친 절대 수준의 변동(표류) 때문이었다(Fig. 10a). 4년 내내 가동된 842개 센서의 연평균 Leq,24h는 49.4(2020)→47.6→47.2→47.1 dB(2023)로 해마다 꾸준히 낮아졌고(85% 센서가 하락 추세), 이 하락이 2020-21년을 상대적으로 '시끄럽게' 보이게 만든 것이다(2023년 자료구조 변경 지점에서는 수준 도약이 없었다, +0.02 dB).", 12)
pB = doc.add_paragraph(); pA._p.addnext(pB._p)
emit(pB,
 "이 하락이 실제 소음 변화가 아니라 센서 드리프트임을, 검교정된 공식 환경소음 측정망(국가소음정보시스템)과 대조해 확인했다. 먼저 절대레벨이 크게 어긋난다. 검교정망 인근의 S-DoT는 표준 LAeq보다 주간 평균 11.7 dB, 도로변에서는 약 16 dB 낮게 읽혀(Fig. 10b), S-DoT 절대값을 표준 소음도로 해석할 수 없음을 정량적으로 보여 준다. 반면 검교정망 자체의 다년 추세는 안정적이거나 오히려 상승한다(Fig. 10a). 일별 자동측정망(도로 9점)은 2020→2023에 +0.04 dB로 사실상 일정했고, 분기 수동측정망은 주거·일반지역 91점이 +1.93 dB, 도로 60점이 +1.91 dB 상승했다(별도의 도로교통 4점 상시 LAeq도 같은 방향이다). 절대레벨의 치우침은 추세 차분에서 상쇄되므로, 검교정망이 안정·상승하는 동안 S-DoT만 2.3 dB 하락했다는 사실은 그 하락이 실제 환경 변화가 아니라 센서의 하향 드리프트임을 뜻한다. 특히 S-DoT 하락이 집중된 비(非)도로 정온지역에서 검교정 주거망이 오히려 상승했다는 점은 그 하락이 실재가 아님을 직접 보여 준다. 결국 S-DoT의 절대 수준과 다년 시계열은 그대로 믿을 수 없으며, 이것이 우리가 절대·시계열 비교 대신 센서 내 상대변화와 '같은 날 동 사이의 차이'(M2)에 의존하는 이유다.", 12)

# ===== Fig 10 캡션 =====
replace_para(find("Fig. 10."),
 "Fig. 10. Drift diagnosis and calibrated validation. (a) Annual noise level relative to 2020: the S-DoT network "
 "falls about 2.3 dB, whereas the calibrated environmental-noise monitoring network (automatic daily and manual "
 "quarterly stations, road and residential) stays flat or rises, identifying the S-DoT decline as a sensor-drift "
 "artifact. (b) Nearby S-DoT reads about 12 dB below the calibrated LAeq (more at roadside), so absolute S-DoT "
 "levels cannot be interpreted as standard noise levels.", 11)

# ===== Data Availability: noiseinfo 추가 =====
replace_para(find("모든 원시데이터는 공개 공공데이터이다"),
 "모든 원시데이터는 공개 공공데이터이다(S-DoT 소음 OA-15969, 생활인구 OA-14991, 지하철 OA-12921, 거리두기 시행연혁 공공데이터포털 15106451, 기상 Open-Meteo ERA5, 행정동 경계 vuski/admdongkor, 검증 4점 도로교통 LAeq OA-15473, 검교정 환경소음 자동·수동 측정망 국가소음정보시스템 noiseinfo.or.kr). 지하철 데이터는 공공누리 제3유형(출처표시+변경금지). 분석 코드는 [저장소 TBD]에 공개 예정이다.", 12)

# ===== 새 Fig 10 재임베드 (index 9) =====
blob = open(os.path.join(FIGD, "fig_drift_validation.png"), "rb").read()
wpx, hpx = struct.unpack(">II", blob[16:24])
sh = doc.inline_shapes[9]
rId = sh._inline.graphic.graphicData.pic.blipFill.blip.embed
doc.part.related_parts[rId]._blob = blob
sh.height = Emu(int(round(int(sh.width) * hpx / wpx)))
print(f"Fig 10 재임베드: {sh.width.inches:.2f} x {sh.height.inches:.2f}")

doc.save(OUT)
print("\n저장:", OUT)
